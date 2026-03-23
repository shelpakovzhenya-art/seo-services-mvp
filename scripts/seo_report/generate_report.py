import argparse
import json
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

BRAND_NAME = "Shelpakov Digital"
COLOR_INK = RGBColor(23, 32, 51)
COLOR_MUTED = RGBColor(84, 96, 115)
COLOR_WARM = RGBColor(240, 143, 73)
COLOR_LINE = RGBColor(217, 227, 239)

PDF_COLOR_INK = colors.HexColor("#172033")
PDF_COLOR_MUTED = colors.HexColor("#546073")
PDF_COLOR_WARM = colors.HexColor("#f08f49")
PDF_COLOR_LINE = colors.HexColor("#d9e3ef")
PDF_COLOR_BG = colors.HexColor("#f7fafd")


def load_report(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def visible_blocks(report: dict) -> list[str]:
    hidden = set(report.get("hiddenBlocks") or [])
    return [item for item in report.get("blockOrder") or [] if item not in hidden]


def add_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    table_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E3EF")
        borders.append(element)
    table_pr.append(borders)


def add_cover(doc: Document, report: dict) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    chip = doc.add_paragraph()
    chip.alignment = WD_ALIGN_PARAGRAPH.LEFT
    chip_run = chip.add_run(BRAND_NAME)
    chip_run.bold = True
    chip_run.font.size = Pt(10)
    chip_run.font.color.rgb = COLOR_WARM

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(report.get("title", "SEO-отчет за месяц"))
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = COLOR_INK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle.add_run(report.get("subtitle", ""))
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = COLOR_MUTED

    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = True
    set_table_borders(meta_table)
    values = [
        ("Проект", report.get("projectName", "—")),
        ("Отчетный период", f"{report.get('periodStart', '—')} — {report.get('periodEnd', '—')}"),
        ("Сравнение", f"{report.get('comparePeriodStart', '—')} — {report.get('comparePeriodEnd', '—')}"),
    ]

    for idx, (label, value) in enumerate(values):
        cell = meta_table.rows[0].cells[idx]
        add_cell_shading(cell, "F7FAFD")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        label_run = paragraph.add_run(f"{label}\n")
        label_run.bold = True
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = COLOR_MUTED
        value_run = paragraph.add_run(str(value))
        value_run.font.size = Pt(11.5)
        value_run.bold = True
        value_run.font.color.rgb = COLOR_INK

    doc.add_paragraph()


def add_heading(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = COLOR_INK


def add_empty_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_MUTED


def add_html_content_doc(doc: Document, html: str) -> None:
    soup = BeautifulSoup(html or "", "html.parser")
    blocks = soup.find_all(["p", "li"])

    if not blocks:
        add_empty_note(doc, "Блок пока пуст.")
        return

    for node in blocks:
        paragraph = doc.add_paragraph(style=None)
        paragraph.paragraph_format.space_after = Pt(6)
        if node.name == "li":
            paragraph.style = "List Bullet"
        run = paragraph.add_run(node.get_text(" ", strip=True))
        run.font.size = Pt(10.8)
        run.font.color.rgb = COLOR_MUTED


def add_metric_table_doc(doc: Document, title: str, rows: list[dict], empty_text: str, include_source: bool = False) -> None:
    add_heading(doc, title)

    if not rows:
        add_empty_note(doc, empty_text)
        return

    columns = 5 if include_source else 4
    table = doc.add_table(rows=1, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    headers = ["Метрика"]
    if include_source:
        headers.append("Источник")
    headers.extend(["Прошлый период", "Текущий период", "Изменение"])

    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        add_cell_shading(cell, "F7FAFD")
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_MUTED

    for row in rows:
        cells = table.add_row().cells
        values = [row.get("label", "—")]
        if include_source:
            values.append(row.get("sourceLabel", "—"))
        values.extend(
            [
                row.get("previousRaw", "—"),
                row.get("currentRaw", "—"),
                row.get("changeRaw", "—"),
            ]
        )

        for idx, value in enumerate(values):
            paragraph = cells[idx].paragraphs[0]
            run = paragraph.add_run(str(value))
            run.font.size = Pt(10)
            if idx == len(values) - 1:
                tone = row.get("changeTone")
                if tone == "positive":
                    run.font.color.rgb = RGBColor(4, 120, 87)
                    run.bold = True
                elif tone == "negative":
                    run.font.color.rgb = RGBColor(185, 28, 28)
                    run.bold = True
                else:
                    run.font.color.rgb = COLOR_MUTED
            else:
                run.font.color.rgb = COLOR_INK if idx == 0 else COLOR_MUTED

        if row.get("notes"):
            note = cells[0].add_paragraph()
            note_run = note.add_run(str(row["notes"]))
            note_run.font.size = Pt(8.5)
            note_run.font.color.rgb = COLOR_MUTED


def build_metric_line(row: dict) -> str:
    label = row.get("label", "Метрика")
    previous = str(row.get("previousRaw", "—")).strip()
    current = str(row.get("currentRaw", "—")).strip()
    change = str(row.get("changeRaw", "—")).strip()
    note = " Требует проверки." if row.get("requiresReview") else ""

    if previous and previous != "—" and current and current != "—" and previous != current:
        suffix = f" ({change})" if change and change != "—" else ""
        return f"{label}: {previous} → {current}{suffix}.{note}".strip()

    if current and current != "—":
        return f"{label}: {current}.{note}".strip()

    return f"{label}: нет данных.{note}".strip()


def group_search_rows(rows: list[dict]) -> list[tuple[str, list[dict]]]:
    grouped: dict[str, list[dict]] = {}
    for row in rows:
        label = row.get("sourceLabel") or "Прочий источник"
        grouped.setdefault(label, []).append(row)

    def sort_key(item: tuple[str, list[dict]]):
        label = item[0]
        if label == "Google":
            return (0, label)
        if label == "Яндекс":
            return (1, label)
        return (2, label)

    metric_order = ["visits", "users", "bounce_rate", "page_depth", "time_on_site"]

    sections: list[tuple[str, list[dict]]] = []
    for label, section_rows in sorted(grouped.items(), key=sort_key):
        ordered = sorted(section_rows, key=lambda row: metric_order.index(row.get("key")) if row.get("key") in metric_order else 99)
        sections.append((label, ordered))
    return sections


def add_search_systems_doc(doc: Document, rows: list[dict], empty_text: str) -> None:
    add_heading(doc, "3. Разрез по поисковым системам")

    if not rows:
        add_empty_note(doc, empty_text)
        return

    for label, section_rows in group_search_rows(rows):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(label)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_INK

        for row in section_rows:
            item = doc.add_paragraph(style="List Bullet")
            item.paragraph_format.space_after = Pt(4)
            run = item.add_run(build_metric_line(row))
            run.font.size = Pt(10.5)
            run.font.color.rgb = COLOR_MUTED


def add_text_block_doc(doc: Document, title: str, html: str) -> None:
    add_heading(doc, title)
    add_html_content_doc(doc, html)


def generate_docx(report: dict, output_path: Path) -> None:
    doc = Document()
    add_cover(doc, report)

    if report.get("parserWarnings") or report.get("missingBlocks"):
        add_heading(doc, "Проверка данных")
        if report.get("missingBlocks"):
            add_empty_note(doc, "Автоматически не заполнены блоки: " + ", ".join(report["missingBlocks"]))
        for item in report.get("parserWarnings") or []:
            add_empty_note(doc, f"• {item}")

    for block_id in visible_blocks(report):
        if block_id == "overview":
            add_text_block_doc(doc, report["textBlocks"]["overview"]["title"], report["textBlocks"]["overview"]["content"])
        elif block_id == "key_metrics":
            add_metric_table_doc(
                doc,
                "2. Ключевые показатели поискового трафика",
                report.get("keyMetrics") or [],
                "Ключевые показатели пока не распознаны.",
            )
        elif block_id == "search_systems":
            add_search_systems_doc(
                doc,
                report.get("searchSystemRows") or [],
                "По Google и Яндексу пока нет подтвержденных значений.",
            )
        elif block_id == "seo_highlights":
            add_metric_table_doc(
                doc,
                "Дополнительные SEO-показатели",
                report.get("seoHighlightRows") or [],
                "Дополнительные SEO-показатели пока не заполнены.",
                include_source=True,
            )
        elif block_id == "analytical_conclusions":
            add_text_block_doc(
                doc,
                report["textBlocks"]["analyticalConclusions"]["title"],
                report["textBlocks"]["analyticalConclusions"]["content"],
            )
        elif block_id == "final_summary":
            add_text_block_doc(doc, report["textBlocks"]["finalSummary"]["title"], report["textBlocks"]["finalSummary"]["content"])
        elif block_id == "recommendations":
            add_text_block_doc(doc, report["textBlocks"]["recommendations"]["title"], report["textBlocks"]["recommendations"]["content"])

    doc.save(str(output_path))


def build_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ReportTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=PDF_COLOR_INK,
            alignment=TA_LEFT,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionTitle",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=19,
            textColor=PDF_COLOR_INK,
            spaceBefore=8,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyTextRu",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14.5,
            textColor=PDF_COLOR_MUTED,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="MetaLabel",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.5,
            leading=11,
            textColor=PDF_COLOR_MUTED,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Subheading",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=PDF_COLOR_INK,
            spaceBefore=6,
            spaceAfter=4,
        )
    )
    return styles


def html_to_pdf_blocks(html: str, styles) -> list:
    soup = BeautifulSoup(html or "", "html.parser")
    nodes = soup.find_all(["p", "li"])
    if not nodes:
        return [Paragraph("Блок пока пуст.", styles["BodyTextRu"])]

    blocks = []
    for node in nodes:
        prefix = "• " if node.name == "li" else ""
        blocks.append(Paragraph(prefix + node.get_text(" ", strip=True), styles["BodyTextRu"]))
    return blocks


def build_metric_table_pdf(rows: list[dict], empty_text: str, include_source: bool, styles) -> list:
    if not rows:
        return [Paragraph(empty_text, styles["BodyTextRu"])]

    header = ["Метрика"]
    widths = [58 * mm]
    if include_source:
        header.append("Источник")
        widths.append(30 * mm)
    header.extend(["Прошлый период", "Текущий период", "Изменение"])
    widths.extend([30 * mm, 30 * mm, 26 * mm])

    data = [header]
    for row in rows:
        values = [row.get("label", "—")]
        if include_source:
          values.append(row.get("sourceLabel", "—"))
        values.extend(
            [
                row.get("previousRaw", "—"),
                row.get("currentRaw", "—"),
                row.get("changeRaw", "—"),
            ]
        )
        data.append(values)

    table = Table(data, repeatRows=1, colWidths=widths)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), PDF_COLOR_BG),
                ("TEXTCOLOR", (0, 0), (-1, 0), PDF_COLOR_MUTED),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.7),
                ("LEADING", (0, 0), (-1, -1), 11),
                ("GRID", (0, 0), (-1, -1), 0.5, PDF_COLOR_LINE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    return [table]


def build_search_systems_pdf(rows: list[dict], empty_text: str, styles) -> list:
    if not rows:
        return [Paragraph(empty_text, styles["BodyTextRu"])]

    blocks = []
    for label, section_rows in group_search_rows(rows):
        blocks.append(Paragraph(label, styles["Subheading"]))
        for row in section_rows:
            blocks.append(Paragraph(f"• {build_metric_line(row)}", styles["BodyTextRu"]))
    return blocks


def generate_pdf(report: dict, output_path: Path) -> None:
    styles = build_pdf_styles()
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )

    story = [
        Paragraph(BRAND_NAME, styles["MetaLabel"]),
        Spacer(1, 4),
        Paragraph(report.get("title", "SEO-отчет за месяц"), styles["ReportTitle"]),
        Paragraph(report.get("subtitle", ""), styles["BodyTextRu"]),
        Spacer(1, 8),
        Paragraph(f"<b>Проект:</b> {report.get('projectName', '—')}", styles["BodyTextRu"]),
        Paragraph(
            f"<b>Отчетный период:</b> {report.get('periodStart', '—')} — {report.get('periodEnd', '—')}",
            styles["BodyTextRu"],
        ),
        Paragraph(
            f"<b>Сравнение:</b> {report.get('comparePeriodStart', '—')} — {report.get('comparePeriodEnd', '—')}",
            styles["BodyTextRu"],
        ),
        Spacer(1, 8),
    ]

    if report.get("parserWarnings") or report.get("missingBlocks"):
        story.append(Paragraph("Проверка данных", styles["SectionTitle"]))
        if report.get("missingBlocks"):
            story.append(
                Paragraph(
                    "Автоматически не заполнены блоки: " + ", ".join(report["missingBlocks"]),
                    styles["BodyTextRu"],
                )
            )
        for item in report.get("parserWarnings") or []:
            story.append(Paragraph(f"• {item}", styles["BodyTextRu"]))
        story.append(Spacer(1, 6))

    for block_id in visible_blocks(report):
        if block_id == "overview":
            story.append(Paragraph(report["textBlocks"]["overview"]["title"], styles["SectionTitle"]))
            story.extend(html_to_pdf_blocks(report["textBlocks"]["overview"]["content"], styles))
        elif block_id == "key_metrics":
            story.append(Paragraph("2. Ключевые показатели поискового трафика", styles["SectionTitle"]))
            story.extend(build_metric_table_pdf(report.get("keyMetrics") or [], "Ключевые показатели пока не распознаны.", False, styles))
        elif block_id == "search_systems":
            story.append(Paragraph("3. Разрез по поисковым системам", styles["SectionTitle"]))
            story.extend(build_search_systems_pdf(report.get("searchSystemRows") or [], "По Google и Яндексу пока нет подтвержденных значений.", styles))
        elif block_id == "seo_highlights":
            story.append(Paragraph("Дополнительные SEO-показатели", styles["SectionTitle"]))
            story.extend(build_metric_table_pdf(report.get("seoHighlightRows") or [], "Дополнительные SEO-показатели пока не заполнены.", True, styles))
        elif block_id == "analytical_conclusions":
            story.append(Paragraph(report["textBlocks"]["analyticalConclusions"]["title"], styles["SectionTitle"]))
            story.extend(html_to_pdf_blocks(report["textBlocks"]["analyticalConclusions"]["content"], styles))
        elif block_id == "final_summary":
            story.append(Paragraph(report["textBlocks"]["finalSummary"]["title"], styles["SectionTitle"]))
            story.extend(html_to_pdf_blocks(report["textBlocks"]["finalSummary"]["content"], styles))
        elif block_id == "recommendations":
            story.append(Paragraph(report["textBlocks"]["recommendations"]["title"], styles["SectionTitle"]))
            story.extend(html_to_pdf_blocks(report["textBlocks"]["recommendations"]["content"], styles))
        story.append(Spacer(1, 8))

    doc.build(story)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate monthly SEO report exports.")
    parser.add_argument("--input", required=True, help="Path to JSON input")
    parser.add_argument("--output-dir", required=True, help="Directory for output files")
    parser.add_argument("--base-name", required=True, help="Base name for docx/pdf")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    report = load_report(input_path)

    docx_path = output_dir / f"{args.base_name}.docx"
    pdf_path = output_dir / f"{args.base_name}.pdf"

    generate_docx(report, docx_path)
    generate_pdf(report, pdf_path)


if __name__ == "__main__":
    main()
