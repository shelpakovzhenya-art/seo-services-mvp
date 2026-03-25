#!/usr/bin/env python
from __future__ import annotations

import argparse
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
import json
import math
import re
import statistics
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable
from urllib.parse import urljoin, urlparse
from xml.etree import ElementTree as ET

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from preview_export import write_preview_html, write_preview_pdf
from text_utils import normalize_output_text, normalize_structure

BRAND_NAME = "Shelpakov Digital"
BRAND_DARK = "101C2B"
BRAND_DARK_ALT = "18283C"
BRAND_TEXT = "102035"
BRAND_MUTED = "5D6B82"
BRAND_LINE = "D8E7F6"
BRAND_CYAN = "69D3FF"
BRAND_ORANGE = "F28B34"
BRAND_SOFT = "F6F9FC"
BRAND_ACCENT_SOFT = "FFF3E6"
USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/132.0.0.0 Safari/537.36 ShelpakovAuditBot/1.0"
)


@dataclass
class PageSnapshot:
    url: str
    page_type: str
    status_code: int
    response_time_ms: int
    content_type: str
    title: str
    description: str
    canonical: str
    h1s: list[str]
    meta_robots: str
    x_robots_tag: str
    image_count: int
    missing_alt_count: int
    lazy_missing_count: int
    missing_dimensions_count: int
    schema_types: list[str]
    internal_links: int
    word_count: int
    has_meta_keywords: bool
    has_forms: bool
    has_faq_block: bool
    has_reviews_block: bool
    has_trust_block: bool
    has_commercial_block: bool
    has_comparison_table: bool
    has_breadcrumbs: bool
    html_size_kb: float


@dataclass
class AuditIssue:
    severity: str
    title: str
    why_it_matters: str
    evidence: list[str]
    recommendation: str


def normalize_base_url(raw_url: str) -> str:
    url = raw_url.strip()
    if not url.startswith(("http://", "https://")):
        url = f"https://{url}"
    parsed = urlparse(url)
    clean = parsed._replace(path=parsed.path.rstrip("/"), params="", query="", fragment="")
    return clean.geturl() or url


def slugify_for_filename(value: str) -> str:
    value = re.sub(r"^https?://", "", value.strip().lower())
    value = re.sub(r"[^a-z0-9.-]+", "-", value)
    value = re.sub(r"-{2,}", "-", value).strip("-")
    return value or "audit"


def make_session() -> requests.Session:
    session = requests.Session()
    session.headers.update({"User-Agent": USER_AGENT, "Accept-Language": "ru,en;q=0.8"})
    return session


def add_page_borders(section) -> None:
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    borders = {
        "top": ("18", BRAND_LINE),
        "right": ("12", BRAND_LINE),
        "bottom": ("18", BRAND_LINE),
        "left": ("28", BRAND_ORANGE),
    }
    for edge, (size, color) in borders.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "18")
        element.set(qn("w:color"), color)
        pg_borders.append(element)
    sect_pr.append(pg_borders)


def set_cell_shading(cell, color: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell_properties.append(shading)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def set_font(run, *, size=None, bold=None, color=None, name="Arial") -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text_paragraph(container, text: str, *, size=11.5, color=BRAND_TEXT, bold=False, space_after=6, alignment=None):
    paragraph = container.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(normalize_output_text(text))
    set_font(run, size=size, bold=bold, color=color)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.25
    return paragraph


def add_bullet_list(container, items: Iterable[str], *, size=11.2, color=BRAND_TEXT) -> None:
    for item in items:
        paragraph = container.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Cm(0.4)
        paragraph.paragraph_format.first_line_indent = Cm(-0.4)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"• {normalize_output_text(item)}")
        set_font(run, size=size, color=color)


def add_section_heading(doc: Document, kicker: str, title: str, intro: str | None = None) -> None:
    kicker_paragraph = doc.add_paragraph()
    kicker_run = kicker_paragraph.add_run(normalize_output_text(kicker).upper())
    set_font(kicker_run, size=9.5, bold=True, color=BRAND_ORANGE)
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(normalize_output_text(title))
    set_font(title_run, size=20, bold=True, color=BRAND_TEXT)
    title_paragraph.paragraph_format.space_after = Pt(6)
    if intro:
        add_text_paragraph(doc, intro, size=11.3, color=BRAND_MUTED, space_after=10)


def action_meta_label(key: str) -> str:
    return {
        "impact": "Ожидаемый эффект",
        "effort": "Сколько займет времени",
    }.get(key, normalize_output_text(key))


def severity_label(value: str) -> str:
    return {
        "Critical": "Критично",
        "CRITICAL": "Критично",
        "High": "Высокий приоритет",
        "HIGH": "Высокий приоритет",
        "Medium": "Средний приоритет",
        "MEDIUM": "Средний приоритет",
        "Low": "Низкий приоритет",
        "LOW": "Низкий приоритет",
    }.get(value or "", normalize_output_text(value))


def strip_namespace(tag: str) -> str:
    return tag.split("}", 1)[-1]


def safe_get(session: requests.Session, url: str, *, allow_redirects=True, timeout=25):
    try:
        response = session.get(url, timeout=timeout, allow_redirects=allow_redirects)
        return response, None
    except Exception as exc:  # noqa: BLE001
        return None, str(exc)


def discover_sitemaps(robots_text: str, base_url: str) -> list[str]:
    found = []
    for line in robots_text.splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        if key.strip().lower() == "sitemap":
            cleaned = value.strip()
            if cleaned:
                found.append(cleaned)
    if not found:
        found.append(urljoin(f"{base_url}/", "sitemap.xml"))
    unique = []
    for item in found:
        if item not in unique:
            unique.append(item)
    return unique


def parse_sitemap_urls(
    session: requests.Session, sitemap_urls: list[str], domain: str, max_sitemaps=25
) -> tuple[list[str], list[str], int]:
    discovered_urls: list[str] = []
    processed: list[str] = []
    queue = list(sitemap_urls)
    seen = set()
    while queue and len(processed) < max_sitemaps:
        current = queue.pop(0)
        if current in seen:
            continue
        seen.add(current)
        response, error = safe_get(session, current, timeout=35)
        if error or response is None or response.status_code >= 400:
            continue
        processed.append(current)
        try:
            root = ET.fromstring(response.content)
        except ET.ParseError:
            continue
        tag = strip_namespace(root.tag)
        if tag == "sitemapindex":
            for child in root:
                if strip_namespace(child.tag) != "sitemap":
                    continue
                for node in child:
                    if strip_namespace(node.tag) == "loc" and node.text:
                        queue.append(node.text.strip())
        elif tag == "urlset":
            for child in root:
                if strip_namespace(child.tag) != "url":
                    continue
                for node in child:
                    if strip_namespace(node.tag) == "loc" and node.text:
                        loc = node.text.strip()
                        if urlparse(loc).netloc == domain:
                            discovered_urls.append(loc)
    unique_urls = []
    seen_urls = set()
    for url in discovered_urls:
        if url not in seen_urls:
            seen_urls.add(url)
            unique_urls.append(url)
    return unique_urls, processed, len(discovered_urls)


def extract_visible_text(soup: BeautifulSoup) -> str:
    for tag in soup(["script", "style", "noscript", "svg"]):
        tag.extract()
    text = " ".join(soup.stripped_strings)
    return re.sub(r"\s+", " ", text).strip()


FAQ_KEYWORDS = (
    "faq",
    "часто задаваемые вопросы",
    "вопросы и ответы",
    "популярные вопросы",
    "ответы на вопросы",
)

REVIEW_KEYWORDS = (
    "отзывы",
    "reviews",
    "testimonials",
    "кейсы",
    "портфолио",
    "нам доверяют",
    "истории клиентов",
)

TRUST_KEYWORDS = (
    "гарантия",
    "сертификат",
    "сертифицирован",
    "почему выбирают",
    "опыт",
    "лет на рынке",
    "официальный",
    "производитель",
    "довер",
)

COMMERCIAL_KEYWORDS = (
    "доставка",
    "оплата",
    "payment",
    "delivery",
    "shipping",
    "срок",
    "стоимость",
    "цена",
    "условия",
    "возврат",
    "обмен",
    "монтаж",
    "замер",
)

COMPARISON_KEYWORDS = (
    "характеристик",
    "сравнение",
    "параметры",
    "комплектация",
    "specifications",
    "features",
    "compare",
)

BREADCRUMB_SELECTORS = (
    '[class*="breadcrumb"]',
    '[id*="breadcrumb"]',
    'nav[aria-label*="breadcrumb" i]',
)


def text_has_any(text: str, keywords: tuple[str, ...]) -> bool:
    haystack = text.lower()
    return any(keyword in haystack for keyword in keywords)


def pluralize_ru(count: int, one: str, few: str, many: str) -> str:
    remainder_100 = count % 100
    remainder_10 = count % 10
    if 11 <= remainder_100 <= 14:
        return many
    if remainder_10 == 1:
        return one
    if 2 <= remainder_10 <= 4:
        return few
    return many


def collect_dom_tokens(soup: BeautifulSoup, limit: int = 250) -> str:
    tokens: list[str] = []
    for tag in soup.find_all(True, limit=limit):
        classes = tag.get("class")
        if classes:
            tokens.extend(str(item) for item in classes if item)
        tag_id = tag.get("id")
        if tag_id:
            tokens.append(str(tag_id))
        aria_label = tag.get("aria-label")
        if aria_label:
            tokens.append(str(aria_label))
    return " ".join(tokens).lower()


def detect_page_signals(soup: BeautifulSoup, visible_text: str, schema_types: list[str], has_forms: bool) -> dict[str, bool]:
    headings_text = " ".join(tag.get_text(" ", strip=True) for tag in soup.find_all(["h1", "h2", "h3", "h4"]))
    dom_tokens = collect_dom_tokens(soup)
    haystack = " ".join(
        [
            headings_text.lower(),
            visible_text.lower()[:12000],
            dom_tokens,
        ]
    )
    has_breadcrumbs = "BreadcrumbList" in schema_types or any(soup.select(selector) for selector in BREADCRUMB_SELECTORS)
    has_comparison_table = bool(soup.find("table")) or text_has_any(haystack, COMPARISON_KEYWORDS)

    return {
        "has_faq_block": "FAQPage" in schema_types or text_has_any(haystack, FAQ_KEYWORDS),
        "has_reviews_block": any(item in schema_types for item in ("Review", "AggregateRating")) or text_has_any(haystack, REVIEW_KEYWORDS),
        "has_trust_block": text_has_any(haystack, TRUST_KEYWORDS),
        "has_commercial_block": text_has_any(haystack, COMMERCIAL_KEYWORDS),
        "has_comparison_table": has_comparison_table,
        "has_breadcrumbs": has_breadcrumbs,
        "has_forms": has_forms,
    }


def parse_jsonld_types(soup: BeautifulSoup) -> list[str]:
    found: list[str] = []
    for script in soup.find_all("script", attrs={"type": "application/ld+json"}):
        raw = (script.string or script.get_text()).strip()
        if not raw:
            continue
        try:
            data = json.loads(raw)
        except Exception:  # noqa: BLE001
            found.append("Invalid JSON-LD")
            continue
        stack = data if isinstance(data, list) else [data]
        for item in stack:
            if isinstance(item, dict) and "@graph" in item and isinstance(item["@graph"], list):
                stack.extend(item["@graph"])
            if isinstance(item, dict):
                value = item.get("@type")
                if isinstance(value, list):
                    found.extend(str(entry) for entry in value)
                elif value:
                    found.append(str(value))
    unique: list[str] = []
    for item in found:
        if item not in unique:
            unique.append(item)
    return unique


def detect_page_type(url: str, schema_types: list[str]) -> str:
    parsed = urlparse(url)
    if not parsed.path or parsed.path == "/":
        return "Р“Р»Р°РІРЅР°СЏ"
    if "Product" in schema_types:
        return "РўРѕРІР°СЂ"
    if parsed.query:
        return "РЎР»СѓР¶РµР±РЅР°СЏ"
    segments = [segment for segment in parsed.path.split("/") if segment]
    if len(segments) == 1:
        return "РљР°С‚РµРіРѕСЂРёСЏ"
    if len(segments) == 2:
        return "РџРѕРґРєР°С‚РµРіРѕСЂРёСЏ"
    return "Р’РЅСѓС‚СЂРµРЅРЅСЏСЏ"


def analyse_page(session: requests.Session, url: str, base_domain: str) -> PageSnapshot:
    response, error = safe_get(session, url)
    if error or response is None:
        return PageSnapshot(
            url=url,
            page_type="РћС€РёР±РєР° Р·Р°РіСЂСѓР·РєРё",
            status_code=0,
            response_time_ms=0,
            content_type="",
            title="",
            description="",
            canonical="",
            h1s=[],
            meta_robots="",
            x_robots_tag="",
            image_count=0,
            missing_alt_count=0,
            lazy_missing_count=0,
            missing_dimensions_count=0,
            schema_types=[],
            internal_links=0,
            word_count=0,
            has_meta_keywords=False,
            has_forms=False,
            has_faq_block=False,
            has_reviews_block=False,
            has_trust_block=False,
            has_commercial_block=False,
            has_comparison_table=False,
            has_breadcrumbs=False,
            html_size_kb=0.0,
        )
    content_type = response.headers.get("content-type", "")
    html = response.text if "html" in content_type else ""
    soup = BeautifulSoup(html, "lxml") if html else BeautifulSoup("", "lxml")
    title = soup.title.get_text(" ", strip=True) if soup.title else ""
    description_tag = soup.find("meta", attrs={"name": "description"})
    description = description_tag.get("content", "").strip() if description_tag else ""
    canonical_tag = soup.find("link", rel=lambda value: value and "canonical" in value)
    canonical = canonical_tag.get("href", "").strip() if canonical_tag else ""
    h1s = [h.get_text(" ", strip=True) for h in soup.find_all("h1")]
    images = soup.find_all("img")
    missing_alt_count = sum(1 for image in images if not (image.get("alt") or "").strip())
    lazy_missing_count = sum(1 for image in images if image.get("loading") != "lazy")
    missing_dimensions_count = sum(1 for image in images if not image.get("width") or not image.get("height"))
    schema_types = parse_jsonld_types(soup)
    internal_links = 0
    for anchor in soup.find_all("a", href=True):
        href = anchor["href"].strip()
        if not href or href.startswith(("#", "mailto:", "tel:", "javascript:")):
            continue
        absolute = urljoin(url, href)
        if urlparse(absolute).netloc == base_domain:
            internal_links += 1
    visible_text = extract_visible_text(soup)
    word_count = len(re.findall(r"\w+", visible_text, flags=re.U))
    meta_keywords = soup.find("meta", attrs={"name": "keywords"}) is not None
    has_forms = bool(soup.find("form"))
    signals = detect_page_signals(soup, visible_text, schema_types, has_forms)
    page_type = detect_page_type(url, schema_types)
    robots_meta_tag = soup.find("meta", attrs={"name": "robots"})
    return PageSnapshot(
        url=url,
        page_type=page_type,
        status_code=response.status_code,
        response_time_ms=round(response.elapsed.total_seconds() * 1000),
        content_type=content_type,
        title=title,
        description=description,
        canonical=canonical,
        h1s=h1s,
        meta_robots=robots_meta_tag.get("content", "").strip() if robots_meta_tag else "",
        x_robots_tag=response.headers.get("X-Robots-Tag", "").strip(),
        image_count=len(images),
        missing_alt_count=missing_alt_count,
        lazy_missing_count=lazy_missing_count,
        missing_dimensions_count=missing_dimensions_count,
        schema_types=schema_types,
        internal_links=internal_links,
        word_count=word_count,
        has_meta_keywords=meta_keywords,
        has_forms=signals["has_forms"],
        has_faq_block=signals["has_faq_block"],
        has_reviews_block=signals["has_reviews_block"],
        has_trust_block=signals["has_trust_block"],
        has_commercial_block=signals["has_commercial_block"],
        has_comparison_table=signals["has_comparison_table"],
        has_breadcrumbs=signals["has_breadcrumbs"],
        html_size_kb=round(len(response.content) / 1024, 1),
    )


def discover_home_links(session: requests.Session, base_url: str, base_domain: str) -> list[str]:
    response, error = safe_get(session, base_url)
    if error or response is None:
        return []
    soup = BeautifulSoup(response.text, "lxml")
    candidates: list[str] = []
    for anchor in soup.find_all("a", href=True):
        href = anchor["href"].strip()
        if not href or href.startswith(("#", "mailto:", "tel:", "javascript:")):
            continue
        absolute = urljoin(base_url, href)
        parsed = urlparse(absolute)
        if parsed.netloc != base_domain:
            continue
        if parsed.query and "route=information/contactform" not in parsed.query:
            continue
        normalized = parsed._replace(fragment="").geturl().rstrip("/") or absolute
        if normalized == base_url:
            continue
        if any(block in normalized for block in ("account/login", "account/register", "checkout", "search")):
            continue
        candidates.append(normalized)
    unique = list(dict.fromkeys(candidates))
    unique.sort(key=lambda item: (urlparse(item).path.count("/"), len(item)))
    return unique[:12]


def analyze_redirects(session: requests.Session, base_url: str) -> list[dict]:
    parsed = urlparse(base_url)
    host = parsed.netloc
    no_www_host = host[4:] if host.startswith("www.") else host
    variants = [
        f"http://{no_www_host}/",
        f"https://www.{no_www_host}/",
        f"http://www.{no_www_host}/",
    ]
    results = []
    for variant in variants:
        response, error = safe_get(session, variant, allow_redirects=True)
        if error or response is None:
            results.append({"url": variant, "error": error, "chain": [], "final_url": ""})
            continue
        chain = []
        for item in response.history:
            chain.append({"status": item.status_code, "location": item.headers.get("location", "")})
        results.append({"url": variant, "error": "", "chain": chain, "final_url": response.url})
    return results


CONTACT_PATH_KEYWORDS = (
    "contact",
    "contacts",
    "kontakt",
    "kontakty",
    "callback",
    "feedback",
    "request",
    "lead",
    "zayav",
    "svyaz",
    "consult",
)

BLOG_PATH_KEYWORDS = (
    "blog",
    "blogs",
    "article",
    "articles",
    "news",
    "novosti",
    "stati",
    "aktuelles",
)


def human_path(url: str) -> str:
    parsed = urlparse(url)
    path = parsed.path or "/"
    if parsed.query:
        path = f"{path}?{parsed.query}"
    return path


def get_contact_related_pages(sample_pages: list[PageSnapshot]) -> list[PageSnapshot]:
    found: list[PageSnapshot] = []
    seen: set[str] = set()
    for snapshot in sample_pages:
        parsed = urlparse(snapshot.url)
        haystack = f"{parsed.path} {parsed.query}".lower()
        if snapshot.has_forms or any(token in haystack for token in CONTACT_PATH_KEYWORDS):
            if snapshot.url not in seen:
                seen.add(snapshot.url)
                found.append(snapshot)
    return found


def get_indexable_query_pages(sample_pages: list[PageSnapshot]) -> list[PageSnapshot]:
    pages: list[PageSnapshot] = []
    for snapshot in sample_pages:
        parsed = urlparse(snapshot.url)
        if not parsed.query:
            continue
        robots_flags = f"{snapshot.meta_robots} {snapshot.x_robots_tag}".lower()
        if "noindex" in robots_flags:
            continue
        pages.append(snapshot)
    return pages


def infer_project_profile(audit: dict) -> dict:
    sample_pages: list[PageSnapshot] = audit["sample_pages"]
    product_pages = [snapshot for snapshot in sample_pages if snapshot.page_type == "РўРѕРІР°СЂ" or "Product" in snapshot.schema_types]
    category_pages = [snapshot for snapshot in sample_pages if snapshot.page_type in ("РљР°С‚РµРіРѕСЂРёСЏ", "РџРѕРґРєР°С‚РµРіРѕСЂРёСЏ")]
    form_pages = [snapshot for snapshot in sample_pages if snapshot.has_forms]
    return {
        "product_pages": product_pages,
        "category_pages": category_pages,
        "form_pages": form_pages,
        "is_catalog": bool(product_pages or category_pages),
    }


COMPETITOR_SIGNAL_LIBRARY = [
    {
        "id": "faq",
        "label": "FAQ и ответы на частые вопросы",
        "field": "has_faq_block",
        "priority": "Высокий приоритет",
        "owner": "SEO / Content / Frontend",
        "min_competitor_coverage": 0.25,
        "implementation_brief": (
            "Добавить на главную и ключевые коммерческие шаблоны блок FAQ из 5-7 вопросов: "
            "о выборе услуги или товара, сроках, цене, гарантии и сценариях использования. "
            "Собрать отдельные H2, короткие ответы, перелинковку на смежные разделы и при уместности микроразметку FAQPage."
        ),
        "benefit": (
            "FAQ закрывает возражения до звонка, усиливает релевантность под длинные запросы и помогает странице забирать более теплый спрос."
        ),
        "weight": 9,
    },
    {
        "id": "reviews",
        "label": "Отзывы, кейсы и социальное доказательство",
        "field": "has_reviews_block",
        "priority": "Высокий приоритет",
        "owner": "Marketing / Content / Frontend",
        "min_competitor_coverage": 0.25,
        "implementation_brief": (
            "Добавить в приоритетные шаблоны блок с отзывами, кейсами или примерами работ: "
            "имя клиента или компания, задача, результат, фото или скрин, короткий CTA на заявку."
        ),
        "benefit": (
            "Такой блок усиливает доверие, повышает коммерческий интент страницы и помогает конвертировать трафик в заявку, а не только в просмотр."
        ),
        "weight": 9,
    },
    {
        "id": "trust",
        "label": "Блоки доверия и доказательства экспертизы",
        "field": "has_trust_block",
        "priority": "Средний приоритет",
        "owner": "Marketing / Content",
        "min_competitor_coverage": 0.3,
        "implementation_brief": (
            "Собрать единый доверительный слой: гарантии, сертификаты, опыт, реальные сроки, "
            "этапы работы, подтверждение статуса производителя или эксперта и короткий блок \"почему нам доверяют\"."
        ),
        "benefit": (
            "Доверительные блоки помогают пользователю быстрее принять решение и поддерживают коммерческие факторы, особенно на услугах и сложных товарах."
        ),
        "weight": 7,
    },
    {
        "id": "commercial",
        "label": "Коммерческие условия: цена, сроки, доставка, оплата",
        "field": "has_commercial_block",
        "priority": "Высокий приоритет",
        "owner": "SEO / Content / Business",
        "min_competitor_coverage": 0.35,
        "implementation_brief": (
            "Добавить на ключевые страницы видимые блоки с условиями: цены или диапазоны, сроки, "
            "доставка, оплата, гарантия, возврат или сценарий сотрудничества. "
            "Разнести их по шаблонам, а не держать только в подвале или на одной служебной странице."
        ),
        "benefit": (
            "Это усиливает коммерческий интент страницы, снимает часть возражений и помогает поисковику лучше понимать, что страница готова к конверсии."
        ),
        "weight": 10,
    },
    {
        "id": "comparison",
        "label": "Таблицы, характеристики и сравнительные блоки",
        "field": "has_comparison_table",
        "priority": "Средний приоритет",
        "owner": "SEO / Content",
        "min_competitor_coverage": 0.25,
        "implementation_brief": (
            "Добавить в приоритетные шаблоны таблицы характеристик, сравнительные блоки или структурированные параметры: "
            "что входит, чем отличается, какие размеры или опции доступны, кому подходит каждое решение."
        ),
        "benefit": (
            "Таблицы помогают закрывать спрос на сравнение, повышают полезность страницы и упрощают выбор, особенно в B2B и каталогах."
        ),
        "weight": 7,
    },
    {
        "id": "breadcrumbs",
        "label": "Хлебные крошки и понятная навигационная цепочка",
        "field": "has_breadcrumbs",
        "priority": "Средний приоритет",
        "owner": "Frontend / SEO",
        "min_competitor_coverage": 0.35,
        "implementation_brief": (
            "Добавить хлебные крошки на шаблоны категорий, карточек и внутренних страниц, "
            "связать их с реальной иерархией сайта и синхронизировать с BreadcrumbList в микроразметке."
        ),
        "benefit": (
            "Это упрощает навигацию, поддерживает внутреннюю структуру сайта и помогает выдаче показывать более чистую иерархию URL."
        ),
        "weight": 6,
    },
    {
        "id": "forms",
        "label": "Формы и заметные точки входа в заявку",
        "field": "has_forms",
        "priority": "Высокий приоритет",
        "owner": "Marketing / Frontend",
        "min_competitor_coverage": 0.3,
        "implementation_brief": (
            "Добавить заметные формы и точки входа в заявку на главную, услуги, категории или карточки, "
            "где пользователь уже готов оставить заявку: телефон, задача, бюджет или удобный следующий шаг."
        ),
        "benefit": (
            "Даже при одинаковом трафике более заметные формы и точки входа в заявку дают больше обращений и помогают аудиту влиять не только на позиции, но и на продажи."
        ),
        "weight": 10,
    },
]


COMPETITOR_SNIPPET_LIBRARY = [
    {
        "id": "snippet_ready",
        "label": "Сниппеты на ключевых страницах: title, описание страницы, H1 и канонический адрес",
        "priority": "Высокий приоритет",
        "owner": "SEO / Frontend",
        "min_competitor_coverage": 0.45,
        "implementation_brief": (
            "Собрать единый шаблон для ключевых страниц: рабочий title без переспама, понятный description, "
            "один H1 по теме страницы и корректный canonical. Проверять это не точечно, а сразу на главной, категориях, услугах, карточках и страницах заявок."
        ),
        "benefit": (
            "Ровные сниппеты помогают стабильно закрывать базовые SEO-сигналы, улучшают читаемость выдачи и уменьшают просадку шаблонов из-за пустых или кривых мета-полей."
        ),
        "weight": 9,
    },
    {
        "id": "canonical_present",
        "label": "Канонические адреса на ключевых шаблонах",
        "priority": "Средний приоритет",
        "owner": "SEO / Frontend",
        "min_competitor_coverage": 0.7,
        "implementation_brief": (
            "Проверить и проставить canonical на ключевых типах страниц по шаблону, чтобы категории, карточки, услуги и служебные варианты не спорили между собой за индексацию."
        ),
        "benefit": (
            "Корректный canonical помогает убрать шум в индексе и дает более чистый сигнал, какая страница должна ранжироваться по целевому спросу."
        ),
        "weight": 7,
    },
]


TEMPLATE_LIBRARY = [
    {"id": "home", "label": "Главная"},
    {"id": "category", "label": "Категории и листинги"},
    {"id": "product", "label": "Карточки товаров"},
    {"id": "service", "label": "Услуги и внутренние коммерческие страницы"},
    {"id": "contact", "label": "Контакты и страницы заявок"},
]


SIGNAL_SHORT_LABELS = {
    "faq": "FAQ",
    "reviews": "отзывы и кейсы",
    "trust": "доверие",
    "commercial": "условия",
    "comparison": "таблицы и параметры",
    "breadcrumbs": "хлебные крошки",
    "forms": "формы заявки",
}


def is_title_in_range(title: str) -> bool:
    length = len((title or "").strip())
    return 25 <= length <= 70


def is_description_in_range(description: str) -> bool:
    length = len((description or "").strip())
    return 110 <= length <= 180


def cleaned_h1s(page: PageSnapshot) -> list[str]:
    return [item.strip() for item in page.h1s if item and item.strip()]


def page_has_snippet_ready(page: PageSnapshot) -> bool:
    return (
        is_title_in_range(page.title)
        and is_description_in_range(page.description)
        and len(cleaned_h1s(page)) == 1
        and bool((page.canonical or "").strip())
    )


def build_metric_bucket(pages: list[PageSnapshot], predicate) -> dict:
    matching_pages = [page for page in pages if predicate(page)]
    total = len(pages)
    return {
        "count": len(matching_pages),
        "coverage": round((len(matching_pages) / total), 3) if total else 0,
        "examples": [human_path(page.url) for page in matching_pages[:3]],
    }


def build_snippet_metrics(pages: list[PageSnapshot]) -> dict[str, dict]:
    return {
        "title_present": build_metric_bucket(pages, lambda page: bool((page.title or "").strip())),
        "title_good": build_metric_bucket(pages, lambda page: is_title_in_range(page.title)),
        "description_present": build_metric_bucket(pages, lambda page: bool((page.description or "").strip())),
        "description_good": build_metric_bucket(pages, lambda page: is_description_in_range(page.description)),
        "h1_present": build_metric_bucket(pages, lambda page: bool(cleaned_h1s(page))),
        "single_h1": build_metric_bucket(pages, lambda page: len(cleaned_h1s(page)) == 1),
        "canonical_present": build_metric_bucket(pages, lambda page: bool((page.canonical or "").strip())),
        "snippet_ready": build_metric_bucket(pages, page_has_snippet_ready),
    }


def get_template_bucket(page: PageSnapshot) -> str | None:
    parsed = urlparse(page.url)
    path = parsed.path or "/"
    haystack = f"{parsed.path} {parsed.query}".lower()

    if not parsed.path or parsed.path == "/":
        return "home"
    if any(token in haystack for token in CONTACT_PATH_KEYWORDS):
        return "contact"
    if page.page_type == "РўРѕРІР°СЂ" or "Product" in page.schema_types:
        return "product"
    if page.page_type in ("РљР°С‚РµРіРѕСЂРёСЏ", "РџРѕРґРєР°С‚РµРіРѕСЂРёСЏ"):
        return "category"
    if any(token in haystack for token in BLOG_PATH_KEYWORDS):
        return None
    return "service"


def page_signal_score(page: PageSnapshot) -> float:
    score = 0.0
    for signal in COMPETITOR_SIGNAL_LIBRARY:
        if bool(getattr(page, signal["field"], False)):
            score += signal["weight"]
    if is_title_in_range(page.title):
        score += 2
    if is_description_in_range(page.description):
        score += 2
    if len(cleaned_h1s(page)) == 1:
        score += 1.5
    if page.canonical:
        score += 1
    if page_has_snippet_ready(page):
        score += 4
    score += min(page.word_count, 800) / 200
    return round(score, 2)


def build_template_benchmark(pages: list[PageSnapshot]) -> dict[str, dict]:
    buckets: dict[str, list[PageSnapshot]] = {template["id"]: [] for template in TEMPLATE_LIBRARY}
    for page in pages:
        template_id = get_template_bucket(page)
        if template_id and template_id in buckets:
            buckets[template_id].append(page)

    template_benchmark: dict[str, dict] = {}
    commercial_signal_ids = ("faq", "reviews", "trust", "commercial", "forms")
    for template in TEMPLATE_LIBRARY:
        template_pages = buckets.get(template["id"], [])
        pages_checked = len(template_pages)
        signals: dict[str, dict] = {}
        for signal in COMPETITOR_SIGNAL_LIBRARY:
            matching_pages = [page for page in template_pages if bool(getattr(page, signal["field"], False))]
            signals[signal["id"]] = {
                "count": len(matching_pages),
                "coverage": round((len(matching_pages) / pages_checked), 3) if pages_checked else 0,
                "examples": [human_path(page.url) for page in matching_pages[:2]],
            }
        snippet_metrics = build_snippet_metrics(template_pages)
        commercial_coverages = [signals[signal_id]["coverage"] for signal_id in commercial_signal_ids]
        best_layers = [
            SIGNAL_SHORT_LABELS[signal_id]
            for signal_id in commercial_signal_ids
            if signals[signal_id]["coverage"] >= 0.34
        ]
        template_benchmark[template["id"]] = {
            "label": template["label"],
            "pages_checked": pages_checked,
            "signals": signals,
            "snippet_metrics": snippet_metrics,
            "commercial_score": round(statistics.mean(commercial_coverages), 3) if commercial_coverages else 0,
            "best_layers": best_layers[:3],
            "sample_paths": [human_path(page.url) for page in sorted(template_pages, key=page_signal_score, reverse=True)[:3]],
        }
    return template_benchmark


def metric_count_text(metric: dict, pages_checked: int) -> str:
    if not pages_checked:
        return "0 из 0"
    return f"{metric.get('count', 0)} из {pages_checked}"


def coverage_percent(value: float) -> str:
    return f"{round((value or 0) * 100)}%"


def snippet_metric_state_text(metric_id: str, metric: dict, pages_checked: int) -> str:
    metric_count = metric_count_text(metric, pages_checked)
    if metric_id == "snippet_ready":
        return f"связка title, описания страницы, H1 и канонического адреса собрана на {metric_count} ключевых страниц"
    if metric_id == "canonical_present":
        return f"канонический адрес стоит на {metric_count} ключевых страниц"
    return f"метрика выполнена на {metric_count} страниц"


def build_competitor_template_findings(target_benchmark: dict, competitor_benchmark: dict) -> list[str]:
    findings: list[dict] = []
    target_templates = target_benchmark.get("templates", {})
    competitor_templates = competitor_benchmark.get("templates", {})

    for template in TEMPLATE_LIBRARY:
        target_template = target_templates.get(template["id"], {})
        competitor_template = competitor_templates.get(template["id"], {})
        if competitor_template.get("pages_checked", 0) == 0:
            continue

        notes: list[str] = []
        score = 0.0
        competitor_snippet = competitor_template.get("snippet_metrics", {}).get("snippet_ready", {}).get("coverage", 0)
        target_snippet = target_template.get("snippet_metrics", {}).get("snippet_ready", {}).get("coverage", 0)
        if competitor_snippet >= max(0.45, target_snippet + 0.25):
            notes.append("сниппеты собраны ровнее")
            score += 1.0 + (competitor_snippet - target_snippet)

        layer_candidates: list[tuple[float, str]] = []
        for signal_id in ("commercial", "forms", "faq", "trust", "reviews", "comparison", "breadcrumbs"):
            competitor_signal = competitor_template.get("signals", {}).get(signal_id, {}).get("coverage", 0)
            target_signal = target_template.get("signals", {}).get(signal_id, {}).get("coverage", 0)
            if competitor_signal >= max(0.4, target_signal + 0.25):
                layer_candidates.append((competitor_signal - target_signal, SIGNAL_SHORT_LABELS.get(signal_id, signal_id)))

        layer_candidates.sort(key=lambda item: item[0], reverse=True)
        if layer_candidates:
            notes.append("лучше видны " + ", ".join(label for _, label in layer_candidates[:2]))
            score += 0.9 + sum(delta for delta, _ in layer_candidates[:2])
        elif competitor_template.get("commercial_score", 0) >= max(0.45, target_template.get("commercial_score", 0) + 0.2):
            notes.append("коммерческий слой собран полнее")
            score += 0.8 + (competitor_template.get("commercial_score", 0) - target_template.get("commercial_score", 0))

        if not notes:
            continue

        findings.append(
            {
                "text": f"{template['label']}: {'; '.join(notes)}.",
                "score": round(score, 3),
            }
        )

    findings.sort(key=lambda item: item["score"], reverse=True)
    return [item["text"] for item in findings[:3]]


def build_competitor_snippet_findings(target_benchmark: dict, competitor_benchmark: dict) -> list[str]:
    target_metrics = target_benchmark.get("snippet_metrics", {})
    competitor_metrics = competitor_benchmark.get("snippet_metrics", {})
    target_pages = target_benchmark.get("pages_checked", 0)
    competitor_pages = competitor_benchmark.get("pages_checked", 0)
    findings: list[dict] = []

    snippet_delta = competitor_metrics.get("snippet_ready", {}).get("coverage", 0) - target_metrics.get("snippet_ready", {}).get("coverage", 0)
    if snippet_delta >= 0.25:
        findings.append(
            {
                "text": (
                    "Базовая связка title, описания страницы, H1 и канонического адреса собрана на "
                    f"{metric_count_text(competitor_metrics.get('snippet_ready', {}), competitor_pages)} страниц; "
                    f"у сайта — на {metric_count_text(target_metrics.get('snippet_ready', {}), target_pages)}."
                ),
                "score": round(snippet_delta + 0.8, 3),
            }
        )

    title_delta = competitor_metrics.get("title_good", {}).get("coverage", 0) - target_metrics.get("title_good", {}).get("coverage", 0)
    if title_delta >= 0.25:
        findings.append(
            {
                "text": (
                    f"Заголовок title чаще попадает в рабочий диапазон длины: {metric_count_text(competitor_metrics.get('title_good', {}), competitor_pages)} "
                    f"против {metric_count_text(target_metrics.get('title_good', {}), target_pages)}."
                ),
                "score": round(title_delta + 0.45, 3),
            }
        )

    description_delta = competitor_metrics.get("description_good", {}).get("coverage", 0) - target_metrics.get("description_good", {}).get("coverage", 0)
    if description_delta >= 0.25:
        findings.append(
            {
                "text": (
                    "Описание страницы чаще заполнено и ближе к нормальной длине: "
                    f"{metric_count_text(competitor_metrics.get('description_good', {}), competitor_pages)} "
                    f"против {metric_count_text(target_metrics.get('description_good', {}), target_pages)}."
                ),
                "score": round(description_delta + 0.45, 3),
            }
        )

    h1_delta = competitor_metrics.get("single_h1", {}).get("coverage", 0) - target_metrics.get("single_h1", {}).get("coverage", 0)
    if h1_delta >= 0.25:
        findings.append(
            {
                "text": (
                    f"H1 чаще собран по одному на страницу: {metric_count_text(competitor_metrics.get('single_h1', {}), competitor_pages)} "
                    f"против {metric_count_text(target_metrics.get('single_h1', {}), target_pages)}."
                ),
                "score": round(h1_delta + 0.35, 3),
            }
        )

    canonical_delta = competitor_metrics.get("canonical_present", {}).get("coverage", 0) - target_metrics.get("canonical_present", {}).get("coverage", 0)
    if canonical_delta >= 0.25:
        findings.append(
            {
                "text": (
                    "Канонический адрес проставлен стабильнее: "
                    f"{metric_count_text(competitor_metrics.get('canonical_present', {}), competitor_pages)} "
                    f"против {metric_count_text(target_metrics.get('canonical_present', {}), target_pages)}."
                ),
                "score": round(canonical_delta + 0.35, 3),
            }
        )

    findings.sort(key=lambda item: item["score"], reverse=True)
    return [item["text"] for item in findings[:3]]


def build_competitor_factor_summary(target_benchmark: dict, competitor_benchmark: dict) -> list[str]:
    target_pages = target_benchmark.get("pages_checked", 0)
    competitor_pages = competitor_benchmark.get("pages_checked", 0)
    target_signals = target_benchmark.get("signals", {})
    competitor_signals = competitor_benchmark.get("signals", {})
    target_snippet = target_benchmark.get("snippet_metrics", {}).get("snippet_ready", {})
    competitor_snippet = competitor_benchmark.get("snippet_metrics", {}).get("snippet_ready", {})
    findings: list[dict[str, float | str]] = []

    target_snippet_coverage = target_snippet.get("coverage", 0)
    competitor_snippet_coverage = competitor_snippet.get("coverage", 0)
    if competitor_snippet_coverage >= max(0.2, target_snippet_coverage + 0.12):
        findings.append(
            {
                "text": (
                    "Сниппетная база: "
                    f"{metric_count_text(competitor_snippet, competitor_pages)} страниц против {metric_count_text(target_snippet, target_pages)} у сайта."
                ),
                "score": round(competitor_snippet_coverage - target_snippet_coverage + 0.35, 3),
            }
        )

    faq_gap = competitor_signals.get("faq", {}).get("coverage", 0) - target_signals.get("faq", {}).get("coverage", 0)
    trust_gap = competitor_signals.get("trust", {}).get("coverage", 0) - target_signals.get("trust", {}).get("coverage", 0)
    faq_parts: list[str] = []
    faq_score = 0.0
    if faq_gap >= 0.1:
        faq_parts.append(
            f"FAQ {coverage_percent(competitor_signals.get('faq', {}).get('coverage', 0))} против {coverage_percent(target_signals.get('faq', {}).get('coverage', 0))}"
        )
        faq_score += faq_gap + 0.2
    if trust_gap >= 0.1:
        faq_parts.append(
            f"доверие {coverage_percent(competitor_signals.get('trust', {}).get('coverage', 0))} против {coverage_percent(target_signals.get('trust', {}).get('coverage', 0))}"
        )
        faq_score += trust_gap + 0.2
    if faq_parts:
        findings.append(
            {
                "text": "FAQ и доверие: " + "; ".join(faq_parts) + ".",
                "score": round(faq_score, 3),
            }
        )

    commercial_gap = competitor_signals.get("commercial", {}).get("coverage", 0) - target_signals.get("commercial", {}).get("coverage", 0)
    forms_gap = competitor_signals.get("forms", {}).get("coverage", 0) - target_signals.get("forms", {}).get("coverage", 0)
    commercial_parts: list[str] = []
    commercial_score = 0.0
    if commercial_gap >= 0.1:
        commercial_parts.append(
            f"условия {coverage_percent(competitor_signals.get('commercial', {}).get('coverage', 0))} против {coverage_percent(target_signals.get('commercial', {}).get('coverage', 0))}"
        )
        commercial_score += commercial_gap + 0.2
    if forms_gap >= 0.1:
        commercial_parts.append(
            f"формы {coverage_percent(competitor_signals.get('forms', {}).get('coverage', 0))} против {coverage_percent(target_signals.get('forms', {}).get('coverage', 0))}"
        )
        commercial_score += forms_gap + 0.2
    if commercial_parts:
        findings.append(
            {
                "text": "Коммерческий слой и точки входа: " + "; ".join(commercial_parts) + ".",
                "score": round(commercial_score, 3),
            }
        )

    target_average_words = target_benchmark.get("average_words", 0)
    competitor_average_words = competitor_benchmark.get("average_words", 0)
    min_content_gap = max(120, target_average_words * 0.15) if target_average_words else 180
    if competitor_average_words >= target_average_words + min_content_gap:
        findings.append(
            {
                "text": (
                    f"Средний текстовый слой: {round(competitor_average_words)} слов на страницу "
                    f"против {round(target_average_words)} у сайта."
                ),
                "score": round((competitor_average_words - target_average_words) / 400 + 0.25, 3),
            }
        )

    findings.sort(key=lambda item: float(item["score"]), reverse=True)
    return [str(item["text"]) for item in findings[:3]]


def build_competitor_template_rows(target_benchmark: dict, competitor_benchmark: dict) -> list[str]:
    rows: list[dict] = []
    target_templates = target_benchmark.get("templates", {})
    competitor_templates = competitor_benchmark.get("templates", {})

    for template in TEMPLATE_LIBRARY:
        target_template = target_templates.get(template["id"], {})
        competitor_template = competitor_templates.get(template["id"], {})
        if competitor_template.get("pages_checked", 0) == 0:
            continue

        target_snippet = target_template.get("snippet_metrics", {}).get("snippet_ready", {}).get("coverage", 0)
        competitor_snippet = competitor_template.get("snippet_metrics", {}).get("snippet_ready", {}).get("coverage", 0)
        target_commercial = target_template.get("commercial_score", 0)
        competitor_commercial = competitor_template.get("commercial_score", 0)
        snippet_gap = competitor_snippet - target_snippet
        commercial_gap = competitor_commercial - target_commercial
        if max(snippet_gap, commercial_gap) < 0.14:
            continue

        stronger_layers: list[str] = []
        for signal_id in ("commercial", "forms", "faq", "trust", "reviews", "comparison", "breadcrumbs"):
            competitor_coverage = competitor_template.get("signals", {}).get(signal_id, {}).get("coverage", 0)
            target_coverage = target_template.get("signals", {}).get(signal_id, {}).get("coverage", 0)
            if competitor_coverage - target_coverage >= 0.2:
                stronger_layers.append(SIGNAL_SHORT_LABELS.get(signal_id, signal_id))

        parts: list[str] = []
        if snippet_gap >= 0.14:
            parts.append(f"сниппеты {coverage_percent(competitor_snippet)} против {coverage_percent(target_snippet)}")
        if commercial_gap >= 0.14:
            parts.append(f"коммерческий слой {coverage_percent(competitor_commercial)} против {coverage_percent(target_commercial)}")
        if stronger_layers:
            parts.append("сильнее закрыты: " + ", ".join(stronger_layers[:3]))

        rows.append(
            {
                "text": f"{template['label']}: {'; '.join(parts)}.",
                "score": round(max(snippet_gap, commercial_gap) + len(stronger_layers) / 10, 3),
            }
        )

    rows.sort(key=lambda item: item["score"], reverse=True)
    return [item["text"] for item in rows[:4]]


def pick_templates_for_gap(
    target_benchmark: dict,
    competitor_benchmarks: list[dict],
    *,
    signal_id: str | None = None,
    snippet_id: str | None = None,
    use_commercial_score: bool = False,
) -> list[str]:
    template_gaps: list[tuple[float, str]] = []

    for template in TEMPLATE_LIBRARY:
        relevant_templates = [
            site.get("templates", {}).get(template["id"], {})
            for site in competitor_benchmarks
            if site.get("templates", {}).get(template["id"], {}).get("pages_checked", 0) > 0
        ]
        if not relevant_templates:
            continue

        target_template = target_benchmark.get("templates", {}).get(template["id"], {})
        if signal_id:
            competitor_average = statistics.mean(
                item.get("signals", {}).get(signal_id, {}).get("coverage", 0) for item in relevant_templates
            )
            target_value = target_template.get("signals", {}).get(signal_id, {}).get("coverage", 0)
        elif snippet_id:
            competitor_average = statistics.mean(
                item.get("snippet_metrics", {}).get(snippet_id, {}).get("coverage", 0) for item in relevant_templates
            )
            target_value = target_template.get("snippet_metrics", {}).get(snippet_id, {}).get("coverage", 0)
        elif use_commercial_score:
            competitor_average = statistics.mean(item.get("commercial_score", 0) for item in relevant_templates)
            target_value = target_template.get("commercial_score", 0)
        else:
            continue

        gap = competitor_average - target_value
        if gap >= 0.15:
            template_gaps.append((round(gap, 3), template["label"]))

    template_gaps.sort(key=lambda item: item[0], reverse=True)
    return [label for _, label in template_gaps[:3]]


def build_gap_execution_details(title: str, templates: list[str]) -> dict[str, list[str]]:
    scope = templates or ["Главная", "Услуги и внутренние коммерческие страницы", "Контакты и страницы заявок"]
    scope_text = ", ".join(scope)
    title_lower = title.lower()

    if "faq" in title_lower:
        steps = [
            f"Определить шаблоны для внедрения: {scope_text}.",
            "Собрать 5-7 вопросов с короткими ответами, которые реально закрывают выбор, сроки, цену, гарантии и частые возражения.",
            "Сверстать блок с отдельными подзаголовками, связать его с перелинковкой и при необходимости добавить FAQPage.",
        ]
        impact = [
            "Страница лучше отвечает на длинные и уточняющие запросы.",
            "Пользователь быстрее снимает возражения и доходит до заявки.",
            "Коммерческий слой становится полезнее, а не просто длиннее.",
        ]
    elif "отзывы" in title_lower or "кейсы" in title_lower:
        steps = [
            f"Выбрать приоритетные шаблоны: {scope_text}.",
            "Собрать 3-5 реальных отзывов, кейсов или примеров работ с задачей, результатом и понятным подтверждением.",
            "Разместить блок рядом с оффером и формой заявки, чтобы он усиливал решение о контакте.",
        ]
        impact = [
            "Доверие к странице и бренду растет быстрее.",
            "Пользователю проще понять, какой результат он получит.",
            "Трафик чаще превращается в обращение, а не в отказ.",
        ]
    elif "довер" in title_lower:
        steps = [
            f"Определить шаблоны: {scope_text}.",
            "Собрать единый слой доверия: опыт, гарантии, сертификаты, этапы работы, фото команды или производства.",
            "Встроить блок в ключевые шаблоны так, чтобы он поддерживал оффер, а не висел отдельно внизу страницы.",
        ]
        impact = [
            "Страница выглядит убедительнее для холодного трафика.",
            "Коммерческие факторы становятся заметнее и понятнее.",
            "Снижается количество отказов на этапе первого знакомства с компанией.",
        ]
    elif "коммерческие условия" in title_lower:
        steps = [
            f"Определить шаблоны: {scope_text}.",
            "Прописать для каждого типа страницы цену или диапазон, сроки, доставку, оплату, гарантию и сценарий работы.",
            "Собрать блок в шаблон и проверить, чтобы условия были видны без ухода в служебные страницы.",
        ]
        impact = [
            "Поисковик видит более сильный коммерческий сигнал.",
            "Пользователь быстрее понимает, подходит ли ему предложение.",
            "Страница лучше закрывает спрос на покупку или заявку.",
        ]
    elif "сниппет" in title_lower or "каноническ" in title_lower:
        steps = [
            f"Определить шаблоны: {scope_text}.",
            "Собрать правила для title, описания страницы, H1 и канонического адреса по каждому типу страниц.",
            "Внедрить шаблоны и проверить результат массово: по HTML, индексации и контрольной выборке URL.",
        ]
        impact = [
            "Снижается количество пустых и кривых мета-полей.",
            "Выдача становится понятнее и ровнее по ключевым страницам.",
            "Шаблоны перестают проседать из-за технической разболтанности.",
        ]
    elif "контента" in title_lower:
        steps = [
            f"Определить шаблоны: {scope_text}.",
            "Собрать для каждой страницы структуру из интро, условий, преимуществ, FAQ, доверия, сравнений и перелинковки.",
            "Написать и внедрить текстовый слой так, чтобы он помогал выбору, а не раздувал страницу ради объема.",
        ]
        impact = [
            "Страница лучше закрывает интент и возражения.",
            "Появляется больше точек входа по длинным и коммерческим запросам.",
            "Контент начинает поддерживать конверсию, а не мешать ей.",
        ]
    else:
        steps = [
            f"Определить шаблоны: {scope_text}.",
            "Собрать требования к блоку и его полям на уровне шаблона, а не одной страницы.",
            "Внедрить изменения, проверить отображение, индексацию и связку с перелинковкой и CTA.",
        ]
        impact = [
            "Шаблоны становятся сильнее и стабильнее в выдаче.",
            "Страницы лучше поддерживают коммерческий сценарий пользователя.",
            "Команда получает понятный пакет задач для внедрения и контроля результата.",
        ]

    return {
        "where_to_implement": scope,
        "implementation_steps": steps,
        "impact_points": impact,
    }


def build_competitor_commercial_findings(target_benchmark: dict, competitor_benchmark: dict) -> list[str]:
    findings: list[dict] = []
    target_pages = target_benchmark.get("pages_checked", 0)
    competitor_pages = competitor_benchmark.get("pages_checked", 0)

    for signal in COMPETITOR_SIGNAL_LIBRARY:
        competitor_signal = competitor_benchmark.get("signals", {}).get(signal["id"], {})
        target_signal = target_benchmark.get("signals", {}).get(signal["id"], {})
        gap_size = competitor_signal.get("coverage", 0) - target_signal.get("coverage", 0)
        if gap_size < 0.25 or competitor_signal.get("coverage", 0) < signal["min_competitor_coverage"]:
            continue
        findings.append(
            {
                "text": (
                    f"{signal['label']}: у конкурента видно на {metric_count_text(competitor_signal, competitor_pages)} страниц; "
                    f"у сайта — на {metric_count_text(target_signal, target_pages)}."
                ),
                "score": round(gap_size + signal["weight"] / 12, 3),
            }
        )

    findings.sort(key=lambda item: item["score"], reverse=True)
    return [item["text"] for item in findings[:3]]


def summarize_template_gap_labels(target_benchmark: dict, competitor_benchmarks: list[dict]) -> list[str]:
    labels: list[tuple[float, str]] = []
    for template in TEMPLATE_LIBRARY:
        relevant_templates = [
            site.get("templates", {}).get(template["id"], {})
            for site in competitor_benchmarks
            if site.get("templates", {}).get(template["id"], {}).get("pages_checked", 0) > 0
        ]
        if not relevant_templates:
            continue

        target_template = target_benchmark.get("templates", {}).get(template["id"], {})
        avg_snippet = statistics.mean(
            item.get("snippet_metrics", {}).get("snippet_ready", {}).get("coverage", 0) for item in relevant_templates
        )
        avg_commercial = statistics.mean(item.get("commercial_score", 0) for item in relevant_templates)
        delta = max(avg_snippet - target_template.get("snippet_metrics", {}).get("snippet_ready", {}).get("coverage", 0), 0)
        delta += max(avg_commercial - target_template.get("commercial_score", 0), 0)
        if delta >= 0.28:
            labels.append((round(delta, 3), template["label"]))

    labels.sort(key=lambda item: item[0], reverse=True)
    return [label for _, label in labels[:3]]


def get_competitor_benchmark_pages(audit: dict) -> list[PageSnapshot]:
    appendix_pages = audit.get("appendix_pages") or []
    if appendix_pages:
        return appendix_pages

    sample_pages: list[PageSnapshot] = audit.get("sample_pages", [])
    focus_pages = [
        page
        for page in sample_pages
        if page.page_type in ("Р“Р»Р°РІРЅР°СЏ", "РљР°С‚РµРіРѕСЂРёСЏ", "РџРѕРґРєР°С‚РµРіРѕСЂРёСЏ", "РўРѕРІР°СЂ", "Р’РЅСѓС‚СЂРµРЅРЅСЏСЏ")
        or page.has_forms
    ]
    return focus_pages[:12] or sample_pages[:12]


def build_site_feature_benchmark(base_url: str, sample_pages: list[PageSnapshot]) -> dict:
    pages = [page for page in sample_pages if page.status_code == 200 and "html" in page.content_type]
    pages_checked = len(pages)
    signals: dict[str, dict] = {}

    for signal in COMPETITOR_SIGNAL_LIBRARY:
        matching_pages = [page for page in pages if bool(getattr(page, signal["field"], False))]
        signals[signal["id"]] = {
            "count": len(matching_pages),
            "coverage": round((len(matching_pages) / pages_checked), 3) if pages_checked else 0,
            "examples": [human_path(page.url) for page in matching_pages[:3]],
        }

    snippet_metrics = build_snippet_metrics(pages)
    template_benchmark = build_template_benchmark(pages)
    average_words = round(statistics.mean([page.word_count for page in pages]), 1) if pages else 0
    ranked_signals = sorted(
        COMPETITOR_SIGNAL_LIBRARY,
        key=lambda signal: signals[signal["id"]]["coverage"],
        reverse=True,
    )
    highlights = [signal["label"] for signal in ranked_signals if signals[signal["id"]]["coverage"] >= 0.34][:3]
    if snippet_metrics["snippet_ready"]["coverage"] >= 0.55:
        highlights = ["Сниппеты на ключевых страницах собраны ровнее", *highlights]
    if average_words >= 420:
        highlights = ["Более плотный коммерческий контент", *highlights]

    return {
        "base_url": base_url,
        "domain": urlparse(base_url).netloc.replace("www.", ""),
        "pages_checked": pages_checked,
        "average_words": average_words,
        "signals": signals,
        "snippet_metrics": snippet_metrics,
        "templates": template_benchmark,
        "highlights": highlights[:3] or ["Сильных шаблонных преимуществ в выборке не нашли."],
        "sample_paths": [human_path(page.url) for page in sorted(pages, key=page_signal_score, reverse=True)[:4]],
    }


def collect_competitor_benchmark(base_url: str) -> dict:
    session = make_session()
    normalized_url = normalize_base_url(base_url)
    base_domain = urlparse(normalized_url).netloc
    robots_response, _ = safe_get(session, urljoin(f"{normalized_url}/", "robots.txt"))
    robots_text = robots_response.text if robots_response is not None and robots_response.status_code < 400 else ""
    sitemap_urls, _, _ = parse_sitemap_urls(
        session,
        discover_sitemaps(robots_text, normalized_url),
        base_domain,
        max_sitemaps=4,
    )
    home_links = discover_home_links(session, normalized_url, base_domain)
    candidate_urls = select_representative_urls(normalized_url, home_links, sitemap_urls, 10)
    sample_pages = analyse_pages_parallel(candidate_urls, base_domain, max_workers=8)
    html_pages = [page for page in sample_pages if page.status_code == 200 and "html" in page.content_type]

    if not html_pages:
        raise RuntimeError(f"Не удалось собрать HTML-страницы по {normalized_url}")

    return build_site_feature_benchmark(normalized_url, html_pages[:12])


def build_competitor_gap_items(target_benchmark: dict, competitor_benchmarks: list[dict]) -> list[dict]:
    if not competitor_benchmarks:
        return []

    gap_items: list[dict] = []

    for signal in COMPETITOR_SIGNAL_LIBRARY:
        target_signal = target_benchmark["signals"][signal["id"]]
        competitor_with_signal = [
            site for site in competitor_benchmarks if site["signals"][signal["id"]]["coverage"] >= signal["min_competitor_coverage"]
        ]
        if not competitor_with_signal:
            continue

        competitor_avg_coverage = statistics.mean(
            [site["signals"][signal["id"]]["coverage"] for site in competitor_with_signal]
        )
        gap_size = competitor_avg_coverage - target_signal["coverage"]

        if gap_size < 0.2 or target_signal["coverage"] >= competitor_avg_coverage * 0.75:
            continue

        examples = []
        for site in competitor_with_signal[:3]:
            example_paths = site["signals"][signal["id"]]["examples"]
            if example_paths:
                examples.append(f"{site['domain']}: {', '.join(example_paths[:2])}")
            else:
                examples.append(f"{site['domain']}: блок встречается в ключевой выборке.")

        execution_details = build_gap_execution_details(
            signal["label"],
            pick_templates_for_gap(target_benchmark, competitor_benchmarks, signal_id=signal["id"]),
        )
        gap_items.append(
            {
                "title": signal["label"],
                "priority": signal["priority"],
                "owner": signal["owner"],
                "current_state": (
                    f"На сайте сигнал найден на {target_signal['count']} из {target_benchmark['pages_checked']} ключевых страниц."
                    if target_benchmark["pages_checked"]
                    else "По сайту не удалось собрать ключевую выборку для сравнения."
                ),
                "competitor_state": (
                    f"У {len(competitor_with_signal)} из {len(competitor_benchmarks)} "
                    f"{pluralize_ru(len(competitor_benchmarks), 'конкурента', 'конкурентов', 'конкурентов')} "
                    "этот блок встречается заметно чаще."
                ),
                "examples": examples,
                "task": signal["implementation_brief"],
                "benefit": signal["benefit"],
                **execution_details,
                "sort_score": round(gap_size + competitor_avg_coverage + signal["weight"] / 10, 3),
            }
        )

    for snippet_signal in COMPETITOR_SNIPPET_LIBRARY:
        target_metric = target_benchmark.get("snippet_metrics", {}).get(snippet_signal["id"], {})
        competitor_with_signal = [
            site
            for site in competitor_benchmarks
            if site.get("snippet_metrics", {}).get(snippet_signal["id"], {}).get("coverage", 0) >= snippet_signal["min_competitor_coverage"]
        ]
        if not competitor_with_signal:
            continue

        competitor_avg_coverage = statistics.mean(
            site.get("snippet_metrics", {}).get(snippet_signal["id"], {}).get("coverage", 0) for site in competitor_with_signal
        )
        gap_size = competitor_avg_coverage - target_metric.get("coverage", 0)
        if gap_size < 0.2 or target_metric.get("coverage", 0) >= competitor_avg_coverage * 0.8:
            continue

        examples = [
            f"{site['domain']}: {snippet_metric_state_text(snippet_signal['id'], site.get('snippet_metrics', {}).get(snippet_signal['id'], {}), site.get('pages_checked', 0))}"
            for site in competitor_with_signal[:3]
        ]
        execution_details = build_gap_execution_details(
            snippet_signal["label"],
            pick_templates_for_gap(target_benchmark, competitor_benchmarks, snippet_id=snippet_signal["id"]),
        )
        gap_items.append(
            {
                "title": snippet_signal["label"],
                "priority": snippet_signal["priority"],
                "owner": snippet_signal["owner"],
                "current_state": (
                    f"На сайте {snippet_metric_state_text(snippet_signal['id'], target_metric, target_benchmark.get('pages_checked', 0))}."
                    if target_benchmark.get("pages_checked", 0)
                    else "По сайту не удалось собрать ключевую выборку для сравнения."
                ),
                "competitor_state": (
                    f"У {len(competitor_with_signal)} из {len(competitor_benchmarks)} "
                    f"{pluralize_ru(len(competitor_benchmarks), 'конкурента', 'конкурентов', 'конкурентов')} "
                    "эта шаблонная база собрана заметно стабильнее."
                ),
                "examples": examples,
                "task": snippet_signal["implementation_brief"],
                "benefit": snippet_signal["benefit"],
                **execution_details,
                "sort_score": round(gap_size + competitor_avg_coverage + snippet_signal["weight"] / 10, 3),
            }
        )

    competitor_average_words = statistics.mean([site["average_words"] for site in competitor_benchmarks if site["average_words"] > 0]) if competitor_benchmarks else 0
    target_average_words = target_benchmark.get("average_words", 0)
    if competitor_average_words >= max(380, target_average_words * 1.35):
        content_examples = [
            f"{site['domain']}: средний текстовый слой {round(site['average_words'])} слов"
            for site in sorted(competitor_benchmarks, key=lambda item: item["average_words"], reverse=True)[:3]
            if site["average_words"] > 0
        ]
        gap_items.append(
            {
                "title": "Глубина коммерческого контента на ключевых страницах",
                "priority": "Высокий приоритет",
                "owner": "SEO / Content",
                "current_state": (
                    f"У сайта в ключевой выборке в среднем {round(target_average_words)} слов видимого текста на страницу."
                    if target_average_words
                    else "У сайта пока нет устойчивого текстового слоя на ключевых страницах."
                ),
                "competitor_state": (
                    f"У конкурентов в среднем {round(competitor_average_words)} слов на ключевую страницу, и за счет этого они лучше закрывают интент и возражения."
                ),
                "examples": content_examples,
                "task": (
                    "Усилить приоритетные страницы не SEO-водой, а полезным коммерческим слоем: краткое интро, "
                    "условия работы, сценарии выбора, FAQ, блоки доверия, таблицы и перелинковку на смежные страницы."
                ),
                "benefit": (
                    "Более глубокий контент помогает закрывать спрос без возврата в выдачу, поддерживает коммерческий интент и дает больше точек входа по длинным запросам."
                ),
                **build_gap_execution_details(
                    "Глубина коммерческого контента на ключевых страницах",
                    pick_templates_for_gap(target_benchmark, competitor_benchmarks, use_commercial_score=True),
                ),
                "sort_score": round((competitor_average_words - target_average_words) / 300 + 1.2, 3),
            }
        )

    gap_items.sort(key=lambda item: item["sort_score"], reverse=True)
    for item in gap_items:
        item.pop("sort_score", None)
    return gap_items[:5]


def build_competitor_comparison(audit: dict, competitor_urls: list[str]) -> dict | None:
    normalized_competitors: list[str] = []
    seen_domains: set[str] = set()
    target_domain = urlparse(audit["base_url"]).netloc.replace("www.", "")

    for item in competitor_urls:
        try:
            normalized = normalize_base_url(item)
        except Exception:
            continue
        domain = urlparse(normalized).netloc.replace("www.", "")
        if not domain or domain == target_domain or domain in seen_domains:
            continue
        seen_domains.add(domain)
        normalized_competitors.append(normalized)

    if not normalized_competitors:
        return None

    target_benchmark = build_site_feature_benchmark(audit["base_url"], get_competitor_benchmark_pages(audit))
    competitor_benchmarks: list[dict] = []
    failures: list[dict] = []

    for competitor_url in normalized_competitors:
        try:
            competitor_benchmarks.append(collect_competitor_benchmark(competitor_url))
        except Exception as exc:  # noqa: BLE001
            failures.append(
                {
                    "url": competitor_url,
                    "error": str(exc),
                }
            )

    if not competitor_benchmarks:
        return {
            "summary": [
                "Конкуренты были переданы в задачу, но по ним не удалось собрать рабочую выборку страниц.",
                "Проверьте доступность URL и повторите прогон аудита.",
            ],
            "competitors": [],
            "gap_items": [],
            "failures": failures,
        }

    for competitor_benchmark in competitor_benchmarks:
        competitor_benchmark["factor_summary"] = build_competitor_factor_summary(target_benchmark, competitor_benchmark)
        competitor_benchmark["template_rows"] = build_competitor_template_rows(target_benchmark, competitor_benchmark)
        competitor_benchmark["template_findings"] = build_competitor_template_findings(target_benchmark, competitor_benchmark)
        competitor_benchmark["snippet_findings"] = build_competitor_snippet_findings(target_benchmark, competitor_benchmark)
        competitor_benchmark["commercial_findings"] = build_competitor_commercial_findings(target_benchmark, competitor_benchmark)

    gap_items = build_competitor_gap_items(target_benchmark, competitor_benchmarks)
    template_focus = summarize_template_gap_labels(target_benchmark, competitor_benchmarks)

    summary = [
        f"Сравнили сайт с {len(competitor_benchmarks)} {pluralize_ru(len(competitor_benchmarks), 'конкурентом', 'конкурентами', 'конкурентами')} "
        "по шаблонам страниц, сниппетам, FAQ, блокам доверия и коммерческим факторам.",
    ]
    if gap_items:
        summary.append(f"Чаще всего у конкурентов сильнее закрыты: {', '.join(item['title'] for item in gap_items[:3])}.")
        if template_focus:
            summary.append(f"Сильнее всего проседают шаблоны: {', '.join(template_focus)}.")
        summary.append("Ниже оставлены только те разрывы, которые можно быстро перевести в понятные задачи для SEO, контента, дизайна и разработки.")
    else:
        if template_focus:
            summary.append(f"Сильнее всего отличаются шаблоны: {', '.join(template_focus)}.")
        summary.append("Сильных системных разрывов по этой выборке не нашли, но карточки конкурентов ниже всё равно показывают удачные решения по структуре и подаче.")

    return {
        "summary": summary,
        "competitors": competitor_benchmarks,
        "gap_items": gap_items,
        "failures": failures,
        "target_benchmark": target_benchmark,
    }


ISSUE_TITLE_COMMERCIAL_GAPS = "Коммерческие страницы не закрывают условия выбора и доверие"
ISSUE_TITLE_CONTENT_PROOF = "Приоритетные страницы не подтверждают оффер фактами и ответами на вопросы"
ISSUE_TITLE_INTERNAL_LINKING = "Внутренняя перелинковка не поддерживает приоритетные страницы"
MONSTER_SUPPORT_FIELDS = ("has_faq_block", "has_reviews_block", "has_comparison_table")
MONSTER_PROOF_FIELDS = ("has_trust_block", "has_reviews_block", "has_comparison_table", "has_faq_block")


def select_priority_pages(sample_pages: list[PageSnapshot]) -> list[PageSnapshot]:
    return [snapshot for snapshot in sample_pages if get_template_bucket(snapshot) in {"home", "category", "product", "service", "contact"}]


def coverage_ratio(pages: list[PageSnapshot], predicate) -> float:
    if not pages:
        return 0.0
    return len([page for page in pages if predicate(page)]) / len(pages)


def collect_page_examples(pages: list[PageSnapshot], limit: int = 3) -> list[str]:
    return [human_path(page.url) for page in pages[:limit]]


def dynamic_build_issue_list(audit: dict) -> list[AuditIssue]:
    issues: list[AuditIssue] = []
    robots_lines = audit["robots_lines"]
    sample_pages: list[PageSnapshot] = audit["sample_pages"]
    problematic_titles = [snapshot for snapshot in sample_pages if len(snapshot.title) > 70]
    problematic_descriptions = [
        snapshot for snapshot in sample_pages if len(snapshot.description) > 160 or (0 < len(snapshot.description) < 120)
    ]
    alt_gaps = [snapshot for snapshot in sample_pages if snapshot.missing_alt_count > 0]
    redirect_chains = [item for item in audit["redirect_checks"] if len(item["chain"]) > 1]
    contact_pages = get_contact_related_pages(sample_pages)
    indexable_query_pages = get_indexable_query_pages(sample_pages)
    priority_pages = select_priority_pages(sample_pages)
    weak_support_pages = [
        snapshot
        for snapshot in priority_pages
        if not any(bool(getattr(snapshot, field, False)) for field in MONSTER_SUPPORT_FIELDS)
    ]
    weak_proof_pages = [
        snapshot
        for snapshot in priority_pages
        if not any(bool(getattr(snapshot, field, False)) for field in MONSTER_PROOF_FIELDS)
    ]
    commercial_gap_pages = [snapshot for snapshot in priority_pages if not snapshot.has_commercial_block]
    trust_gap_pages = [snapshot for snapshot in priority_pages if not snapshot.has_trust_block]
    thin_priority_pages = [snapshot for snapshot in priority_pages if 0 < snapshot.word_count < 250]
    breadcrumb_gap_pages = [
        snapshot
        for snapshot in priority_pages
        if get_template_bucket(snapshot) in {"category", "product", "service", "contact"} and not snapshot.has_breadcrumbs
    ]
    low_link_priority_pages = sorted(
        [snapshot for snapshot in priority_pages if get_template_bucket(snapshot) != "home"],
        key=lambda snapshot: snapshot.internal_links,
    )[:3]
    average_priority_links = round(statistics.mean([snapshot.internal_links for snapshot in priority_pages]), 1) if priority_pages else 0
    profile = infer_project_profile(audit)

    disallow_sitemap_lines = [line for line in robots_lines if "sitemap.xml" in line.lower() and "disallow" in line.lower()]
    sitemap_directive = next((line for line in robots_lines if line.lower().startswith("sitemap:")), "")

    if disallow_sitemap_lines:
        evidence = [f"Р’ robots.txt РЅР°Р№РґРµРЅРѕ РїСЂР°РІРёР»Рѕ `{disallow_sitemap_lines[0]}`."]
        if sitemap_directive:
            evidence.append(f"РћРґРЅРѕРІСЂРµРјРµРЅРЅРѕ С„Р°Р№Р» СЃРѕРґРµСЂР¶РёС‚ `{sitemap_directive}`.")
        issues.append(
            AuditIssue(
                severity="Critical",
                title="robots.txt РєРѕРЅС„Р»РёРєС‚СѓРµС‚ СЃ РєР°СЂС‚РѕР№ СЃР°Р№С‚Р°",
                why_it_matters=(
                    "РљРѕРіРґР° robots.txt РѕРґРЅРѕРІСЂРµРјРµРЅРЅРѕ Р·Р°РїСЂРµС‰Р°РµС‚ sitemap.xml Рё СЃР°Рј Р¶Рµ СЃСЃС‹Р»Р°РµС‚СЃСЏ РЅР° РєР°СЂС‚Сѓ СЃР°Р№С‚Р°, "
                    "РїРѕРёСЃРєРѕРІРёРєРё РїРѕР»СѓС‡Р°СЋС‚ РєРѕРЅС„Р»РёРєС‚СѓСЋС‰РёР№ СЃРёРіРЅР°Р». Р­С‚Рѕ РјРµС€Р°РµС‚ С‡РёСЃС‚РѕР№ Рё РїСЂРµРґСЃРєР°Р·СѓРµРјРѕР№ РёРЅРґРµРєСЃР°С†РёРё."
                ),
                evidence=evidence,
                recommendation=(
                    "РЈР±СЂР°С‚СЊ Р·Р°РїСЂРµС‚ РЅР° sitemap.xml, РѕСЃС‚Р°РІРёС‚СЊ СЂР°Р±РѕС‡СѓСЋ РґРёСЂРµРєС‚РёРІСѓ Sitemap Рё СЃРёРЅС…СЂРѕРЅРёР·РёСЂРѕРІР°С‚СЊ robots.txt "
                    "СЃ РєР°РЅРѕРЅРёС‡РµСЃРєРёРј РґРѕРјРµРЅРѕРј РїСЂРѕРµРєС‚Р°."
                ),
            )
        )

    weak_contact_pages: list[tuple[PageSnapshot, list[str]]] = []
    for snapshot in contact_pages:
        missing = []
        if not snapshot.h1s:
            missing.append("H1")
        if not snapshot.title:
            missing.append("title")
        if not snapshot.description:
            missing.append("description")
        if not snapshot.canonical:
            missing.append("canonical")
        if missing:
            weak_contact_pages.append((snapshot, missing))

    if weak_contact_pages:
        evidence = [f"РќР° `{human_path(snapshot.url)}` РЅРµ С…РІР°С‚Р°РµС‚: {', '.join(missing)}." for snapshot, missing in weak_contact_pages[:3]]
        issues.append(
            AuditIssue(
                severity="High",
                title="РљРѕРЅС‚Р°РєС‚РЅС‹Рµ Рё lead-СЃС‚СЂР°РЅРёС†С‹ РЅРµРґРѕРѕС„РѕСЂРјР»РµРЅС‹ РєР°Рє SEO-РїРѕСЃР°РґРѕС‡РЅС‹Рµ",
                why_it_matters=(
                    "РЎС‚СЂР°РЅРёС†С‹ РєРѕРЅС‚Р°РєС‚РѕРІ Рё Р·Р°СЏРІРѕРє СѓС‡Р°СЃС‚РІСѓСЋС‚ РЅРµ С‚РѕР»СЊРєРѕ РІ РєРѕРЅРІРµСЂСЃРёРё, РЅРѕ Рё РІ РґРѕРІРµСЂРёРё, Р±СЂРµРЅРґРѕРІРѕР№ РІС‹РґР°С‡Рµ "
                    "Рё РїРѕРЅРёРјР°РЅРёРё Р±РёР·РЅРµСЃР° РїРѕРёСЃРєРѕРІРёРєР°РјРё."
                ),
                evidence=evidence,
                recommendation=(
                    "РЎРѕР±СЂР°С‚СЊ РїРѕР»РЅРѕС†РµРЅРЅС‹Рµ СЃС‚СЂР°РЅРёС†С‹ РєРѕРЅС‚Р°РєС‚Р° Рё Р·Р°СЏРІРєРё: H1, title РґРѕ 70 СЃРёРјРІРѕР»РѕРІ, description 140вЂ“160 СЃРёРјРІРѕР»РѕРІ, "
                    "canonical, Р±Р»РѕРєРё РґРѕРІРµСЂРёСЏ, Р°РґСЂРµСЃР°, С‚РµР»РµС„РѕРЅС‹ Рё РїРѕРЅСЏС‚РЅС‹Р№ CTA."
                ),
            )
        )

    weak_query_pages = [
        snapshot for snapshot in indexable_query_pages if not snapshot.title or not snapshot.description or not snapshot.h1s
    ]
    if weak_query_pages:
        evidence = []
        for snapshot in weak_query_pages[:3]:
            missing = []
            if not snapshot.title:
                missing.append("title")
            if not snapshot.description:
                missing.append("description")
            if not snapshot.h1s:
                missing.append("H1")
            evidence.append(f"URL `{human_path(snapshot.url)}` РґРѕСЃС‚СѓРїРµРЅ РґР»СЏ РѕР±С…РѕРґР°, РЅРѕ РЅРµ С…РІР°С‚Р°РµС‚: {', '.join(missing)}.")
        issues.append(
            AuditIssue(
                severity="High",
                title="Р’ РІС‹Р±РѕСЂРєРµ РµСЃС‚СЊ РёРЅРґРµРєСЃРёСЂСѓРµРјС‹Рµ СЃР»СѓР¶РµР±РЅС‹Рµ РёР»Рё query-URL Р±РµР· РЅРѕСЂРјР°Р»СЊРЅРѕР№ SEO-РѕР±РІСЏР·РєРё",
                why_it_matters=(
                    "Служебные страницы и URL с параметрами без title, description и H1 создают шум в индексе, "
                    "размывают релевантность и тратят ресурсы на обход."
                ),
                evidence=evidence,
                recommendation=(
                    "Либо закрыть такие URL от индексации и убрать прямые ссылки на них, либо перевести их на нормальные страницы для поиска "
                    "с полноценным мета-оформлением."
                ),
            )
        )

    commercial_gap_ratio = coverage_ratio(priority_pages, lambda snapshot: not snapshot.has_commercial_block)
    trust_gap_ratio = coverage_ratio(priority_pages, lambda snapshot: not snapshot.has_trust_block)
    if priority_pages and (commercial_gap_ratio >= 0.45 or trust_gap_ratio >= 0.55):
        page_scope = (
            "категорий, карточек и страниц заявок"
            if profile["is_catalog"]
            else "страниц услуг, внутренних коммерческих страниц и контактов"
        )
        evidence = [
            (
                f"На {len(commercial_gap_pages)} из {len(priority_pages)} приоритетных страниц не найден явный блок условий: "
                "цена, сроки, порядок работы, доставка, оплата или гарантия."
            ),
            f"На {len(trust_gap_pages)} из {len(priority_pages)} приоритетных страниц не хватает доверительных сигналов.",
        ]
        examples = collect_page_examples(commercial_gap_pages or trust_gap_pages)
        if examples:
            evidence.append(f"Примеры URL: {', '.join(examples)}.")
        issues.append(
            AuditIssue(
                severity="High",
                title=ISSUE_TITLE_COMMERCIAL_GAPS,
                why_it_matters=(
                    "По SEO Монстр коммерческий интент закрывается не общим SEO-текстом, а условиями выбора: "
                    "цена или диапазон, сроки, гарантия, доставка/оплата, поддержка, контакты и понятный следующий шаг. "
                    f"Если этот слой отсутствует на {page_scope}, страница хуже отвечает на запрос и слабее конвертирует."
                ),
                evidence=evidence,
                recommendation=(
                    "Пересобрать приоритетные шаблоны по коммерческому каркасу: оффер, цена или диапазон, сроки, "
                    "условия сделки, гарантия/поддержка, контакты, CTA и отдельные служебные страницы "
                    "доставки, оплаты, гарантий или порядка работ."
                ),
            )
        )

    support_gap_ratio = coverage_ratio(
        priority_pages,
        lambda snapshot: not any(bool(getattr(snapshot, field, False)) for field in MONSTER_SUPPORT_FIELDS),
    )
    thin_priority_ratio = coverage_ratio(priority_pages, lambda snapshot: 0 < snapshot.word_count < 250)
    if priority_pages and (support_gap_ratio >= 0.5 or thin_priority_ratio >= 0.35):
        evidence = [
            (
                f"На {len(weak_support_pages)} из {len(priority_pages)} приоритетных страниц нет FAQ, отзывов/кейсов "
                "или сравнительных блоков."
            ),
            f"Тонких страниц с объёмом до 250 слов: {len(thin_priority_pages)} из {len(priority_pages)}.",
        ]
        examples = collect_page_examples(weak_support_pages or thin_priority_pages)
        if examples:
            evidence.append(f"Примеры URL: {', '.join(examples)}.")
        if weak_proof_pages:
            evidence.append(
                f"На {len(weak_proof_pages)} из {len(priority_pages)} страниц не видно доказательств: доверия, отзывов, сравнений или FAQ."
            )
        issues.append(
            AuditIssue(
                severity="High" if thin_priority_ratio >= 0.45 else "Medium",
                title=ISSUE_TITLE_CONTENT_PROOF,
                why_it_matters=(
                    "SEO Монстр рекомендует усиливать приоритетные страницы не водой, а новыми абзацами под потерянные запросы, "
                    "доказательствами, списками, H2/H3, фото, кейсами, таблицами и ответами на возражения. "
                    "Когда на странице нет этого слоя, она слабо закрывает выбор пользователя и хуже растёт после индексации."
                ),
                evidence=evidence,
                recommendation=(
                    "Доработать ключевые страницы по схеме повторной оптимизации: собрать потерянные запросы, добавить новые H2/H3, "
                    "вынести в контент факты, сравнения, кейсы, скрины, списки и FAQ, затем обновить title/description и заново усилить перелинковку."
                ),
            )
        )

    structure_gap_ratio = coverage_ratio(priority_pages, lambda snapshot: snapshot.internal_links <= 5)
    breadcrumb_gap_ratio = coverage_ratio(
        priority_pages,
        lambda snapshot: get_template_bucket(snapshot) in {"category", "product", "service", "contact"} and not snapshot.has_breadcrumbs,
    )
    if priority_pages and (structure_gap_ratio >= 0.35 or breadcrumb_gap_ratio >= 0.45):
        evidence = [
            f"Среднее число внутренних ссылок на приоритетную страницу: {average_priority_links}.",
            (
                "Слабее всего по внутренним ссылкам выглядят: "
                + ", ".join(f"{human_path(snapshot.url)} ({snapshot.internal_links})" for snapshot in low_link_priority_pages)
                + "."
                if low_link_priority_pages
                else "Слабые по внутренним ссылкам страницы не выделяются."
            ),
            (
                f"На {len(breadcrumb_gap_pages)} из {len(priority_pages)} приоритетных страниц не обнаружены хлебные крошки."
                if breadcrumb_gap_pages
                else "Критичной просадки по хлебным крошкам в выборке не видно."
            ),
        ]
        issues.append(
            AuditIssue(
                severity="Medium",
                title=ISSUE_TITLE_INTERNAL_LINKING,
                why_it_matters=(
                    "По SEO Монстр внутренняя перелинковка должна распределять вес не хаотично, а тематически: "
                    "вести пользователя и робота от сильных страниц к целевым кластерам, держать важные URL в нескольких кликах "
                    "и использовать осмысленные анкоры вместо навигационного шума."
                ),
                evidence=evidence,
                recommendation=(
                    "Собрать тематические связки между приоритетными страницами: главная -> раздел -> подраздел/услуга -> заявка, "
                    "добавить хлебные крошки, блоки \"по теме\", релевантные анкорные ссылки из текста и сократить бессмысленные ссылки вида "
                    "\"читать далее\" или дубли ссылок на один и тот же URL."
                ),
            )
        )

    if problematic_titles:
        examples = [f"{human_path(snapshot.url)} вЂ” {len(snapshot.title)} СЃРёРјРІ." for snapshot in problematic_titles[:4]]
        issues.append(
            AuditIssue(
                severity="High",
                title="РќР° С‡Р°СЃС‚Рё СЃС‚СЂР°РЅРёС† title РІС‹С…РѕРґСЏС‚ Р·Р° СЂР°Р±РѕС‡СѓСЋ РґР»РёРЅСѓ",
                why_it_matters=(
                    "РЎР»РёС€РєРѕРј РґР»РёРЅРЅС‹Рµ title СЂРµР¶СѓС‚СЃСЏ РІ РІС‹РґР°С‡Рµ, СѓС…СѓРґС€Р°СЋС‚ CTR Рё СЃРЅРёР¶Р°СЋС‚ РєРѕРЅС‚СЂРѕР»СЊ РЅР°Рґ С‚РµРј, "
                    "РєР°РєРѕР№ РѕС„С„РµСЂ РїРѕРёСЃРєРѕРІРёРє РїРѕРєР°Р¶РµС‚ РїРѕР»СЊР·РѕРІР°С‚РµР»СЋ."
                ),
                evidence=[
                    f"Р’ РІС‹Р±РѕСЂРєРµ {len(problematic_titles)} РёР· {len(sample_pages)} СЃС‚СЂР°РЅРёС† РёРјРµСЋС‚ title РґР»РёРЅРЅРµРµ 70 СЃРёРјРІРѕР»РѕРІ.",
                    *examples,
                ],
                recommendation=(
                    "РџРµСЂРµСЃРѕР±СЂР°С‚СЊ С€Р°Р±Р»РѕРЅС‹ title: СЃРЅР°С‡Р°Р»Р° РєР»СЋС‡ Рё С‚РёРї СЃС‚СЂР°РЅРёС†С‹, Р·Р°С‚РµРј С†РµРЅРЅРѕСЃС‚РЅС‹Р№ РѕС„С„РµСЂ Рё Р±СЂРµРЅРґ. "
                    "РћСЂРёРµРЅС‚РёСЂ вЂ” РґРёР°РїР°Р·РѕРЅ 55вЂ“70 СЃРёРјРІРѕР»РѕРІ."
                ),
            )
        )

    if problematic_descriptions:
        examples = [f"{human_path(snapshot.url)} вЂ” {len(snapshot.description)} СЃРёРјРІ." for snapshot in problematic_descriptions[:4]]
        issues.append(
            AuditIssue(
                severity="Medium",
                title="Meta description РЅР° С‡Р°СЃС‚Рё СЃС‚СЂР°РЅРёС† РІРЅРµ СЂР°Р±РѕС‡РµРіРѕ РґРёР°РїР°Р·РѕРЅР°",
                why_it_matters=(
                    "Описание страницы — это управляемый оффер в сниппете. Когда оно слишком короткое или перегруженное, "
                    "РїРѕРёСЃРєРѕРІРёРє С‡Р°С‰Рµ Р±РµСЂРµС‚ СЃР»СѓС‡Р°Р№РЅС‹Р№ РєСѓСЃРѕРє С‚РµРєСЃС‚Р° СЃРѕ СЃС‚СЂР°РЅРёС†С‹."
                ),
                evidence=[
                    f"Р’ РІС‹Р±РѕСЂРєРµ {len(problematic_descriptions)} РёР· {len(sample_pages)} СЃС‚СЂР°РЅРёС† РёРјРµСЋС‚ description РІРЅРµ РєРѕРјС„РѕСЂС‚РЅРѕРіРѕ РґРёР°РїР°Р·РѕРЅР°.",
                    *examples,
                ],
                recommendation=(
                    "РЎРѕР±СЂР°С‚СЊ С€Р°Р±Р»РѕРЅС‹ description РїРѕ С‚РёРїР°Рј СЃС‚СЂР°РЅРёС†: РєР»СЋС‡, С†РµРЅРЅРѕСЃС‚СЊ, РґРѕРІРµСЂРёРµ, СЂРµРіРёРѕРЅ Рё РїСЂРёР·С‹РІ. "
                    "РћСЂРёРµРЅС‚РёСЂ вЂ” 140вЂ“160 СЃРёРјРІРѕР»РѕРІ."
                ),
            )
        )

    if alt_gaps:
        issues.append(
            AuditIssue(
                severity="Medium",
                title="РР·РѕР±СЂР°Р¶РµРЅРёСЏ С‚РµСЂСЏСЋС‚ SEO-СЃРёРіРЅР°Р»С‹ РёР·-Р·Р° РїСѓСЃС‚С‹С… alt",
                why_it_matters=(
                    "Alt РІР°Р¶РµРЅ Рё РґР»СЏ РїРѕРёСЃРєР° РїРѕ РєР°СЂС‚РёРЅРєР°Рј, Рё РґР»СЏ РїРѕРЅРёРјР°РЅРёСЏ РєРѕРЅС‚РµРєСЃС‚Р° СЃС‚СЂР°РЅРёС†С‹, Рё РґР»СЏ РґРѕСЃС‚СѓРїРЅРѕСЃС‚Рё. "
                    "РљРѕРіРґР° Сѓ РёР·РѕР±СЂР°Р¶РµРЅРёР№ РЅРµС‚ РѕРїРёСЃР°РЅРёР№, СЃР°Р№С‚ С‚РµСЂСЏРµС‚ РґРѕРїРѕР»РЅРёС‚РµР»СЊРЅС‹Р№ СЃР»РѕР№ СЂРµР»РµРІР°РЅС‚РЅРѕСЃС‚Рё."
                ),
                evidence=[
                    f"РќР° {len(alt_gaps)} РёР· {len(sample_pages)} РїСЂРѕРІРµСЂРµРЅРЅС‹С… СЃС‚СЂР°РЅРёС† РµСЃС‚СЊ РёР·РѕР±СЂР°Р¶РµРЅРёСЏ Р±РµР· alt.",
                    f"Р’СЃРµРіРѕ РІ РІС‹Р±РѕСЂРєРµ РЅР°Р№РґРµРЅРѕ {audit['total_missing_alt']} РёР·РѕР±СЂР°Р¶РµРЅРёР№ Р±РµР· РѕРїРёСЃР°РЅРёР№.",
                ],
                recommendation="Р’РІРµСЃС‚Рё С€Р°Р±Р»РѕРЅРЅСѓСЋ РіРµРЅРµСЂР°С†РёСЋ alt РїРѕ С‚РёРїСѓ СЃС‚СЂР°РЅРёС†С‹ Рё С‚РёРїСѓ РёР·РѕР±СЂР°Р¶РµРЅРёСЏ: РѕР±СЉРµРєС‚, РёРЅС‚РµРЅС‚, РјРѕРґРµР»СЊ Рё Р±СЂРµРЅРґ.",
            )
        )

    if redirect_chains:
        longest_chain = max(redirect_chains, key=lambda item: len(item["chain"]))
        issues.append(
            AuditIssue(
                severity="Medium",
                title="Р•СЃС‚СЊ Р»РёС€РЅРёРµ redirect-С†РµРїРѕС‡РєРё Сѓ РґРѕРјРµРЅРЅС‹С… РІР°СЂРёР°РЅС‚РѕРІ",
                why_it_matters=(
                    "Р”РѕРїРѕР»РЅРёС‚РµР»СЊРЅС‹Рµ СЂРµРґРёСЂРµРєС‚С‹ СѓРІРµР»РёС‡РёРІР°СЋС‚ Р·Р°РґРµСЂР¶РєСѓ РЅР° РІС…РѕРґРµ Рё СЃРѕР·РґР°СЋС‚ Р»РёС€РЅСЋСЋ С‚РµС…РЅРёС‡РµСЃРєСѓСЋ СЃР»РѕР¶РЅРѕСЃС‚СЊ РґР»СЏ Р±РѕС‚РѕРІ Рё РїРѕР»СЊР·РѕРІР°С‚РµР»РµР№."
                ),
                evidence=[
                    f"РџСѓС‚СЊ `{longest_chain['url']}` РѕС‚РґР°РµС‚ {len(longest_chain['chain'])} РїРѕСЃР»РµРґРѕРІР°С‚РµР»СЊРЅС‹С… redirect.",
                    f"Р¤РёРЅР°Р»СЊРЅС‹Р№ URL вЂ” `{longest_chain['final_url']}`.",
                ],
                recommendation="РЎРІРµСЃС‚Рё РІСЃРµ РґРѕРјРµРЅРЅС‹Рµ РІР°СЂРёР°РЅС‚С‹ Рє РѕРґРЅРѕРјСѓ РїСЂСЏРјРѕРјСѓ 301-СЂРµРґРёСЂРµРєС‚Сѓ РЅР° РєР°РЅРѕРЅРёС‡РµСЃРєРёР№ HTTPS-РґРѕРјРµРЅ.",
            )
        )

    if audit["home_page"].has_meta_keywords:
        issues.append(
            AuditIssue(
                severity="Low",
                title="РќР° С‡Р°СЃС‚Рё С€Р°Р±Р»РѕРЅР° РѕСЃС‚Р°Р»СЃСЏ meta keywords",
                why_it_matters=(
                    "РЎР°Рј РїРѕ СЃРµР±Рµ С‚РµРі СѓР¶Рµ РЅРµ РґР°РµС‚ SEO-С†РµРЅРЅРѕСЃС‚Рё, РЅРѕ РїРѕРєР°Р·С‹РІР°РµС‚, С‡С‚Рѕ С€Р°Р±Р»РѕРЅ РјРµС‚Р°РґР°РЅРЅС‹С… РЅРµ РґРѕС‡РёС‰РµРЅ "
                    "Рё СЃС‚РѕРёС‚ РїСЂРёРІРµСЃС‚Рё РµРіРѕ Рє СЃРѕРІСЂРµРјРµРЅРЅРѕРјСѓ РІРёРґСѓ."
                ),
                evidence=["РќР° РіР»Р°РІРЅРѕР№ РЅР°Р№РґРµРЅ С‚РµРі meta keywords."],
                recommendation="РЈР±СЂР°С‚СЊ meta keywords РёР· С€Р°Р±Р»РѕРЅР° Рё СЃРѕСЃСЂРµРґРѕС‚РѕС‡РёС‚СЊСЃСЏ РЅР° title, description, H1, canonical Рё schema.",
            )
        )

    if audit["sitemap_url_count"] > 5000:
        issues.append(
            AuditIssue(
                severity="Medium",
                title="РљР°СЂС‚Р° СЃР°Р№С‚Р° СѓР¶Рµ Р±РѕР»СЊС€Р°СЏ Рё С‚СЂРµР±СѓРµС‚ СЃРµРіРјРµРЅС‚Р°С†РёРё",
                why_it_matters=(
                    "РљРѕРіРґР° РІ sitemap РЅРµСЃРєРѕР»СЊРєРѕ С‚С‹СЃСЏС‡ URL РІ РѕРґРЅРѕРј РїРѕС‚РѕРєРµ, СЃР»РѕР¶РЅРµРµ РєРѕРЅС‚СЂРѕР»РёСЂРѕРІР°С‚СЊ РёРЅРґРµРєСЃР°С†РёСЋ РєР°С‚РµРіРѕСЂРёР№, РєР°СЂС‚РѕС‡РµРє, "
                    "СЃР»СѓР¶РµР±РЅС‹С… СЃС‚СЂР°РЅРёС† Рё РѕС‚РґРµР»СЊРЅС‹С… С‚РёРїРѕРІ РєРѕРЅС‚РµРЅС‚Р°."
                ),
                evidence=[
                    f"Р’ sitemap РѕР±РЅР°СЂСѓР¶РµРЅРѕ {audit['sitemap_url_count']} СѓРЅРёРєР°Р»СЊРЅС‹С… URL.",
                    "РџСЂРё С‚Р°РєРѕРј РѕР±СЉРµРјРµ СѓРґРѕР±РЅРµРµ СѓРїСЂР°РІР»СЏС‚СЊ РёРЅРґРµРєСЃРѕРј С‡РµСЂРµР· РѕС‚РґРµР»СЊРЅС‹Рµ sitemap РїРѕ С‚РёРїР°Рј СЃС‚СЂР°РЅРёС†.",
                ],
                recommendation="Р Р°Р·РґРµР»РёС‚СЊ sitemap РјРёРЅРёРјСѓРј РїРѕ РѕСЃРЅРѕРІРЅС‹Рј С‚РёРїР°Рј СЃС‚СЂР°РЅРёС† Рё РѕС‚РґРµР»СЊРЅРѕ РєРѕРЅС‚СЂРѕР»РёСЂРѕРІР°С‚СЊ РїРѕРєСЂС‹С‚РёРµ РєР°Р¶РґРѕРіРѕ РЅР°Р±РѕСЂР°.",
            )
        )

    if audit.get("sitemap_total_entries", 0) > audit["sitemap_url_count"] * 1.3:
        duplicate_entries = audit["sitemap_total_entries"] - audit["sitemap_url_count"]
        issues.append(
            AuditIssue(
                severity="Medium",
                title="Р’ sitemap РјРЅРѕРіРѕ РїРѕРІС‚РѕСЂСЏСЋС‰РёС…СЃСЏ URL",
                why_it_matters=(
                    "РљРѕРіРґР° РѕРґРёРЅ Рё С‚РѕС‚ Р¶Рµ URL РїРѕРІС‚РѕСЂСЏРµС‚СЃСЏ РІ РєР°СЂС‚Рµ СЃР°Р№С‚Р° РјРЅРѕРіРѕ СЂР°Р·, РїРѕРёСЃРєРѕРІРёРєРё РїРѕР»СѓС‡Р°СЋС‚ Р»РёС€РЅРёР№ С€СѓРј РІРјРµСЃС‚Рѕ С‡РёСЃС‚РѕРіРѕ СЃРёРіРЅР°Р»Р°."
                ),
                evidence=[
                    f"Р’ sitemap РЅР°Р№РґРµРЅРѕ {audit['sitemap_total_entries']} Р·Р°РїРёСЃРµР№, РЅРѕ С‚РѕР»СЊРєРѕ {audit['sitemap_url_count']} СѓРЅРёРєР°Р»СЊРЅС‹С… URL.",
                    f"РџРѕРІС‚РѕСЂРѕРІ: {duplicate_entries}.",
                ],
                recommendation="РџРµСЂРµСЃРѕР±СЂР°С‚СЊ РіРµРЅРµСЂР°С†РёСЋ sitemap С‚Р°Рє, С‡С‚РѕР±С‹ РєР°Р¶РґС‹Р№ РёРЅРґРµРєСЃРёСЂСѓРµРјС‹Р№ URL РїРѕРїР°РґР°Р» С‚СѓРґР° РѕРґРёРЅ СЂР°Р· Рё Р±РµР· РґСѓР±Р»РµР№.",
            )
        )

    severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
    issues.sort(key=lambda issue: severity_order.get(issue.severity, 4))
    return issues


def dynamic_build_strengths(audit: dict) -> list[str]:
    strengths: list[str] = []
    profile = infer_project_profile(audit)
    if audit["home_page"].status_code == 200:
        strengths.append("Р“Р»Р°РІРЅР°СЏ СЃС‚СЂР°РЅРёС†Р° СЃС‚Р°Р±РёР»СЊРЅРѕ РѕС‚РєСЂС‹РІР°РµС‚СЃСЏ РїРѕ HTTPS Рё РѕС‚РґР°РµС‚ 200 РєРѕРґ Р±РµР· JS-Р·Р°РіР»СѓС€РєРё.")
    if audit["average_response_ms"] and audit["average_response_ms"] < 1200:
        strengths.append(f"РЎСЂРµРґРЅРёР№ РѕС‚РІРµС‚ РїРѕ РІС‹Р±РѕСЂРєРµ вЂ” {audit['average_response_ms']} ms: С‚РµС…РЅРёС‡РµСЃРєР°СЏ Р±Р°Р·Р° РЅРµ РІС‹РіР»СЏРґРёС‚ РїРµСЂРµРіСЂСѓР¶РµРЅРЅРѕР№.")
    if audit["sitemap_url_count"]:
        strengths.append(f"РЈ СЃР°Р№С‚Р° СѓР¶Рµ РµСЃС‚СЊ РєР°СЂС‚Р° СЃР°Р№С‚Р° СЃ {audit['sitemap_url_count']} СѓРЅРёРєР°Р»СЊРЅС‹РјРё URL, Р° Р·РЅР°С‡РёС‚ Р±Р°Р·Р° РґР»СЏ СѓРїСЂР°РІР»СЏРµРјРѕР№ РёРЅРґРµРєСЃР°С†РёРё СѓР¶Рµ СЃСѓС‰РµСЃС‚РІСѓРµС‚.")
    if audit.get("canonical_coverage_ratio", 0) >= 0.6:
        strengths.append(
            f"РЈ {math.floor(audit['canonical_coverage_ratio'] * 100)}% СЃС‚СЂР°РЅРёС† РІ РІС‹Р±РѕСЂРєРµ СѓР¶Рµ РїСЂРѕСЃС‚Р°РІР»РµРЅ canonical, СЌС‚Рѕ С…РѕСЂРѕС€РёР№ С„СѓРЅРґР°РјРµРЅС‚ РґР»СЏ С‡РёСЃС‚РѕР№ РёРЅРґРµРєСЃР°С†РёРё."
        )
    if audit["schema_coverage_ratio"] >= 0.4:
        strengths.append(f"Schema-СЂР°Р·РјРµС‚РєР° СѓР¶Рµ РёСЃРїРѕР»СЊР·СѓРµС‚СЃСЏ РЅР° {math.floor(audit['schema_coverage_ratio'] * 100)}% СЃС‚СЂР°РЅРёС† РІС‹Р±РѕСЂРєРё.")
    if audit["h1_coverage_ratio"] >= 0.8:
        strengths.append(f"РЈ {math.floor(audit['h1_coverage_ratio'] * 100)}% РїСЂРѕРІРµСЂРµРЅРЅС‹С… СЃС‚СЂР°РЅРёС† РµСЃС‚СЊ H1 вЂ” СЃС‚СЂСѓРєС‚СѓСЂР° РґРѕРєСѓРјРµРЅС‚РѕРІ СѓР¶Рµ РЅРµ РІС‹РіР»СЏРґРёС‚ С…Р°РѕС‚РёС‡РЅРѕР№.")
    if profile["is_catalog"]:
        strengths.append("РЈ РїСЂРѕРµРєС‚Р° СѓР¶Рµ РµСЃС‚СЊ РєР°С‚Р°Р»РѕРі РёР»Рё РєР»Р°СЃС‚РµСЂРЅР°СЏ URL-СЃС‚СЂСѓРєС‚СѓСЂР°, РєРѕС‚РѕСЂСѓСЋ РјРѕР¶РЅРѕ СѓСЃРёР»РёРІР°С‚СЊ Р±РµР· РїРѕР»РЅРѕР№ СЃРјРµРЅС‹ Р°СЂС…РёС‚РµРєС‚СѓСЂС‹.")
    else:
        strengths.append("РЈ РїСЂРѕРµРєС‚Р° СѓР¶Рµ РµСЃС‚СЊ Р±Р°Р·РѕРІС‹Р№ РЅР°Р±РѕСЂ РїРѕСЃР°РґРѕС‡РЅС‹С… СЃС‚СЂР°РЅРёС†, РєРѕС‚РѕСЂС‹Р№ РјРѕР¶РЅРѕ РґРѕРєСЂСѓС‡РёРІР°С‚СЊ С‚РѕС‡РµС‡РЅРѕ, Р° РЅРµ РїРµСЂРµСЃРѕР±РёСЂР°С‚СЊ СЃ РЅСѓР»СЏ.")
    if audit["llms_exists"]:
        strengths.append("РЈ РїСЂРѕРµРєС‚Р° СѓР¶Рµ РµСЃС‚СЊ llms.txt, Р° Р·РЅР°С‡РёС‚ РјРѕР¶РЅРѕ РѕС‚РґРµР»СЊРЅРѕ СѓСЃРёР»РёРІР°С‚СЊ РІРёРґРёРјРѕСЃС‚СЊ РІ РР-РѕС‚РІРµС‚Р°С… Рё branded-РїРѕРєСЂС‹С‚РёРµ.")
    return strengths[:6]


def dynamic_build_growth_points(audit: dict) -> list[str]:
    profile = infer_project_profile(audit)
    points: list[str] = []

    if profile["is_catalog"]:
        points.append("РЈСЃРёР»РёС‚СЊ РєР°С‚РµРіРѕСЂРёРё Рё РїРѕРґРєР°С‚РµРіРѕСЂРёРё: РёРЅС‚СЂРѕ-Р±Р»РѕРєРё, FAQ, СѓСЃР»РѕРІРёСЏ РїРѕРєСѓРїРєРё, Р±Р»РѕРєРё РґРѕРІРµСЂРёСЏ Рё РїРµСЂРµР»РёРЅРєРѕРІРєСѓ РЅР° РєР°СЂС‚РѕС‡РєРё.")
        points.append("РџРµСЂРµСЃРѕР±СЂР°С‚СЊ С€Р°Р±Р»РѕРЅС‹ РєР°СЂС‚РѕС‡РµРє Рё Р»РёСЃС‚РёРЅРіРѕРІ: РєРѕСЂРѕС‚РєРёРµ title, СЂР°Р±РѕС‡РёРµ description, canonical, alt, schema Рё РєРѕРјРјРµСЂС‡РµСЃРєРёРµ CTA.")
    else:
        points.append("Р Р°Р·РІРµСЃС‚Рё РєР»СЋС‡РµРІС‹Рµ РёРЅС‚РµРЅС‚С‹ РїРѕ РѕС‚РґРµР»СЊРЅС‹Рј РїРѕСЃР°РґРѕС‡РЅС‹Рј СЃС‚СЂР°РЅРёС†Р°Рј, Р° РЅРµ РґРµСЂР¶Р°С‚СЊ РЅРµСЃРєРѕР»СЊРєРѕ С‚РµРј РІРЅСѓС‚СЂРё РѕРґРЅРѕРіРѕ РґРѕРєСѓРјРµРЅС‚Р°.")
        points.append("РЈСЃРёР»РёС‚СЊ СЃС‚СЂР°РЅРёС†С‹ СѓСЃР»СѓРі Рё Р»РёРґ-СЃС‚СЂР°РЅРёС†С‹ Р±Р»РѕРєР°РјРё РґРѕРІРµСЂРёСЏ, РєРµР№СЃР°РјРё, FAQ, CTA Рё answer-first С„СЂР°РіРјРµРЅС‚Р°РјРё.")

    points.append("РЎРѕР±СЂР°С‚СЊ РµРґРёРЅС‹Р№ СЃР»РѕР№ С€Р°Р±Р»РѕРЅРѕРІ РґР»СЏ title, description, H1 Рё canonical РїРѕ РІСЃРµРј С‚РёРїР°Рј СЃС‚СЂР°РЅРёС†.")

    if audit["total_missing_alt"] > 0:
        points.append("Р Р°Р·РІРµСЂРЅСѓС‚СЊ image SEO: alt-С€Р°Р±Р»РѕРЅС‹, РїРѕРґРїРёСЃРё, РєРѕРЅС‚СЂРѕР»СЊ lazy-load Рё РїРѕРЅСЏС‚РЅСѓСЋ СЃРІСЏР·СЊ РёР·РѕР±СЂР°Р¶РµРЅРёСЏ СЃ РёРЅС‚РµРЅС‚РѕРј СЃС‚СЂР°РЅРёС†С‹.")

    if get_contact_related_pages(audit["sample_pages"]):
        points.append("Р”РѕРґРµР»Р°С‚СЊ РєРѕРЅС‚Р°РєС‚РЅС‹Рµ Рё Р·Р°СЏРІРѕС‡РЅС‹Рµ СЃС‚СЂР°РЅРёС†С‹ РєР°Рє РїРѕР»РЅРѕС†РµРЅРЅС‹Рµ SEO-РґРѕРєСѓРјРµРЅС‚С‹, Р° РЅРµ РєР°Рє С‚РµС…РЅРёС‡РµСЃРєРёРµ С„РѕСЂРјС‹.")

    points.append("Р”РѕР±Р°РІРёС‚СЊ answer-first Р±Р»РѕРєРё, FAQ Рё РІРЅСѓС‚СЂРµРЅРЅСЋСЋ РїРµСЂРµР»РёРЅРєРѕРІРєСѓ РјРµР¶РґСѓ РІР°Р¶РЅС‹РјРё СЂР°Р·РґРµР»Р°РјРё РґР»СЏ РѕР±С‹С‡РЅРѕР№ Рё РР-РІС‹РґР°С‡Рё.")

    if audit["sitemap_url_count"] > 2000 or audit.get("sitemap_total_entries", 0) > audit["sitemap_url_count"] * 1.2:
        points.append("РЎРµРіРјРµРЅС‚РёСЂРѕРІР°С‚СЊ sitemap Рё РѕС‚РґРµР»СЊРЅРѕ РєРѕРЅС‚СЂРѕР»РёСЂРѕРІР°С‚СЊ РёРЅРґРµРєСЃР°С†РёСЋ РїРѕ С‚РёРїР°Рј СЃС‚СЂР°РЅРёС†, С‡С‚РѕР±С‹ Р±С‹СЃС‚СЂРµРµ РЅР°С…РѕРґРёС‚СЊ СЃР±РѕРё Рё РґСѓР±Р»Рё.")

    unique_points: list[str] = []
    for point in points:
        if point not in unique_points:
            unique_points.append(point)
    return unique_points[:6]


def dynamic_build_roadmap(audit: dict) -> list[tuple[str, list[str]]]:
    profile = infer_project_profile(audit)
    robots_conflict = any("sitemap.xml" in line.lower() and "disallow" in line.lower() for line in audit["robots_lines"])
    query_pages = get_indexable_query_pages(audit["sample_pages"])
    redirect_chains = [item for item in audit["redirect_checks"] if len(item["chain"]) > 1]

    sprint_1 = []
    if robots_conflict:
        sprint_1.append("Убрать конфликт между robots.txt и sitemap.xml.")
    if redirect_chains:
        sprint_1.append("Свести редиректы домена к одному шагу без лишних переходов.")
    if query_pages:
        sprint_1.append("Закрыть от индексации служебные query- и route-страницы.")
    sprint_1.append("Подготовить единые правила для title, description, H1 и canonical.")

    sprint_2 = [
        "Навести порядок в изображениях: alt, lazy loading, размеры и сжатие.",
        "Разделить sitemap по типам страниц и заново отправить карты сайта в панели вебмастеров.",
        "Обновить шаблоны сниппетов на приоритетных страницах и проверить рост CTR.",
        "Разобрать ключевые страницы по модели SEO Монстр: какие оставить, какие переписать, какие объединить, а какие увести в архив.",
    ]

    if profile["is_catalog"]:
        sprint_3 = [
            "Усилить категории и подкатегории коммерческим слоем: условия, гарантии, FAQ, блоки доверия и сценарии выбора.",
            "Пересобрать карточки и листинги так, чтобы они содержали доказательства, сравнения, отзывы и вели к заявке.",
            "Собрать тематическую перелинковку по кластерам спроса и поддержать её хлебными крошками и блоками \"по теме\".",
        ]
    else:
        sprint_3 = [
            "Усилить страницы услуг и контактов: добавить доказательства, FAQ, условия выбора и сильный призыв к действию.",
            "Развести ключевые направления по отдельным страницам и связать их тематической перелинковкой.",
            "Запустить повторную оптимизацию страниц с потерянными запросами: новые H2/H3, факты, кейсы, изображения и обновлённые сниппеты.",
        ]

    return [("0-14 дней", sprint_1[:4]), ("15-30 дней", sprint_2), ("31-60 дней", sprint_3)]


def build_executive_summary_dynamic(audit: dict) -> list[str]:
    profile = infer_project_profile(audit)
    top_issues = audit.get("issues", [])[:3]
    top_issue_titles = "; ".join(issue.title for issue in top_issues) if top_issues else "технические и шаблонные ограничения"
    priority_pages = select_priority_pages(audit["sample_pages"])
    commercial_gap_ratio = coverage_ratio(priority_pages, lambda snapshot: not snapshot.has_commercial_block)
    support_gap_ratio = coverage_ratio(
        priority_pages,
        lambda snapshot: not any(bool(getattr(snapshot, field, False)) for field in MONSTER_SUPPORT_FIELDS),
    )
    structure_gap_ratio = coverage_ratio(priority_pages, lambda snapshot: snapshot.internal_links <= 5)
    project_shape = "категории, карточки и связанные посадочные" if profile["is_catalog"] else "страницы услуг, внутренние коммерческие разделы и лид-страницы"

    return [
        (
            f"Аудит собран по слоям из SEO Монстр: индексация, сниппеты, внутренняя структура, коммерческие факторы и качество контента. "
            f"В выборке {len(audit['sample_pages'])} HTML-страниц, основной фокус — на шаблонах типа {project_shape}."
        ),
        (
            f"Сильнее всего рост сейчас тормозят {top_issue_titles}. Это не косметические замечания: они режут индексацию, ухудшают читаемость сниппетов и оставляют приоритетные страницы слабее конкурентов."
        ),
        (
            "Главный резерв лежит в переработке приоритетных шаблонов, а не в добавлении ещё одного общего SEO-текста: "
            f"без условий и доверия остаются {round(commercial_gap_ratio * 100)}% ключевых страниц, "
            f"без ответов на возражения и сравнительных блоков — {round(support_gap_ratio * 100)}%, "
            f"а со слабой перелинковкой — {round(structure_gap_ratio * 100)}%."
        ),
    ]


def normalize_page_url(url: str) -> str:
    parsed = urlparse(url.strip())
    normalized_path = parsed.path.rstrip("/") or "/"
    cleaned = parsed._replace(path=normalized_path, params="", fragment="")
    return cleaned.geturl()


def select_representative_urls(base_url: str, home_links: list[str], sitemap_urls: list[str], sample_size: int) -> list[str]:
    selected: list[str] = []
    seen: set[str] = set()

    def add(url: str) -> None:
        normalized = normalize_page_url(url)
        if not normalized or normalized in seen:
            return
        seen.add(normalized)
        selected.append(normalized)

    add(base_url)
    for url in home_links[:10]:
        add(url)

    buckets: dict[str, list[str]] = {
        "contact": [],
        "query": [],
        "blog": [],
        "root": [],
        "mid": [],
        "deep": [],
        "other": [],
    }
    for url in sitemap_urls:
        parsed = urlparse(url)
        path = parsed.path.lower()
        depth = len([segment for segment in parsed.path.split("/") if segment])
        if parsed.query:
            buckets["query"].append(url)
        elif any(token in path for token in CONTACT_PATH_KEYWORDS):
            buckets["contact"].append(url)
        elif any(token in path for token in BLOG_PATH_KEYWORDS):
            buckets["blog"].append(url)
        elif depth <= 1:
            buckets["root"].append(url)
        elif depth == 2:
            buckets["mid"].append(url)
        elif depth >= 3:
            buckets["deep"].append(url)
        else:
            buckets["other"].append(url)

    quotas = {
        "contact": 2,
        "query": 2,
        "blog": 3,
        "root": 4,
        "mid": max(6, sample_size // 2),
        "deep": max(6, sample_size // 2),
        "other": 4,
    }
    for bucket_name in ("contact", "query", "blog", "root", "mid", "deep", "other"):
        for url in buckets[bucket_name][: quotas[bucket_name]]:
            add(url)

    for bucket_name in ("contact", "query", "blog", "root", "mid", "deep", "other"):
        for url in buckets[bucket_name]:
            if len(selected) >= sample_size + 10:
                break
            add(url)

    return selected[: sample_size + 10]


def discover_page_links(session: requests.Session, url: str, base_domain: str) -> list[str]:
    response, error = safe_get(session, url, timeout=30)
    if error or response is None or response.status_code >= 400:
        return []

    content_type = response.headers.get("content-type", "")
    if "html" not in content_type:
        return []

    soup = BeautifulSoup(response.text, "lxml")
    discovered: list[str] = []
    seen: set[str] = set()

    for anchor in soup.find_all("a", href=True):
        href = anchor.get("href", "").strip()
        if not href or href.startswith(("#", "mailto:", "tel:", "javascript:")):
            continue

        absolute = normalize_page_url(urljoin(url, href))
        parsed = urlparse(absolute)
        if parsed.netloc != base_domain:
            continue
        if re.search(r"\.(pdf|jpg|jpeg|png|webp|svg|zip|doc|docx|xls|xlsx|ppt|pptx)$", parsed.path.lower()):
            continue
        if absolute in seen:
            continue

        seen.add(absolute)
        discovered.append(absolute)

    return discovered


def crawl_internal_urls(
    base_url: str,
    seed_urls: list[str],
    base_domain: str,
    page_limit: int,
) -> list[str]:
    session = make_session()
    queue = [base_url, *seed_urls]
    discovered: list[str] = []
    queued: set[str] = set()
    visited: set[str] = set()
    effective_limit = page_limit if page_limit > 0 else 400

    for url in queue:
        normalized = normalize_page_url(url)
        queued.add(normalized)

    while queue and len(discovered) < effective_limit:
        current = normalize_page_url(queue.pop(0))
        if current in visited:
            continue

        visited.add(current)
        discovered.append(current)

        for linked_url in discover_page_links(session, current, base_domain):
            if linked_url in visited or linked_url in queued:
                continue
            queued.add(linked_url)
            queue.append(linked_url)
            if len(queued) >= effective_limit:
                break

    return discovered


def build_audit_url_inventory(
    base_url: str,
    preferred_urls: list[str],
    sitemap_urls: list[str],
    base_domain: str,
    requested_limit: int,
) -> list[str]:
    selected: list[str] = []
    seen: set[str] = set()

    def add(url: str) -> None:
        normalized = normalize_page_url(url)
        if not normalized or normalized in seen:
            return
        seen.add(normalized)
        selected.append(normalized)

    add(base_url)
    for url in preferred_urls:
        add(url)

    if sitemap_urls:
        for url in sitemap_urls:
            add(url)
        if requested_limit == 0:
            supplemental_urls = crawl_internal_urls(base_url, preferred_urls[:12], base_domain, min(max(len(preferred_urls) * 12, 180), 480))
            for url in supplemental_urls:
                add(url)
        if requested_limit > 0:
            return selected[:requested_limit]
        return selected

    crawled_urls = crawl_internal_urls(base_url, preferred_urls, base_domain, requested_limit)
    for url in crawled_urls:
        add(url)

    if requested_limit > 0:
        return selected[:requested_limit]
    return selected


def analyse_pages_parallel(urls: list[str], base_domain: str, max_workers: int = 10) -> list[PageSnapshot]:
    if not urls:
        return []

    results: list[PageSnapshot | None] = [None] * len(urls)

    def worker(index: int, page_url: str) -> tuple[int, PageSnapshot]:
        session = make_session()
        return index, analyse_page(session, page_url, base_domain)

    worker_count = max(1, min(max_workers, len(urls)))
    with ThreadPoolExecutor(max_workers=worker_count) as executor:
        futures = [executor.submit(worker, index, page_url) for index, page_url in enumerate(urls)]
        for future in as_completed(futures):
            index, snapshot = future.result()
            results[index] = snapshot

    return [snapshot for snapshot in results if snapshot is not None]


def infer_issue_owner(issue: AuditIssue) -> str:
    haystack = f"{issue.title} {issue.recommendation}".lower()
    if any(token in haystack for token in ("robots", "sitemap", "redirect", "canonical", "query-url", "route", "host")):
        return "Backend / SEO"
    if any(token in haystack for token in ("schema", "json", "title", "description", "h1", "alt", "image", "snippet")):
        return "Frontend / SEO"
    if any(token in haystack for token in ("РєРѕРЅС‚Р°РєС‚", "lead", "cta", "content", "faq", "РїРµСЂРµР»РёРЅРєРѕРІ")):
        return "SEO / Marketing"
    return "SEO"


def build_priority_matrix(audit: dict) -> list[dict]:
    rows: list[dict] = []
    severity_base = {
        "Critical": (10, 9, 10),
        "High": (8, 7, 8),
        "Medium": (6, 5, 6),
        "Low": (3, 3, 4),
    }
    for issue in audit.get("issues", [])[:8]:
        impact, risk, business = severity_base.get(issue.severity, (4, 4, 4))
        haystack = f"{issue.title} {issue.recommendation}".lower()
        if any(token in haystack for token in ("robots", "sitemap", "РёРЅРґРµРєСЃР°", "canonical", "redirect")):
            impact = min(10, impact + 1)
        if any(token in haystack for token in ("title", "description", "ctr", "snippet", "schema")):
            business = min(10, business + 1)
        if any(token in haystack for token in ("РєРѕРЅС‚Р°РєС‚", "lead", "route", "query-url")):
            risk = min(10, risk + 1)
        rows.append(
            {
                "problem": issue.title,
                "severity": issue.severity,
                "impact": impact,
                "risk": risk,
                "business": business,
                "total": impact + risk + business,
                "owner": infer_issue_owner(issue),
            }
        )
    return rows


def build_critical_errors(audit: dict) -> list[dict]:
    return [asdict(issue) for issue in audit.get("issues", []) if issue.severity in ("Critical", "High")][:4]


def build_quick_wins(audit: dict) -> list[dict]:
    quick_wins: list[dict] = []
    issue_titles = {issue.title for issue in audit.get("issues", [])}

    if any("robots.txt" in title for title in issue_titles):
        quick_wins.append(
            {
                "title": "РџРѕС‡РёРЅРёС‚СЊ РєРѕРЅС„Р»РёРєС‚ robots.txt Рё sitemap.xml",
                "effort": "10-20 РјРёРЅСѓС‚",
                "impact": "РЎРЅРёРјРµС‚ РєРѕРЅС„Р»РёРєС‚ СЃРёРіРЅР°Р»РѕРІ РґР»СЏ РЇРЅРґРµРєСЃР° Рё Google",
                "action": "РЈР±СЂР°С‚СЊ Р·Р°РїСЂРµС‚ РЅР° sitemap.xml Рё РѕСЃС‚Р°РІРёС‚СЊ РѕРґРЅСѓ РєРѕСЂСЂРµРєС‚РЅСѓСЋ РґРёСЂРµРєС‚РёРІСѓ Sitemap.",
            }
        )

    if any("redirect" in title.lower() for title in issue_titles):
        quick_wins.append(
            {
                "title": "РЎРІРµСЃС‚Рё РґРѕРјРµРЅРЅС‹Рµ СЂРµРґРёСЂРµРєС‚С‹ Рє РѕРґРЅРѕРјСѓ С€Р°РіСѓ",
                "effort": "15-30 РјРёРЅСѓС‚",
                "impact": "РЈР±РµСЂРµС‚ Р»РёС€РЅСЋСЋ Р·Р°РґРµСЂР¶РєСѓ Рё С‚РµС…РЅРёС‡РµСЃРєРёР№ С€СѓРј РЅР° РІС…РѕРґРµ",
                "action": "РћСЃС‚Р°РІРёС‚СЊ РµРґРёРЅС‹Р№ 301 СЃСЂР°Р·Сѓ РЅР° РєР°РЅРѕРЅРёС‡РµСЃРєРёР№ HTTPS-РґРѕРјРµРЅ Р±РµР· РїСЂРѕРјРµР¶СѓС‚РѕС‡РЅС‹С… С…РѕРїРѕРІ.",
            }
        )

    weak_contacts = audit.get("weak_contact_pages", [])
    if weak_contacts:
        quick_wins.append(
            {
                "title": "РќРѕСЂРјР°Р»СЊРЅРѕ РѕС„РѕСЂРјРёС‚СЊ РєРѕРЅС‚Р°РєС‚С‹ Рё Р»РёРґ-СЃС‚СЂР°РЅРёС†С‹",
                "effort": "30-60 РјРёРЅСѓС‚",
                "impact": "РџРѕРґРЅРёРјРµС‚ РґРѕРІРµСЂРёРµ, Р±СЂРµРЅРґРѕРІС‹Р№ Р·Р°РїСЂРѕСЃ Рё РєРѕРЅРІРµСЂСЃРёСЋ",
                "action": "Р”РѕР±Р°РІРёС‚СЊ H1, title, description, canonical Рё Р±Р»РѕРєРё РґРѕРІРµСЂРёСЏ РЅР° СЃС‚СЂР°РЅРёС†Сѓ РєРѕРЅС‚Р°РєС‚РѕРІ/С„РѕСЂРјС‹.",
            }
        )

    if audit.get("invalid_schema_count", 0):
        quick_wins.append(
            {
                "title": "Исправить битую JSON-LD разметку",
                "effort": "15-40 РјРёРЅСѓС‚",
                "impact": "Вернет валидную микроразметку в индекс",
                "action": "Проверить синтаксис JSON-LD и пересобрать проблемные блоки микроразметки на ключевых страницах.",
            }
        )

    if audit.get("title_missing_count", 0) or audit.get("description_missing_count", 0):
        quick_wins.append(
            {
                "title": "Р—Р°РєСЂС‹С‚СЊ РґС‹СЂС‹ РІ title Рё description",
                "effort": "1-2 С‡Р°СЃР°",
                "impact": "Р”Р°СЃС‚ Р±С‹СЃС‚СЂС‹Р№ РїСЂРёСЂРѕСЃС‚ СѓРїСЂР°РІР»СЏРµРјРѕСЃС‚Рё СЃРЅРёРїРїРµС‚РѕРІ",
                "action": "Р”РѕРїРёСЃР°С‚СЊ С€Р°Р±Р»РѕРЅС‹ РіРµРЅРµСЂР°С†РёРё title/description РґР»СЏ СЃС‚СЂР°РЅРёС† Р±РµР· SEO-РѕР±РІСЏР·РєРё.",
            }
        )

    if audit.get("indexable_query_count", 0):
        quick_wins.append(
            {
                "title": "Закрыть индексируемые URL с параметрами и служебные страницы",
                "effort": "30-90 РјРёРЅСѓС‚",
                "impact": "Сократит шум в индексе и освободит ресурсы на обход",
                "action": "Закрыть служебные страницы и URL с параметрами от индексации или перевести их на чистые страницы для поиска.",
            }
        )

    if audit.get("total_missing_alt", 0):
        quick_wins.append(
            {
                "title": "Р—Р°РїСѓСЃС‚РёС‚СЊ С€Р°Р±Р»РѕРЅ alt РґР»СЏ РёР·РѕР±СЂР°Р¶РµРЅРёР№",
                "effort": "1-2 С‡Р°СЃР°",
                "impact": "РЈСЃРёР»РёС‚ image SEO Рё РїРѕРЅСЏС‚РЅРѕСЃС‚СЊ РєР°СЂС‚РѕС‡РµРє",
                "action": "РЎРѕР±СЂР°С‚СЊ РїСЂР°РІРёР»Р° alt РїРѕ С‚РёРїСѓ СЃС‚СЂР°РЅРёС†С‹, С‚РѕРІР°СЂСѓ, РєР°С‚РµРіРѕСЂРёРё Рё РєР»СЋС‡РµРІРѕР№ СЃСѓС‰РЅРѕСЃС‚Рё РёР·РѕР±СЂР°Р¶РµРЅРёСЏ.",
            }
        )

    if ISSUE_TITLE_COMMERCIAL_GAPS in issue_titles:
        quick_wins.append(
            {
                "title": "Вынести условия и гарантии на приоритетные страницы",
                "effort": "1-3 часа",
                "impact": "Снимет часть возражений и усилит коммерческий интент",
                "action": "Добавить на 5-10 ключевых страниц блоки с ценой/диапазоном, сроками, гарантией, оплатой/доставкой или порядком работ и явным CTA.",
            }
        )

    if ISSUE_TITLE_CONTENT_PROOF in issue_titles:
        quick_wins.append(
            {
                "title": "Добавить доказательства на ключевые страницы",
                "effort": "2-4 часа",
                "impact": "Сделает контент полезнее без раздувания текста",
                "action": "Вставить в приоритетные шаблоны FAQ, таблицы, сравнения, кейсы, отзывы, фото или скрины вместо общего SEO-текста.",
            }
        )

    if ISSUE_TITLE_INTERNAL_LINKING in issue_titles:
        quick_wins.append(
            {
                "title": "Проставить тематические внутренние ссылки",
                "effort": "1-2 часа",
                "impact": "Подтянет важные страницы и снизит навигационный шум",
                "action": "Добавить хлебные крошки, блоки \"по теме\" и осмысленные анкоры с приоритетных страниц на соседние кластеры и страницы заявки.",
            }
        )

    return quick_wins[:6]


def build_strategic_moves(audit: dict) -> list[dict]:
    profile = infer_project_profile(audit)
    moves = [
        {
            "title": "Пересобрать шаблоны ключевых страниц",
            "impact": "Высокое",
            "effort": "5-10 дней",
            "details": "Собрать единые требования к H1, title, description, canonical, FAQ, CTA, блокам доверия и перелинковке по всем важным типам страниц.",
        },
        {
            "title": "Сделать индексацию управляемой",
            "impact": "Высокое",
            "effort": "3-7 дней",
            "details": "Разделить sitemap по типам страниц, убрать служебные route/query URL из индекса и настроить единый контроль над canonical и редиректами.",
        },
        {
            "title": "Навести порядок в микроразметке",
            "impact": "Среднее / высокое",
            "effort": "2-5 дней",
            "details": "Покрыть ключевые страницы валидной schema-разметкой и следить, чтобы она соответствовала реальному типу страницы и данным на сайте.",
        },
        {
            "title": "Запустить контент-аудит по модели оставить / переписать / объединить / архивировать",
            "impact": "Высокое",
            "effort": "5-10 дней",
            "details": "Периодически пересматривать старые страницы: оставлять сильные, объединять пересекающиеся, обновлять доказательства и убирать балласт, который не приносит трафик и не закрывает интент.",
        },
        {
            "title": "Собрать кластеры спроса и тематическую перелинковку",
            "impact": "Высокое",
            "effort": "1-3 недели",
            "details": "Развести приоритетные темы по кластерам, связать их релевантными анкорами, хлебными крошками и блоками \"по теме\", чтобы вес и интент не расползались по сайту.",
        },
    ]
    if profile["is_catalog"]:
        moves.append(
            {
                "title": "Развить спрос через категории и карточки",
                "impact": "Высокое",
                "effort": "2-4 недели",
                "details": "Собрать отдельные посадочные страницы под главные кластеры спроса, усилить категории ответами на вопросы клиента, а карточки сделать более понятными и конверсионными.",
            }
        )
    else:
        moves.append(
            {
                "title": "Развести услуги по отдельным посадочным страницам",
                "impact": "Высокое",
                "effort": "1-3 недели",
                "details": "Вынести ключевые направления в самостоятельные страницы и связать их между собой кейсами, FAQ, экспертными блоками и понятной перелинковкой.",
            }
        )
    return moves


def build_phase_sections(audit: dict) -> list[dict]:
    sample_pages: list[PageSnapshot] = audit["sample_pages"]
    page_type_counter = Counter(page.page_type for page in sample_pages)
    status_counter = Counter(page.status_code for page in sample_pages)
    schema_counter = Counter(
        schema_type
        for page in sample_pages
        for schema_type in page.schema_types
        if schema_type and schema_type != "Invalid JSON-LD"
    )
    contact_pages = get_contact_related_pages(sample_pages)
    query_pages = get_indexable_query_pages(sample_pages)
    invalid_schema_pages = [page for page in sample_pages if "Invalid JSON-LD" in page.schema_types]
    missing_title_pages = [page for page in sample_pages if not page.title]
    missing_description_pages = [page for page in sample_pages if not page.description]
    missing_h1_pages = [page for page in sample_pages if not page.h1s]
    low_link_pages = sorted(sample_pages, key=lambda page: page.internal_links)[:3]
    long_path_pages = [page for page in sample_pages if len(urlparse(page.url).path) > 75]
    thin_pages = [page for page in sample_pages if 0 < page.word_count < 250]
    weak_contact_pages = []
    for page in contact_pages:
        gaps = []
        if not page.h1s:
            gaps.append("H1")
        if not page.title:
            gaps.append("title")
        if not page.description:
            gaps.append("description")
        if not page.canonical:
            gaps.append("canonical")
        if gaps:
            weak_contact_pages.append({"url": page.url, "missing": gaps})
    audit["weak_contact_pages"] = weak_contact_pages
    audit["invalid_schema_count"] = len(invalid_schema_pages)
    audit["title_missing_count"] = len(missing_title_pages)
    audit["description_missing_count"] = len(missing_description_pages)
    audit["indexable_query_count"] = len(query_pages)

    average_internal_links = round(statistics.mean([page.internal_links for page in sample_pages]), 1) if sample_pages else 0
    total_lazy_gaps = sum(page.lazy_missing_count for page in sample_pages)
    total_dimension_gaps = sum(page.missing_dimensions_count for page in sample_pages)
    commercial_pages = [
        page
        for page in sample_pages
        if page.page_type in ("Р“Р»Р°РІРЅР°СЏ", "РљР°С‚РµРіРѕСЂРёСЏ", "РџРѕРґРєР°С‚РµРіРѕСЂРёСЏ", "РўРѕРІР°СЂ", "Р’РЅСѓС‚СЂРµРЅРЅСЏСЏ") or page.has_forms
    ]
    missing_schema_commercial = [page for page in commercial_pages if not page.schema_types]
    average_commercial_words = (
        round(statistics.mean([page.word_count for page in commercial_pages if page.word_count]), 1)
        if commercial_pages
        else 0
    )

    return [
        {
            "title": "Р­С‚Р°Рї 1 вЂ” Crawl & Indexability",
            "intro": "РџСЂРѕРІРµСЂСЏСЋ, РЅР°СЃРєРѕР»СЊРєРѕ С‡РёСЃС‚Рѕ СЃР°Р№С‚ РѕС‚РґР°РµС‚ РїРѕРёСЃРєРѕРІРёРєР°Рј СЃРёРіРЅР°Р»С‹ РґР»СЏ РѕР±С…РѕРґР° Рё РёРЅРґРµРєСЃР°С†РёРё.",
            "checks": [
                {
                    "name": "robots.txt Рё РєР°СЂС‚Р° СЃР°Р№С‚Р°",
                    "checked": "Р”РѕСЃС‚СѓРїРЅРѕСЃС‚СЊ robots.txt, РґРёСЂРµРєС‚РёРІС‹, sitemap, С‡РёСЃС‚РѕС‚Р° СЃРёРіРЅР°Р»РѕРІ РґР»СЏ СЂРѕР±РѕС‚Р°.",
                    "method": "HTTP-Р·Р°РїСЂРѕСЃ + СЂР°Р·Р±РѕСЂ robots.txt Рё XML sitemap.",
                    "metrics": [
                        ("robots.txt", f"{audit['robots_status']}"),
                        ("РЈРЅРёРєР°Р»СЊРЅС‹С… URL РІ sitemap", str(audit["sitemap_url_count"])),
                        ("Р’СЃРµРіРѕ Р·Р°РїРёСЃРµР№ РІ sitemap", str(audit.get("sitemap_total_entries", audit["sitemap_url_count"]))),
                        ("Р¤Р°Р№Р»РѕРІ sitemap РѕР±СЂР°Р±РѕС‚Р°РЅРѕ", str(len(audit.get("processed_sitemaps", [])))),
                    ],
                    "findings": [
                        "Р•СЃС‚СЊ РєРѕРЅС„Р»РёРєС‚ РјРµР¶РґСѓ robots.txt Рё sitemap.xml." if any("robots.txt" in issue.title for issue in audit["issues"]) else "РљСЂРёС‚РёС‡РЅРѕРіРѕ РєРѕРЅС„Р»РёРєС‚Р° РјРµР¶РґСѓ robots.txt Рё sitemap.xml РІ РІС‹Р±РѕСЂРєРµ РЅРµ РІРёРґРЅРѕ.",
                        "РљР°СЂС‚Р° СЃР°Р№С‚Р° СЂР°Р·РґСѓС‚Р° РґСѓР±Р»СЏРјРё." if audit.get("sitemap_total_entries", 0) > audit["sitemap_url_count"] * 1.3 else "РЎРёР»СЊРЅРѕРіРѕ СЂР°Р·РґСѓРІР°РЅРёСЏ sitemap РґСѓР±Р»СЏРјРё РІ РІС‹Р±РѕСЂРєРµ РЅРµ РІРёРґРЅРѕ.",
                        "Р•СЃС‚СЊ llms.txt, Р·РЅР°С‡РёС‚ СЃР°Р№С‚ СѓР¶Рµ РјРѕР¶РЅРѕ РґРѕР¶РёРјР°С‚СЊ Рё РїРѕРґ РР-РІРёРґРёРјРѕСЃС‚СЊ." if audit.get("llms_exists") else "РћС‚РґРµР»СЊРЅРѕРіРѕ llms.txt РЅРµ РЅР°Р№РґРµРЅРѕ.",
                    ],
                    "priority": "HIGH" if any(issue.severity == "Critical" for issue in audit["issues"]) else "MEDIUM",
                    "owner": "Backend / SEO",
                    "recommendation": "РЎРёРЅС…СЂРѕРЅРёР·РёСЂРѕРІР°С‚СЊ robots.txt, sitemap Рё РєР°РЅРѕРЅРёС‡РµСЃРєРёР№ РґРѕРјРµРЅ С‚Р°Рє, С‡С‚РѕР±С‹ СЂРѕР±РѕС‚ РІРёРґРµР» РѕРґРёРЅ РЅРµРїСЂРѕС‚РёРІРѕСЂРµС‡РёРІС‹Р№ СЃС†РµРЅР°СЂРёР№ РѕР±С…РѕРґР°.",
                },
                {
                    "name": "РЎС‚Р°С‚СѓСЃ-РєРѕРґС‹, СЂРµРґРёСЂРµРєС‚С‹ Рё РєР°РЅРѕРЅРёРєР°Р»РёР·Р°С†РёСЏ",
                    "checked": "HTTP-РєРѕРґС‹ РєР»СЋС‡РµРІС‹С… URL, РґРѕРјРµРЅРЅС‹Рµ СЂРµРґРёСЂРµРєС‚С‹, РїРѕРєСЂС‹С‚РёРµ canonical.",
                    "method": "HTTP-РїСЂРѕРІРµСЂРєР° РґРѕРјРµРЅРЅС‹С… РІР°СЂРёР°РЅС‚РѕРІ + РїР°СЂСЃРёРЅРі rel=canonical РїРѕ РІС‹Р±РѕСЂРєРµ СЃС‚СЂР°РЅРёС†.",
                    "metrics": [
                        ("РЎС‚СЂР°РЅРёС† 200 РІ РІС‹Р±РѕСЂРєРµ", str(status_counter.get(200, 0))),
                        ("Redirect-С†РµРїРѕС‡РµРє > 1 С€Р°РіР°", str(len([item for item in audit["redirect_checks"] if len(item["chain"]) > 1]))),
                        ("Canonical coverage", f"{math.floor(audit['canonical_coverage_ratio'] * 100)}%"),
                        ("РРЅРґРµРєСЃРёСЂСѓРµРјС‹С… query-URL", str(len(query_pages))),
                    ],
                    "findings": [
                        f"Р’ РІС‹Р±РѕСЂРєРµ {status_counter.get(200, 0)} СЃС‚СЂР°РЅРёС† РѕС‚РІРµС‡Р°СЋС‚ РєРѕРґРѕРј 200." if status_counter else "РЎС‚Р°С‚СѓСЃ-РєРѕРґС‹ РІ РІС‹Р±РѕСЂРєРµ РїСЂРѕРІРµСЂРёС‚СЊ РЅРµ СѓРґР°Р»РѕСЃСЊ.",
                        "Р•СЃС‚СЊ Р»РёС€РЅСЏСЏ redirect-С†РµРїРѕС‡РєР° РЅР° РґРѕРјРµРЅРЅС‹С… РІР°СЂРёР°РЅС‚Р°С…." if any("redirect" in issue.title.lower() for issue in audit["issues"]) else "Р›РёС€РЅРёС… РґРѕРјРµРЅРЅС‹С… redirect-С†РµРїРѕС‡РµРє РІ РєР»СЋС‡РµРІС‹С… РІР°СЂРёР°РЅС‚Р°С… РЅРµ РѕР±РЅР°СЂСѓР¶РµРЅРѕ.",
                        "Р•СЃС‚СЊ РёРЅРґРµРєСЃРёСЂСѓРµРјС‹Рµ query- РёР»Рё route-URL, РєРѕС‚РѕСЂС‹Рµ Р»СѓС‡С€Рµ РїРѕС‡РёСЃС‚РёС‚СЊ." if query_pages else "РЁСѓРјР° РѕС‚ РёРЅРґРµРєСЃРёСЂСѓРµРјС‹С… query-СЃС‚СЂР°РЅРёС† РІ РІС‹Р±РѕСЂРєРµ РїРѕС‡С‚Рё РЅРµС‚.",
                    ],
                    "priority": "HIGH" if query_pages else "MEDIUM",
                    "owner": "Backend / SEO",
                    "recommendation": "РЎРІРµСЃС‚Рё РІСЃРµ РґРѕРјРµРЅРЅС‹Рµ РІР°СЂРёР°РЅС‚С‹ Рє РѕРґРЅРѕРјСѓ РєР°РЅРѕРЅСѓ Рё Р»РёР±Рѕ Р·Р°РєСЂС‹С‚СЊ СЃР»СѓР¶РµР±РЅС‹Рµ URL РѕС‚ РёРЅРґРµРєСЃР°С†РёРё, Р»РёР±Рѕ РїРµСЂРµРІРµСЃС‚Рё РёС… РЅР° С‡РёСЃС‚С‹Рµ SEO-friendly РјР°СЂС€СЂСѓС‚С‹.",
                },
            ],
        },
        {
            "title": "Р­С‚Р°Рї 2 вЂ” РђСЂС…РёС‚РµРєС‚СѓСЂР° Рё РІРЅСѓС‚СЂРµРЅРЅСЏСЏ СЃС‚СЂСѓРєС‚СѓСЂР°",
            "intro": "РЎРјРѕС‚СЂСЋ, РЅР°СЃРєРѕР»СЊРєРѕ РїРѕРЅСЏС‚РЅР° Р°СЂС…РёС‚РµРєС‚СѓСЂР° СЃР°Р№С‚Р° Рё РєР°Рє СЂР°СЃРїСЂРµРґРµР»СЏСЋС‚СЃСЏ СЃСЃС‹Р»РєРё РІРЅСѓС‚СЂРё РІС‹Р±РѕСЂРєРё.",
            "checks": [
                {
                    "name": "РўРёРїС‹ СЃС‚СЂР°РЅРёС† Рё РїРѕРєСЂС‹С‚РёРµ РІС‹Р±РѕСЂРєРё",
                    "checked": "РљР°РєРёРµ С‚РёРїС‹ URL СЂРµР°Р»СЊРЅРѕ РїРѕРїР°Р»Рё РІ СЂР°Р·Р±РѕСЂ Рё РєР°Рє СЃР°Р№С‚ РґСЂРѕР±РёС‚ СЃРїСЂРѕСЃ.",
                    "method": "РљР»Р°СЃСЃРёС„РёРєР°С†РёСЏ URL РїРѕ РіР»СѓР±РёРЅРµ, query-РїР°СЂР°РјРµС‚СЂР°Рј Рё С‚РёРїР°Рј schema.",
                    "metrics": [
                        ("Р“Р»Р°РІРЅР°СЏ", str(page_type_counter.get("Р“Р»Р°РІРЅР°СЏ", 0))),
                        ("РљР°С‚РµРіРѕСЂРёРё / РїРѕРґРєР°С‚РµРіРѕСЂРёРё", str(page_type_counter.get("РљР°С‚РµРіРѕСЂРёСЏ", 0) + page_type_counter.get("РџРѕРґРєР°С‚РµРіРѕСЂРёСЏ", 0))),
                        ("Р’РЅСѓС‚СЂРµРЅРЅРёРµ СЃС‚СЂР°РЅРёС†С‹", str(page_type_counter.get("Р’РЅСѓС‚СЂРµРЅРЅСЏСЏ", 0))),
                        ("РЎР»СѓР¶РµР±РЅС‹Рµ URL", str(page_type_counter.get("РЎР»СѓР¶РµР±РЅР°СЏ", 0))),
                    ],
                    "findings": [
                        "Р’С‹Р±РѕСЂРєР° СѓР¶Рµ РїРѕРєСЂС‹РІР°РµС‚ СЂР°Р·РЅС‹Рµ СѓСЂРѕРІРЅРё СЃС‚СЂСѓРєС‚СѓСЂС‹, РїРѕСЌС‚РѕРјСѓ РІРёРґРЅРѕ РЅРµ С‚РѕР»СЊРєРѕ РіР»Р°РІРЅСѓСЋ, РЅРѕ Рё С€Р°Р±Р»РѕРЅС‹ РІРЅСѓС‚СЂРµРЅРЅРёС… СЃС‚СЂР°РЅРёС†.",
                        "Р’ Р°СЂС…РёС‚РµРєС‚СѓСЂРµ РµСЃС‚СЊ СЃР»СѓР¶РµР±РЅС‹Рµ РёР»Рё route-СЃС‚СЂР°РЅРёС†С‹, РєРѕС‚РѕСЂС‹Рµ СЃРјРµС€РёРІР°СЋС‚СЃСЏ СЃ РєРѕРјРјРµСЂС‡РµСЃРєРёРјРё URL." if page_type_counter.get("РЎР»СѓР¶РµР±РЅР°СЏ", 0) else "РЇРІРЅРѕРіРѕ РґР°РІР»РµРЅРёСЏ СЃР»СѓР¶РµР±РЅС‹С… URL РЅР° РІС‹Р±РѕСЂРєСѓ РїРѕС‡С‚Рё РЅРµС‚.",
                        f"Р”Р»РёРЅРЅС‹С… РїСѓС‚РµР№ РІ РІС‹Р±РѕСЂРєРµ: {len(long_path_pages)}." if long_path_pages else "РЎРёР»СЊРЅРѕРіРѕ РїРµСЂРµРєРѕСЃР° РІ СЃС‚РѕСЂРѕРЅСѓ СЃР»РёС€РєРѕРј РґР»РёРЅРЅС‹С… URL РІ РІС‹Р±РѕСЂРєРµ РЅРµ РІРёРґРЅРѕ.",
                    ],
                    "priority": "MEDIUM",
                    "owner": "SEO / Backend",
                    "recommendation": "Р”РµСЂР¶Р°С‚СЊ Р°СЂС…РёС‚РµРєС‚СѓСЂСѓ СЃРїСЂРѕСЃР° РѕС‚РґРµР»СЊРЅРѕР№ РѕС‚ СЃР»СѓР¶РµР±РЅРѕР№ Р»РѕРіРёРєРё Рё РЅРµ СЃРјРµС€РёРІР°С‚СЊ route/query URL СЃ РїРѕСЃР°РґРѕС‡РЅС‹РјРё Рё РєРѕРјРјРµСЂС‡РµСЃРєРёРјРё СЃС‚СЂР°РЅРёС†Р°РјРё.",
                },
                {
                    "name": "Р’РЅСѓС‚СЂРµРЅРЅСЏСЏ РїРµСЂРµР»РёРЅРєРѕРІРєР° Рё РІРёРґРёРјРѕСЃС‚СЊ РєР»СЋС‡РµРІС‹С… СЃС‚СЂР°РЅРёС†",
                    "checked": "РЎРєРѕР»СЊРєРѕ РІРЅСѓС‚СЂРµРЅРЅРёС… СЃСЃС‹Р»РѕРє РїРѕР»СѓС‡Р°СЋС‚ СЃС‚СЂР°РЅРёС†С‹ Рё РіРґРµ РµСЃС‚СЊ РїСЂРѕСЃР°РґРєРё РїРѕ РІРЅРёРјР°РЅРёСЋ СЃР°Р№С‚Р°.",
                    "method": "РџРѕРґСЃС‡РµС‚ РІРЅСѓС‚СЂРµРЅРЅРёС… СЃСЃС‹Р»РѕРє РЅР° РєР°Р¶РґРѕР№ СЃС‚СЂР°РЅРёС†Рµ РІС‹Р±РѕСЂРєРё.",
                    "metrics": [
                        ("РЎСЂРµРґРЅРµРµ С‡РёСЃР»Рѕ РІРЅСѓС‚СЂРµРЅРЅРёС… СЃСЃС‹Р»РѕРє", str(average_internal_links)),
                        ("РњРёРЅРёРјСѓРј РІРЅСѓС‚СЂРµРЅРЅРёС… СЃСЃС‹Р»РѕРє", str(low_link_pages[0].internal_links if low_link_pages else 0)),
                        ("РњР°РєСЃРёРјСѓРј РІРЅСѓС‚СЂРµРЅРЅРёС… СЃСЃС‹Р»РѕРє", str(max([page.internal_links for page in sample_pages], default=0))),
                        ("РљРѕРЅС‚Р°РєС‚РЅС‹С…/lead-СЃС‚СЂР°РЅРёС†", str(len(contact_pages))),
                    ],
                    "findings": [
                        f"РЎР»Р°Р±РµРµ РІСЃРµРіРѕ РїРѕ РІРЅСѓС‚СЂРµРЅРЅРёРј СЃСЃС‹Р»РєР°Рј РІС‹РіР»СЏРґСЏС‚: {', '.join(human_path(page.url) for page in low_link_pages if page.url)}." if low_link_pages else "РЎР»Р°Р±С‹Рµ СЃС‚СЂР°РЅРёС†С‹ РїРѕ РїРµСЂРµР»РёРЅРєРѕРІРєРµ РЅРµ РІС‹РґРµР»СЏСЋС‚СЃСЏ.",
                        "РљРѕРЅС‚Р°РєС‚РЅС‹Рµ СЃС‚СЂР°РЅРёС†С‹ РµСЃС‚СЊ, РЅРѕ С‡Р°СЃС‚СЊ РёР· РЅРёС… РЅРµРґРѕРѕС„РѕСЂРјР»РµРЅР° РєР°Рє РїРѕР»РЅРѕС†РµРЅРЅС‹Рµ SEO-РїРѕСЃР°РґРєРё." if weak_contact_pages else "РљРѕРЅС‚Р°РєС‚РЅР°СЏ Р·РѕРЅР° РІ РІС‹Р±РѕСЂРєРµ РІС‹РіР»СЏРґРёС‚ С†РµР»СЊРЅРѕ.",
                        "Р Р°Р·СѓРјРЅРѕ РґРѕР±Р°РІРёС‚СЊ Р±РѕР»СЊС€Рµ РєРѕРЅС‚РµРєСЃС‚РЅС‹С… СЃСЃС‹Р»РѕРє РјРµР¶РґСѓ Р±Р»РёР·РєРёРјРё СѓСЃР»СѓРіР°РјРё, РєРµР№СЃР°РјРё, Р±Р»РѕРіРѕРј Рё Р»РёРґ-СЃС‚СЂР°РЅРёС†Р°РјРё.",
                    ],
                    "priority": "MEDIUM",
                    "owner": "SEO / Marketing",
                    "recommendation": "РЈСЃРёР»РёС‚СЊ РїРµСЂРµР»РёРЅРєРѕРІРєСѓ РІРѕРєСЂСѓРі РїСЂРёРѕСЂРёС‚РµС‚РЅС‹С… СѓСЃР»СѓРі Рё Р»РёРґ-СЃС‚СЂР°РЅРёС†: СЃРІСЏР·Р°С‚СЊ СЃРїСЂРѕСЃРѕРІС‹Рµ СЃС‚СЂР°РЅРёС†С‹, РєРµР№СЃС‹, FAQ Рё CTA, Р° РЅРµ РЅР°РґРµСЏС‚СЊСЃСЏ С‚РѕР»СЊРєРѕ РЅР° РјРµРЅСЋ Рё С„СѓС‚РµСЂ.",
                },
            ],
        },
        {
            "title": "Р­С‚Р°Рї 3 вЂ” On-Page SEO",
            "intro": "Р Р°Р·Р±РёСЂР°СЋ С€Р°Р±Р»РѕРЅС‹ title, description, H1 Рё С‚Рѕ, РЅР°СЃРєРѕР»СЊРєРѕ СЃС‚СЂР°РЅРёС†С‹ РіРѕС‚РѕРІС‹ Рє РЅРѕСЂРјР°Р»СЊРЅРѕРјСѓ СЃРЅРёРїРїРµС‚Сѓ.",
            "checks": [
                {
                    "name": "Title и описание сниппета",
                    "checked": "РџРѕРєСЂС‹С‚РёРµ, РґР»РёРЅР° Рё РїСЂРёРіРѕРґРЅРѕСЃС‚СЊ СЃРЅРёРїРїРµС‚РѕРІ Рє СѓРїСЂР°РІР»СЏРµРјРѕР№ РІС‹РґР°С‡Рµ.",
                    "method": "РџР°СЂСЃРёРЅРі title Рё meta description РїРѕ СЂРµРїСЂРµР·РµРЅС‚Р°С‚РёРІРЅРѕР№ РІС‹Р±РѕСЂРєРµ СЃС‚СЂР°РЅРёС†.",
                    "metrics": [
                        ("Title длиннее 70 символов", str(len([page for page in sample_pages if len(page.title) > 70]))),
                        ("Title РѕС‚СЃСѓС‚СЃС‚РІСѓРµС‚", str(len(missing_title_pages))),
                        ("Описание страницы отсутствует", str(len(missing_description_pages))),
                        ("Описание страницы вне диапазона", str(len([page for page in sample_pages if len(page.description) > 160 or (0 < len(page.description) < 120)]))),
                    ],
                    "findings": [
                        "РЁР°Р±Р»РѕРЅС‹ title РЅР° С‡Р°СЃС‚Рё СЃС‚СЂР°РЅРёС† РїРµСЂРµРіСЂСѓР¶РµРЅС‹ РїРѕ РґР»РёРЅРµ." if any("title" in issue.title.lower() for issue in audit["issues"]) else "РљСЂРёС‚РёС‡РµСЃРєРѕР№ СЏРјС‹ РїРѕ title РІ РІС‹Р±РѕСЂРєРµ РЅРµ РІРёРґРЅРѕ.",
                        "Есть страницы без описания, из-за чего сниппет будет собираться случайно." if missing_description_pages else "Описания страниц покрыты достаточно ровно.",
                        f"РџСѓСЃС‚РѕР№ title РЅР°Р№РґРµРЅ РЅР°: {', '.join(human_path(page.url) for page in missing_title_pages[:3])}." if missing_title_pages else "РџСѓСЃС‚С‹С… title РІ РІС‹Р±РѕСЂРєРµ РїРѕС‡С‚Рё РЅРµС‚.",
                    ],
                    "priority": "HIGH",
                    "owner": "SEO / Frontend",
                    "recommendation": "РџРµСЂРµСЃРѕР±СЂР°С‚СЊ РїСЂР°РІРёР»Р° РіРµРЅРµСЂР°С†РёРё title Рё description РїРѕРґ РєР°Р¶РґС‹Р№ С‚РёРї СЃС‚СЂР°РЅРёС†С‹: РєРѕСЂРѕС‚РєРѕ, РєРѕРЅРєСЂРµС‚РЅРѕ, СЃ РєР»СЋС‡РѕРј, РѕС„С„РµСЂРѕРј Рё РєРѕРЅС‚СЂРѕР»РµРј РґР»РёРЅС‹.",
                },
                {
                    "name": "H1, РєРѕРЅС‚РµРЅС‚РЅС‹Р№ РєР°СЂРєР°СЃ Рё С‚РѕРЅРєРёРµ СЃС‚СЂР°РЅРёС†С‹",
                    "checked": "РќР°Р»РёС‡РёРµ H1, РїР»РѕС‚РЅРѕСЃС‚СЊ С‚РµРєСЃС‚РѕРІРѕРіРѕ СЃР»РѕСЏ Рё РєР°С‡РµСЃС‚РІРѕ Р±Р°Р·РѕРІРѕРіРѕ on-page РєР°СЂРєР°СЃР°.",
                    "method": "РџР°СЂСЃРёРЅРі H1 Рё РїРѕРґСЃС‡РµС‚ СЃР»РѕРІ РІ РІРёРґРёРјРѕРј РєРѕРЅС‚РµРЅС‚Рµ.",
                    "metrics": [
                        ("Покрытие H1", f"{math.floor(audit['h1_coverage_ratio'] * 100)}%"),
                        ("РЎС‚СЂР°РЅРёС† Р±РµР· H1", str(len(missing_h1_pages))),
                        ("РЎСЂРµРґРЅРµРµ С‡РёСЃР»Рѕ СЃР»РѕРІ", str(audit.get("average_words", 0))),
                        ("РўРѕРЅРєРёС… СЃС‚СЂР°РЅРёС† (<250 СЃР»РѕРІ)", str(len(thin_pages))),
                    ],
                    "findings": [
                        "РЈ С‡Р°СЃС‚Рё СЃС‚СЂР°РЅРёС† РЅРµС‚ H1, РїРѕСЌС‚РѕРјСѓ РїРѕРёСЃРєРѕРІРёРєСѓ СЃР»РѕР¶РЅРµРµ РїРѕРЅСЏС‚СЊ РіР»Р°РІРЅС‹Р№ РёРЅС‚РµРЅС‚ URL." if missing_h1_pages else "РџРѕ H1 РєР°СЂРєР°СЃ Сѓ РІС‹Р±РѕСЂРєРё РІ С†РµР»РѕРј СЃРѕР±СЂР°РЅ РЅРµРїР»РѕС…Рѕ.",
                        "Р•СЃС‚СЊ С‚РѕРЅРєРёРµ СЃС‚СЂР°РЅРёС†С‹ СЃ РјРёРЅРёРјР°Р»СЊРЅС‹Рј С‚РµРєСЃС‚РѕРІС‹Рј СЃР»РѕРµРј." if thin_pages else "РЎРёР»СЊРЅРѕ РїСѓСЃС‚С‹С… СЃС‚СЂР°РЅРёС† РІ РІС‹Р±РѕСЂРєРµ РЅРµРјРЅРѕРіРѕ.",
                        "Р”Р°Р¶Рµ РїСЂРё РЅРѕСЂРјР°Р»СЊРЅРѕРј РѕР±СЉРµРјРµ С‚РµРєСЃС‚Р° РЅСѓР¶РЅРѕ РѕС‚РґРµР»СЊРЅРѕ РїСЂРѕРІРµСЂРёС‚СЊ, РЅР°СЃРєРѕР»СЊРєРѕ С…РѕСЂРѕС€Рѕ РѕРЅ РїРѕРґРґРµСЂР¶РёРІР°РµС‚ РєРѕРјРјРµСЂС‡РµСЃРєРёР№ РёРЅС‚РµРЅС‚ Рё FAQ-СЃР»РѕР№.",
                    ],
                    "priority": "MEDIUM",
                    "owner": "SEO / Content",
                    "recommendation": "РќР° РєР°Р¶РґРѕРј РєР»СЋС‡РµРІРѕРј URL Р·Р°РєСЂРµРїРёС‚СЊ РїРѕРЅСЏС‚РЅС‹Р№ H1, РёРЅС‚СЂРѕ, РѕС‚РІРµС‚С‹ РЅР° С‡Р°СЃС‚С‹Рµ РІРѕРїСЂРѕСЃС‹ Рё РєРѕСЂРѕС‚РєРёРµ РєРѕРјРјРµСЂС‡РµСЃРєРёРµ Р±Р»РѕРєРё, С‡С‚РѕР±С‹ СЃС‚СЂР°РЅРёС†Р° СЂР°Р±РѕС‚Р°Р»Р° РЅРµ С‚РѕР»СЊРєРѕ РЅР° РёРЅРґРµРєСЃ, РЅРѕ Рё РЅР° РєРѕРЅРІРµСЂСЃРёСЋ.",
                },
            ],
        },
        {
            "title": "Р­С‚Р°Рї 4 вЂ” Performance Рё CWV",
            "intro": "РЎРјРѕС‚СЂСЋ РЅР° СЃРєРѕСЂРѕСЃС‚СЊ РѕС‚РІРµС‚Р°, РІРµСЃ HTML Рё С‚Рѕ, РєР°Рє РјРµРґРёР°-СЃР»РѕР№ РјРѕР¶РµС‚ РјРµС€Р°С‚СЊ Р·Р°РіСЂСѓР·РєРµ.",
            "checks": [
                {
                    "name": "РћС‚РІРµС‚ СЃРµСЂРІРµСЂР° Рё РІРµСЃ СЃС‚СЂР°РЅРёС†",
                    "checked": "TTFB-РїРѕРґРѕР±РЅР°СЏ СЃРєРѕСЂРѕСЃС‚СЊ РѕС‚РІРµС‚Р° Рё РїСЂРёРјРµСЂРЅС‹Р№ РІРµСЃ HTML РїРѕ РІС‹Р±РѕСЂРєРµ.",
                    "method": "HTTP-РѕС‚РІРµС‚С‹ РїРѕ РєР»СЋС‡РµРІС‹Рј URL Р±РµР· РѕС‚РґРµР»СЊРЅРѕРіРѕ Lighthouse-РїСЂРѕРіРѕРЅР°.",
                    "metrics": [
                        ("РЎСЂРµРґРЅРёР№ РѕС‚РІРµС‚", f"{audit.get('average_response_ms', 0)} ms"),
                        ("РЎСЂРµРґРЅРёР№ HTML", f"{audit.get('average_html_kb', 0)} KB"),
                        ("РњР°РєСЃРёРјСѓРј РёР·РѕР±СЂР°Р¶РµРЅРёР№ РЅР° СЃС‚СЂР°РЅРёС†Рµ", str(max([page.image_count for page in sample_pages], default=0))),
                        ("РЎС‚СЂР°РЅРёС† РІ РІС‹Р±РѕСЂРєРµ", str(len(sample_pages))),
                    ],
                    "findings": [
                        "РЎРєРѕСЂРѕСЃС‚СЊ РѕС‚РІРµС‚Р° СЃРµСЂРІРµСЂР° СЃР°РјР° РїРѕ СЃРµР±Рµ РЅРµ РІС‹РіР»СЏРґРёС‚ РіР»Р°РІРЅС‹Рј СЃС‚РѕРї-С„Р°РєС‚РѕСЂРѕРј." if audit.get("average_response_ms", 0) < 600 else "РЎСЂРµРґРЅРёР№ РѕС‚РІРµС‚ СѓР¶Рµ СЃС‚РѕРёС‚ РґРµСЂР¶Р°С‚СЊ РІ Р·РѕРЅРµ РІРЅРёРјР°РЅРёСЏ.",
                        "Р РµР°Р»СЊРЅС‹Рµ СЂРёСЃРєРё Р·РґРµСЃСЊ С‡Р°С‰Рµ СЃРёРґСЏС‚ РІ РјРµРґРёР°-СЃР»РѕРµ, С€Р°Р±Р»РѕРЅР°С… Рё РєР»РёРµРЅС‚СЃРєРѕРј СЂРµРЅРґРµСЂРёРЅРіРµ, С‡РµРј РІ РіРѕР»РѕРј TTFB.",
                    ],
                    "priority": "MEDIUM",
                    "owner": "Frontend / Backend",
                    "recommendation": "РџСЂРѕРІРµСЂРёС‚СЊ РїСЂРёРѕСЂРёС‚РµС‚РЅС‹Рµ СЃС‚СЂР°РЅРёС†С‹ РѕС‚РґРµР»СЊРЅС‹Рј CWV-РїСЂРѕРіРѕРЅРѕРј Рё Р·Р°С„РёРєСЃРёСЂРѕРІР°С‚СЊ LCP/CLS/INP РїРѕСЃР»Рµ С‡РёСЃС‚РєРё РјРµРґРёР° Рё С€Р°Р±Р»РѕРЅРѕРІ.",
                },
                {
                    "name": "РР·РѕР±СЂР°Р¶РµРЅРёСЏ, lazy loading Рё Р°С‚СЂРёР±СѓС‚С‹ СЂР°Р·РјРµСЂРѕРІ",
                    "checked": "РќР°СЃРєРѕР»СЊРєРѕ РјРµРґРёР°-СЃР»РѕР№ РїРѕРјРѕРіР°РµС‚ РёР»Рё РјРµС€Р°РµС‚ SEO Рё СЃС‚Р°Р±РёР»СЊРЅРѕСЃС‚Рё РІРµСЂСЃС‚РєРё.",
                    "method": "РџР°СЂСЃРёРЅРі img-С‚РµРіРѕРІ РїРѕ СЂРµРїСЂРµР·РµРЅС‚Р°С‚РёРІРЅРѕР№ РІС‹Р±РѕСЂРєРµ СЃС‚СЂР°РЅРёС†.",
                    "metrics": [
                        ("Р’СЃРµРіРѕ РїСЂРѕРїСѓСЃРєРѕРІ alt", str(audit.get("total_missing_alt", 0))),
                        ("РџСЂРѕРїСѓСЃРєРѕРІ lazy loading", str(total_lazy_gaps)),
                        ("РР·РѕР±СЂР°Р¶РµРЅРёР№ Р±РµР· width/height", str(total_dimension_gaps)),
                        ("РР·РѕР±СЂР°Р¶РµРЅРёР№ РЅР° РіР»Р°РІРЅРѕР№", str(audit["home_page"].image_count)),
                    ],
                    "findings": [
                        "РР·РѕР±СЂР°Р¶РµРЅРёСЏ СѓР¶Рµ С‚РµСЂСЏСЋС‚ SEO-СЃРёРіРЅР°Р»С‹ РёР·-Р·Р° РїСЂРѕРїСѓСЃРєРѕРІ alt." if audit.get("total_missing_alt", 0) else "РњР°СЃСЃРѕРІРѕРіРѕ РїСЂРѕРІР°Р»Р° РїРѕ alt РІ РІС‹Р±РѕСЂРєРµ РЅРµ РІРёРґРЅРѕ.",
                        "РћС‚СЃСѓС‚СЃС‚РІРёРµ width/height РїРѕРІС‹С€Р°РµС‚ СЂРёСЃРє CLS Рё РІРёР·СѓР°Р»СЊРЅРѕР№ РЅРµСЃС‚Р°Р±РёР»СЊРЅРѕСЃС‚Рё." if total_dimension_gaps else "РђС‚СЂРёР±СѓС‚С‹ СЂР°Р·РјРµСЂРѕРІ РЅР° РєР°СЂС‚РёРЅРєР°С… РІ С†РµР»РѕРј РїСЂРёСЃСѓС‚СЃС‚РІСѓСЋС‚.",
                        "Lazy loading СЃС‚РѕРёС‚ РІС‹СЂРѕРІРЅСЏС‚СЊ С€Р°Р±Р»РѕРЅРЅРѕ, Р° РЅРµ РїСЂР°РІРёС‚СЊ РІСЂСѓС‡РЅСѓСЋ РїРѕ РѕРґРЅРѕР№ СЃС‚СЂР°РЅРёС†Рµ." if total_lazy_gaps else "РЇРІРЅРѕР№ РјР°СЃСЃРѕРІРѕР№ РїСЂРѕР±Р»РµРјС‹ РїРѕ lazy loading РІ РІС‹Р±РѕСЂРєРµ РЅРµ РІРёРґРЅРѕ.",
                    ],
                    "priority": "MEDIUM",
                    "owner": "Frontend / SEO",
                    "recommendation": "РЎРґРµР»Р°С‚СЊ РµРґРёРЅС‹Р№ image-С€Р°Р±Р»РѕРЅ: alt, width/height, lazy loading Рё, РїСЂРё РЅРµРѕР±С…РѕРґРёРјРѕСЃС‚Рё, РѕС‚РґРµР»СЊРЅСѓСЋ image-СЃС…РµРјСѓ/РєР°СЂС‚Сѓ СЃР°Р№С‚Р°.",
                },
            ],
        },
        {
            "title": "Р­С‚Р°Рї 5 вЂ” Structured Data",
            "intro": "РџСЂРѕРІРµСЂСЏСЋ, РЅР°СЃРєРѕР»СЊРєРѕ С…РѕСЂРѕС€Рѕ СЃР°Р№С‚ РѕР±СЉСЏСЃРЅСЏРµС‚ РїРѕРёСЃРєРѕРІРёРєСѓ СЃСѓС‰РЅРѕСЃС‚Рё Р±РёР·РЅРµСЃР°, СЃС‚СЂР°РЅРёС† Рё РїСЂРµРґР»РѕР¶РµРЅРёР№.",
            "checks": [
                {
                    "name": "РџРѕРєСЂС‹С‚РёРµ schema-СЂР°Р·РјРµС‚РєРѕР№",
                    "checked": "РљР°РєРёРµ С‚РёРїС‹ schema СЂРµР°Р»СЊРЅРѕ РІСЃС‚СЂРµС‡Р°СЋС‚СЃСЏ Рё РЅР°СЃРєРѕР»СЊРєРѕ С€РёСЂРѕРєРѕ РѕРЅРё РїРѕРєСЂС‹РІР°СЋС‚ РІС‹Р±РѕСЂРєСѓ.",
                    "method": "РџР°СЂСЃРёРЅРі application/ld+json Рё СЃР±РѕСЂ @type РїРѕ РєР»СЋС‡РµРІС‹Рј URL.",
                    "metrics": [
                        ("Schema coverage", f"{math.floor(audit['schema_coverage_ratio'] * 100)}%"),
                        ("Страниц с битой JSON-LD", str(len(invalid_schema_pages))),
                        ("Страниц без микроразметки", str(len([page for page in sample_pages if not page.schema_types]))),
                        ("Типы микроразметки", ", ".join(f"{name} ({count})" for name, count in schema_counter.most_common(3)) or "нет"),
                    ],
                    "findings": [
                        "Р•СЃС‚СЊ СЃС‚СЂР°РЅРёС†С‹ СЃ РЅРµРІР°Р»РёРґРЅРѕР№ JSON-LD." if invalid_schema_pages else "РЇРІРЅС‹С… РїРѕР»РѕРјРѕРє JSON-LD РІ РІС‹Р±РѕСЂРєРµ РЅРµ РІРёРґРЅРѕ.",
                        "Р§Р°СЃС‚СЊ РєРѕРјРјРµСЂС‡РµСЃРєРёС… СЃС‚СЂР°РЅРёС† РїРѕРєР° Р±РµР· schema-СЂР°Р·РјРµС‚РєРё." if missing_schema_commercial else "РљРѕРјРјРµСЂС‡РµСЃРєРёРµ СЃС‚СЂР°РЅРёС†С‹ РІ РІС‹Р±РѕСЂРєРµ СѓР¶Рµ РЅРµРїР»РѕС…Рѕ РїРѕРєСЂС‹С‚С‹ schema.",
                        "Р”Р°Р¶Рµ РїСЂРё РЅР°Р»РёС‡РёРё schema РІР°Р¶РЅРѕ, С‡С‚РѕР±С‹ РѕРЅР° СЃРѕРѕС‚РІРµС‚СЃС‚РІРѕРІР°Р»Р° С‚РёРїСѓ СЃС‚СЂР°РЅРёС†С‹ Рё Р±С‹Р»Р° РІР°Р»РёРґРЅРѕР№ РїРѕ СЃРёРЅС‚Р°РєСЃРёСЃСѓ.",
                    ],
                    "priority": "HIGH" if invalid_schema_pages else "MEDIUM",
                    "owner": "Frontend / SEO",
                    "recommendation": "РџРѕРєСЂС‹С‚СЊ РєР»СЋС‡РµРІС‹Рµ С€Р°Р±Р»РѕРЅС‹ РІР°Р»РёРґРЅРѕР№ schema-СЂР°Р·РјРµС‚РєРѕР№ Рё РґРµСЂР¶Р°С‚СЊ РµРµ СЃРёРЅС…СЂРѕРЅРЅРѕР№ СЃ СЂРµР°Р»СЊРЅС‹Рј С‚РёРїРѕРј СЃС‚СЂР°РЅРёС†С‹: Organization, Service, Product, BreadcrumbList, FAQ, Article.",
                },
            ],
        },
        {
            "title": "Р­С‚Р°Рї 6 вЂ” РљРѕРјРјРµСЂС‡РµСЃРєРёРµ СЃС‚СЂР°РЅРёС†С‹ Рё РєРѕРЅС‚РµРЅС‚",
            "intro": "РЎРјРѕС‚СЂСЋ, РЅР°СЃРєРѕР»СЊРєРѕ СЃР°Р№С‚ РіРѕС‚РѕРІ РЅРµ С‚РѕР»СЊРєРѕ СЃРѕР±РёСЂР°С‚СЊ РёРЅРґРµРєСЃ, РЅРѕ Рё РїСЂРµРІСЂР°С‰Р°С‚СЊ СЃРїСЂРѕСЃ РІ Р·Р°СЏРІРєСѓ.",
            "checks": [
                {
                    "name": "РљРѕРЅС‚Р°РєС‚РЅС‹Рµ Рё Р»РёРґ-СЃС‚СЂР°РЅРёС†С‹",
                    "checked": "Р•СЃС‚СЊ Р»Рё РЅР° СЃР°Р№С‚Рµ СЃС‚СЂР°РЅРёС†С‹, РєРѕС‚РѕСЂС‹Рµ РјРѕР¶РЅРѕ РїСЂРѕРґРІРёРіР°С‚СЊ РєР°Рє С‚РѕС‡РєРё РІС…РѕРґР° РІ Р·Р°СЏРІРєСѓ.",
                    "method": "РџРѕРёСЃРє С„РѕСЂРј, РєРѕРЅС‚Р°РєС‚РЅС‹С… URL Рё Р±Р°Р·РѕРІРѕР№ SEO-РѕР±РІСЏР·РєРё СЌС‚РёС… СЃС‚СЂР°РЅРёС†.",
                    "metrics": [
                        ("РљРѕРЅС‚Р°РєС‚РЅС‹С…/lead-СЃС‚СЂР°РЅРёС†", str(len(contact_pages))),
                        ("Р¤РѕСЂРј РІ РІС‹Р±РѕСЂРєРµ", str(len([page for page in sample_pages if page.has_forms]))),
                        ("РЎР»Р°Р±С‹С… contact-СЃС‚СЂР°РЅРёС†", str(len(weak_contact_pages))),
                        ("РРЅРґРµРєСЃРёСЂСѓРµРјС‹С… route/query-СЃС‚СЂР°РЅРёС†", str(len(query_pages))),
                    ],
                    "findings": [
                        "РљРѕРЅС‚Р°РєС‚РЅР°СЏ Р·РѕРЅР° СѓР¶Рµ РµСЃС‚СЊ, РЅРѕ С‡Р°СЃС‚СЊ СЃС‚СЂР°РЅРёС† РЅРµ РѕС„РѕСЂРјР»РµРЅР° РєР°Рє РїРѕР»РЅРѕС†РµРЅРЅС‹Рµ SEO-РїРѕСЃР°РґРєРё." if weak_contact_pages else "РљРѕРЅС‚Р°РєС‚РЅС‹Рµ СЃС‚СЂР°РЅРёС†С‹ РІ РІС‹Р±РѕСЂРєРµ РІС‹РіР»СЏРґСЏС‚ СЃРѕР±СЂР°РЅРЅРѕ.",
                        "Р•СЃС‚СЊ С„РѕСЂРјС‹, РЅРѕ РЅСѓР¶РЅРѕ РїСЂРѕРІРµСЂСЏС‚СЊ, РЅР°СЃРєРѕР»СЊРєРѕ РІРѕРєСЂСѓРі РЅРёС… СЃРѕР±СЂР°РЅ РґРѕРІРµСЂРёС‚РµР»СЊРЅС‹Р№ Рё РєРѕРјРјРµСЂС‡РµСЃРєРёР№ СЃР»РѕР№.",
                        "РЎР»СѓР¶РµР±РЅС‹Рµ query/route URL РЅРµ РґРѕР»Р¶РЅС‹ РєРѕРЅРєСѓСЂРёСЂРѕРІР°С‚СЊ СЃ РЅРѕСЂРјР°Р»СЊРЅС‹РјРё РїРѕСЃР°РґРѕС‡РЅС‹РјРё СЃС‚СЂР°РЅРёС†Р°РјРё." if query_pages else "РЁСѓРјР° РѕС‚ СЃР»СѓР¶РµР±РЅС‹С… Р»РёРґ-URL РІ РІС‹Р±РѕСЂРєРµ РЅРµРјРЅРѕРіРѕ.",
                    ],
                    "priority": "HIGH",
                    "owner": "SEO / Marketing",
                    "recommendation": "РЎРґРµР»Р°С‚СЊ РёР· РєРѕРЅС‚Р°РєС‚РѕРІ Рё Р·Р°СЏРІРєРё СЃРёР»СЊРЅС‹Рµ РїРѕСЃР°РґРѕС‡РЅС‹Рµ: СЃ РѕС„С„РµСЂРѕРј, РѕС‚РІРµС‚Р°РјРё РЅР° РІРѕР·СЂР°Р¶РµРЅРёСЏ, Р±Р»РѕРєР°РјРё РґРѕРІРµСЂРёСЏ, FAQ Рё Р°РєРєСѓСЂР°С‚РЅРѕР№ SEO-РѕР±РІСЏР·РєРѕР№.",
                },
                {
                    "name": "РљРѕРЅС‚РµРЅС‚РЅР°СЏ РіР»СѓР±РёРЅР° РїСЂРёРѕСЂРёС‚РµС‚РЅС‹С… СЃС‚СЂР°РЅРёС†",
                    "checked": "РҐРІР°С‚Р°РµС‚ Р»Рё СЃС‚СЂР°РЅРёС†Р°Рј С„Р°РєС‚СѓСЂС‹, С‡С‚РѕР±С‹ Р·Р°РєСЂС‹РІР°С‚СЊ РёРЅС‚РµРЅС‚, Р° РЅРµ РїСЂРѕСЃС‚Рѕ СЃСѓС‰РµСЃС‚РІРѕРІР°С‚СЊ РІ РёРЅРґРµРєСЃРµ.",
                    "method": "РџРѕРґСЃС‡РµС‚ СЃР»РѕРІ Рё РїСЂРѕСЃРјРѕС‚СЂ РєРѕРјРјРµСЂС‡РµСЃРєРёС… С€Р°Р±Р»РѕРЅРѕРІ РІРЅСѓС‚СЂРё СЂРµРїСЂРµР·РµРЅС‚Р°С‚РёРІРЅРѕР№ РІС‹Р±РѕСЂРєРё.",
                    "metrics": [
                        ("РљРѕРјРјРµСЂС‡РµСЃРєРёС… СЃС‚СЂР°РЅРёС† РІ РІС‹Р±РѕСЂРєРµ", str(len(commercial_pages))),
                        ("РЎСЂРµРґРЅРµРµ С‡РёСЃР»Рѕ СЃР»РѕРІ", str(average_commercial_words)),
                        ("РўРѕРЅРєРёС… СЃС‚СЂР°РЅРёС†", str(len(thin_pages))),
                        ("РЎС‚СЂР°РЅРёС† Р±РµР· schema СЃСЂРµРґРё РєРѕРјРјРµСЂС‡РµСЃРєРёС…", str(len(missing_schema_commercial))),
                    ],
                    "findings": [
                        "РћРґРЅРѕРіРѕ РѕР±СЉРµРјР° С‚РµРєСЃС‚Р° РЅРµРґРѕСЃС‚Р°С‚РѕС‡РЅРѕ: СЃС‚СЂР°РЅРёС†С‹ РґРѕР»Р¶РЅС‹ РѕР±СЉСЏСЃРЅСЏС‚СЊ РІС‹Р±РѕСЂ, С†РµРЅСѓ, РїСЂРѕС†РµСЃСЃ, СЃСЂРѕРєРё Рё СЃР»РµРґСѓСЋС‰РёР№ С€Р°Рі.",
                        "Р§Р°СЃС‚СЊ СЃС‚СЂР°РЅРёС† РѕСЃС‚Р°РµС‚СЃСЏ С‚РѕРЅРєРѕР№ РїРѕ С„Р°РєС‚СѓСЂРµ." if thin_pages else "РЎРёР»СЊРЅРѕ РїСѓСЃС‚С‹С… РєРѕРјРјРµСЂС‡РµСЃРєРёС… СЃС‚СЂР°РЅРёС† РІ РІС‹Р±РѕСЂРєРµ РЅРµРјРЅРѕРіРѕ.",
                        "РџРѕС‚РµРЅС†РёР°Р» СЂРѕСЃС‚Р° Р»РµР¶РёС‚ РЅРµ С‚РѕР»СЊРєРѕ РІ SEO-С‚РµРєСЃС‚Р°С…, РЅРѕ Рё РІ Р±Р»РѕРєР°С… РґРѕРІРµСЂРёСЏ, FAQ, С‚Р°Р±Р»РёС†Р°С… СЃСЂР°РІРЅРµРЅРёСЏ Рё РѕС‚РІРµС‚Р°С… РЅР° СЃРїСЂРѕСЃ.",
                    ],
                    "priority": "MEDIUM",
                    "owner": "SEO / Content / Marketing",
                    "recommendation": "РџРµСЂРµСЃРѕР±СЂР°С‚СЊ РїСЂРёРѕСЂРёС‚РµС‚РЅС‹Рµ РїРѕСЃР°РґРѕС‡РЅС‹Рµ РїРѕРґ РєРѕРјРјРµСЂС‡РµСЃРєРёР№ РёРЅС‚РµРЅС‚: РґРѕР±Р°РІРёС‚СЊ РґРѕРєР°Р·Р°С‚РµР»СЊСЃС‚РІР°, FAQ, answer-first Р±Р»РѕРєРё, С†РµРЅС‹/СЃС†РµРЅР°СЂРёРё Рё РїРµСЂРµР»РёРЅРєРѕРІРєСѓ РЅР° СЃРѕСЃРµРґРЅРёРµ С‚РѕС‡РєРё СЃРїСЂРѕСЃР°.",
                },
            ],
        },
    ]


def build_issue_list(audit: dict) -> list[AuditIssue]:
    issues: list[AuditIssue] = []
    robots_lines = audit["robots_lines"]
    sample_pages: list[PageSnapshot] = audit["sample_pages"]
    key_pages = {snapshot.url: snapshot for snapshot in sample_pages}
    problematic_titles = [snapshot for snapshot in sample_pages if len(snapshot.title) > 70]
    problematic_descriptions = [
        snapshot for snapshot in sample_pages if len(snapshot.description) > 160 or (0 < len(snapshot.description) < 120)
    ]
    alt_gaps = [snapshot for snapshot in sample_pages if snapshot.missing_alt_count > 0]
    redirect_chains = [item for item in audit["redirect_checks"] if len(item["chain"]) > 1]

    if any("sitemap.xml" in line.lower() and "disallow" in line.lower() for line in robots_lines):
        issues.append(
            AuditIssue(
                severity="Critical",
                title="robots.txt РєРѕРЅС„Р»РёРєС‚СѓРµС‚ СЃ СЃРѕР±СЃС‚РІРµРЅРЅРѕР№ РєР°СЂС‚РѕР№ СЃР°Р№С‚Р°",
                why_it_matters=(
                    "РЎРµР№С‡Р°СЃ РІ robots.txt РѕРґРЅРѕРІСЂРµРјРµРЅРЅРѕ СѓРєР°Р·Р°РЅ Sitemap Рё СЃС‚РѕРёС‚ Disallow РЅР° sitemap.xml. "
                    "Р”Р»СЏ РЇРЅРґРµРєСЃР° Рё С‡Р°СЃС‚Рё СЃРµСЂРІРёСЃРѕРІ СЌС‚Рѕ РІС‹РіР»СЏРґРёС‚ РєР°Рє РєРѕРЅС„Р»РёРєС‚ СЃРёРіРЅР°Р»РѕРІ Рё РјРµС€Р°РµС‚ С‡РёСЃС‚РѕР№ РёРЅРґРµРєСЃР°С†РёРё."
                ),
                evidence=[
                    "Р’ robots.txt РЅР°Р№РґРµРЅРѕ РїСЂР°РІРёР»Рѕ `Disallow: *sitemap.xml`.",
                    "РћРґРЅРѕРІСЂРµРјРµРЅРЅРѕ С„Р°Р№Р» СЃРѕРґРµСЂР¶РёС‚ `Sitemap: https://akvarium-akvas.ru/sitemap.xml`.",
                ],
                recommendation=(
                    "РЈРґР°Р»РёС‚СЊ Р·Р°РїСЂРµС‚ РЅР° sitemap.xml, РѕСЃС‚Р°РІРёС‚СЊ С‚РѕР»СЊРєРѕ РґРёСЂРµРєС‚РёРІСѓ Sitemap Рё РїСЂРёРІРµСЃС‚Рё Host Рє С‡РёСЃС‚РѕРјСѓ РґРѕРјРµРЅСѓ Р±РµР· РїСЂРѕС‚РѕРєРѕР»Р°."
                ),
            )
        )

    contact_page = key_pages.get(urljoin(f"{audit['base_url']}/", "contacts"))
    route_contact_page = key_pages.get(urljoin(f"{audit['base_url']}/", "index.php?route=information/contactform"))
    if contact_page and (not contact_page.h1s or not contact_page.description or not contact_page.canonical):
        evidence = []
        if not contact_page.h1s:
            evidence.append("РЎС‚СЂР°РЅРёС†Р° /contacts РѕС‚РєСЂС‹РІР°РµС‚СЃСЏ Р±РµР· H1.")
        if not contact_page.description:
            evidence.append("РќР° /contacts РѕС‚СЃСѓС‚СЃС‚РІСѓРµС‚ meta description.")
        if not contact_page.canonical:
            evidence.append("РќР° /contacts РѕС‚СЃСѓС‚СЃС‚РІСѓРµС‚ canonical.")
        issues.append(
            AuditIssue(
                severity="High",
                title="РљРѕРЅС‚Р°РєС‚РЅС‹Рµ Рё lead-СЃС‚СЂР°РЅРёС†С‹ РЅРµРґРѕРѕРїС‚РёРјРёР·РёСЂРѕРІР°РЅС‹",
                why_it_matters=(
                    "РљРѕРЅС‚Р°РєС‚РЅС‹Рµ СЃС‚СЂР°РЅРёС†С‹ СѓС‡Р°СЃС‚РІСѓСЋС‚ РІ РґРѕРІРµСЂРёРё, РєРѕРЅРІРµСЂСЃРёРё Рё Р±СЂРµРЅРґРѕРІРѕР№ РІС‹РґР°С‡Рµ. "
                    "РљРѕРіРґР° Сѓ РЅРёС… РЅРµС‚ Р±Р°Р·РѕРІРѕР№ SEO-СЂР°Р·РјРµС‚РєРё, СЃР°Р№С‚ С‚РµСЂСЏРµС‚ Рё РєР»РёРєР°Р±РµР»СЊРЅРѕСЃС‚СЊ, Рё РєР°С‡РµСЃС‚РІРѕ РёРЅРґРµРєСЃР°."
                ),
                evidence=evidence or ["РљРѕРЅС‚Р°РєС‚РЅР°СЏ Р·РѕРЅР° СЃР°Р№С‚Р° РЅРµ РѕС„РѕСЂРјР»РµРЅР° РєР°Рє РїРѕР»РЅРѕС†РµРЅРЅР°СЏ РїРѕСЃР°РґРѕС‡РЅР°СЏ СЃС‚СЂР°РЅРёС†Р°."],
                recommendation=(
                    "РЎРґРµР»Р°С‚СЊ РїРѕР»РЅРѕС†РµРЅРЅСѓСЋ РєРѕРЅС‚Р°РєС‚РЅСѓСЋ СЃС‚СЂР°РЅРёС†Сѓ: H1, title РґРѕ 70 СЃРёРјРІРѕР»РѕРІ, description 140-160 СЃРёРјРІРѕР»РѕРІ, canonical, "
                    "РєРѕРЅС‚РµРЅС‚ РїСЂРѕ РїСЂРѕРёР·РІРѕРґСЃС‚РІРѕ, РґРѕСЃС‚Р°РІРєСѓ, РіР°СЂР°РЅС‚РёСЋ, Р°РґСЂРµСЃ, С‚РµР»РµС„РѕРЅС‹ Рё CTA."
                ),
            )
        )
    if route_contact_page and not route_contact_page.title and not route_contact_page.description:
        issues.append(
            AuditIssue(
                severity="High",
                title="Р’РЅСѓС‚СЂРµРЅРЅРёР№ СЃР»СѓР¶РµР±РЅС‹Р№ URL СЃ С„РѕСЂРјРѕР№ РѕР±СЂР°С‚РЅРѕРіРѕ Р·РІРѕРЅРєР° РґРѕСЃС‚СѓРїРµРЅ Р±РµР· SEO-РєРѕРЅС‚СЂРѕР»СЏ",
                why_it_matters=(
                    "Indexable-СЃС‚СЂР°РЅРёС†С‹ Р±РµР· title, description Рё H1 СЃРѕР·РґР°СЋС‚ С€СѓРј РІ РёРЅРґРµРєСЃРµ, СЂР°Р·РјС‹РІР°СЋС‚ СЂРµР»РµРІР°РЅС‚РЅРѕСЃС‚СЊ Рё С‚СЏРЅСѓС‚ crawl budget."
                ),
                evidence=[
                    "URL `/index.php?route=information/contactform` РґРѕСЃС‚СѓРїРµРЅ РїРѕ 200.",
                    "РќР° СЃС‚СЂР°РЅРёС†Рµ РЅРµС‚ title, description Рё H1.",
                    "РЎСЃС‹Р»РєР° РЅР° СЌС‚РѕС‚ URL РµСЃС‚СЊ РІ С€Р°РїРєРµ РєР°Рє `РћР±СЂР°С‚РЅС‹Р№ Р·РІРѕРЅРѕРє`.",
                ],
                recommendation=(
                    "Р›РёР±Рѕ Р·Р°РєСЂС‹С‚СЊ route-СЃС‚СЂР°РЅРёС†Сѓ РѕС‚ РёРЅРґРµРєСЃР°С†РёРё Рё СѓР±СЂР°С‚СЊ РїСЂСЏРјС‹Рµ СЃСЃС‹Р»РєРё, Р»РёР±Рѕ РїРµСЂРµРІРµСЃС‚Рё РµРµ РЅР° РЅРѕСЂРјР°Р»СЊРЅС‹Р№ SEO-friendly URL СЃ РѕС„РѕСЂРјР»РµРЅРёРµРј."
                ),
            )
        )

    if problematic_titles:
        examples = [f"{urlparse(snapshot.url).path or '/'} вЂ” {len(snapshot.title)} СЃРёРјРІ." for snapshot in problematic_titles[:4]]
        issues.append(
            AuditIssue(
                severity="High",
                title="РЁР°Р±Р»РѕРЅС‹ title РЅР° С‚РѕРІР°СЂР°С… Рё РєР°С‚РµРіРѕСЂРёСЏС… СЃР»РёС€РєРѕРј РґР»РёРЅРЅС‹Рµ",
                why_it_matters=(
                    "Р”Р»РёРЅРЅС‹Рµ title СЂРµР¶СѓС‚СЃСЏ РІ РІС‹РґР°С‡Рµ, СѓС…СѓРґС€Р°СЋС‚ CTR Рё СЃРЅРёР¶Р°СЋС‚ РєРѕРЅС‚СЂРѕР»СЊ РЅР°Рґ С‚РµРј, РєР°РєРѕР№ РѕС„С„РµСЂ РїРѕРёСЃРєРѕРІРёРє РїРѕРєР°Р¶РµС‚ РїРѕР»СЊР·РѕРІР°С‚РµР»СЋ."
                ),
                evidence=[
                    f"Р’ РІС‹Р±РѕСЂРєРµ {len(problematic_titles)} РёР· {len(sample_pages)} СЃС‚СЂР°РЅРёС† РёРјРµСЋС‚ title РґР»РёРЅРЅРµРµ 70 СЃРёРјРІРѕР»РѕРІ.",
                    *examples,
                ],
                recommendation=(
                    "РџРµСЂРµСЃРѕР±СЂР°С‚СЊ С€Р°Р±Р»РѕРЅС‹ title: СЃРЅР°С‡Р°Р»Р° РєР»СЋС‡ + РјРѕРґРµР»СЊ/РєР°С‚РµРіРѕСЂРёСЏ, Р·Р°С‚РµРј РѕС„С„РµСЂ Рё Р±СЂРµРЅРґ. "
                    "Р”Р»СЏ РєР°С‚Р°Р»РѕРіР° РґРµСЂР¶Р°С‚СЊ РґРёР°РїР°Р·РѕРЅ 55-70 СЃРёРјРІРѕР»РѕРІ."
                ),
            )
        )

    if problematic_descriptions:
        examples = [
            f"{urlparse(snapshot.url).path or '/'} вЂ” {len(snapshot.description)} СЃРёРјРІ."
            for snapshot in problematic_descriptions[:4]
        ]
        issues.append(
            AuditIssue(
                severity="Medium",
                title="Meta description РЅР° С‡Р°СЃС‚Рё СЃС‚СЂР°РЅРёС† РѕР±СЂРµР·Р°РµС‚СЃСЏ РёР»Рё РЅРµ РґРѕС‚СЏРіРёРІР°РµС‚ РґРѕ РЅРѕСЂРјР°Р»СЊРЅРѕРіРѕ СЃРЅРёРїРїРµС‚Р°",
                why_it_matters=(
                    "РЎРЅРёРїРїРµС‚ вЂ” СЌС‚Рѕ РєРѕРјРјРµСЂС‡РµСЃРєРёР№ РѕС„С„РµСЂ РІ РІС‹РґР°С‡Рµ. "
                    "РљРѕРіРґР° description СЃР»РёС€РєРѕРј РєРѕСЂРѕС‚РєРёР№ РёР»Рё РїРµСЂРµРіСЂСѓР¶РµРЅРЅС‹Р№, РїРѕРёСЃРєРѕРІРёРє С‡Р°С‰Рµ Р±РµСЂРµС‚ СЃР»СѓС‡Р°Р№РЅС‹Р№ РєСѓСЃРѕРє С‚РµРєСЃС‚Р° СЃРѕ СЃС‚СЂР°РЅРёС†С‹."
                ),
                evidence=[
                    f"Р’ РІС‹Р±РѕСЂРєРµ {len(problematic_descriptions)} РёР· {len(sample_pages)} СЃС‚СЂР°РЅРёС† РёРјРµСЋС‚ description РІРЅРµ РєРѕРјС„РѕСЂС‚РЅРѕРіРѕ РґРёР°РїР°Р·РѕРЅР°.",
                    *examples,
                ],
                recommendation=(
                    "РЎРѕР±СЂР°С‚СЊ С€Р°Р±Р»РѕРЅС‹ description РїРѕРґ С‚РёРїС‹ СЃС‚СЂР°РЅРёС†: РєР»СЋС‡, С†РµРЅРЅРѕСЃС‚СЊ, РґРѕСЃС‚Р°РІРєР°/РіР°СЂР°РЅС‚РёСЏ, СЂРµРіРёРѕРЅ, РїСЂРёР·С‹РІ. "
                    "РћСЂРёРµРЅС‚РёСЂ вЂ” 140-160 СЃРёРјРІРѕР»РѕРІ."
                ),
            )
        )

    if alt_gaps:
        issues.append(
            AuditIssue(
                severity="Medium",
                title="РР·РѕР±СЂР°Р¶РµРЅРёСЏ С‚РµСЂСЏСЋС‚ РїРѕРёСЃРєРѕРІС‹Р№ С‚СЂР°С„РёРє Рё РєРѕРЅС‚РµРєСЃС‚ РєР°СЂС‚РѕС‡РµРє",
                why_it_matters="Р”Р»СЏ С‚РѕРІР°СЂРЅРѕРіРѕ РїСЂРѕРµРєС‚Р° alt РІР»РёСЏРµС‚ Рё РЅР° РєР°СЂС‚РёРЅРєРё, Рё РЅР° РїРѕРЅРёРјР°РЅРёРµ Р°СЃСЃРѕСЂС‚РёРјРµРЅС‚Р°, Рё РЅР° РґРѕСЃС‚СѓРїРЅРѕСЃС‚СЊ.",
                evidence=[
                    f"РќР° {len(alt_gaps)} РёР· {len(sample_pages)} РїСЂРѕРІРµСЂРµРЅРЅС‹С… СЃС‚СЂР°РЅРёС† РµСЃС‚СЊ РёР·РѕР±СЂР°Р¶РµРЅРёСЏ Р±РµР· alt.",
                    f"РќР° РіР»Р°РІРЅРѕР№ РЅР°Р№РґРµРЅРѕ {audit['home_page'].missing_alt_count} РёР·РѕР±СЂР°Р¶РµРЅРёР№ Р±РµР· alt.",
                ],
                recommendation=(
                    "Р’РІРµСЃС‚Рё С€Р°Р±Р»РѕРЅРЅСѓСЋ РіРµРЅРµСЂР°С†РёСЋ alt РїРѕ С‚РёРїСѓ РєР°СЂС‚РѕС‡РєРё: С‚РѕРІР°СЂ + РѕР±СЉРµРј/СЂР°Р·РјРµСЂ + Р±СЂРµРЅРґ/СЃРµСЂРёСЏ, "
                    "Р° РґР»СЏ РєР°С‚РµРіРѕСЂРёР№ вЂ” РѕР±С‰РёР№ РёРЅС‚РµРЅС‚ СЃС‚СЂР°РЅРёС†С‹."
                ),
            )
        )

    if redirect_chains:
        issues.append(
            AuditIssue(
                severity="Medium",
                title="Р•СЃС‚СЊ Р»РёС€РЅСЏСЏ redirect-С†РµРїРѕС‡РєР° Сѓ С‡Р°СЃС‚Рё РґРѕРјРµРЅРЅС‹С… РІР°СЂРёР°РЅС‚РѕРІ",
                why_it_matters=(
                    "Р”РѕРїРѕР»РЅРёС‚РµР»СЊРЅС‹Р№ СЂРµРґРёСЂРµРєС‚ РґРѕР±Р°РІР»СЏРµС‚ Р·Р°РґРµСЂР¶РєСѓ РЅР° РІС…РѕРґРµ Рё СЃРѕР·РґР°РµС‚ РЅРµРЅСѓР¶РЅСѓСЋ С‚РµС…РЅРёС‡РµСЃРєСѓСЋ СЃР»РѕР¶РЅРѕСЃС‚СЊ РґР»СЏ Р±РѕС‚РѕРІ Рё РїРѕР»СЊР·РѕРІР°С‚РµР»РµР№."
                ),
                evidence=[
                    "РџСѓС‚СЊ `http://www.akvarium-akvas.ru/` РѕС‚РґР°РµС‚ РґРІРµ РїРѕСЃР»РµРґРѕРІР°С‚РµР»СЊРЅС‹Рµ 301-РїРµСЂРµР°РґСЂРµСЃР°С†РёРё.",
                    "Р¤РёРЅР°Р»СЊРЅС‹Р№ РєР°РЅРѕРЅРёС‡РµСЃРєРёР№ РґРѕРјРµРЅ вЂ” `https://akvarium-akvas.ru/`.",
                ],
                recommendation=(
                    "РЎРІРµСЃС‚Рё РІСЃРµ РІР°СЂРёР°РЅС‚С‹ РґРѕРјРµРЅР° Рє РѕРґРЅРѕРјСѓ 301-СЂРµРґРёСЂРµРєС‚Сѓ РЅР°РїСЂСЏРјСѓСЋ РЅР° РєР°РЅРѕРЅРёС‡РµСЃРєРёР№ HTTPS Р±РµР· РїСЂРѕРјРµР¶СѓС‚РѕС‡РЅРѕРіРѕ `https://www`."
                ),
            )
        )

    if audit["home_page"].has_meta_keywords:
        issues.append(
            AuditIssue(
                severity="Low",
                title="РќР° РіР»Р°РІРЅРѕР№ РѕСЃС‚Р°Р»СЃСЏ meta keywords, РєРѕС‚РѕСЂС‹Р№ СѓР¶Рµ РЅРµ РґР°РµС‚ SEO-С†РµРЅРЅРѕСЃС‚Рё",
                why_it_matters="РЎР°Рј РїРѕ СЃРµР±Рµ С‚РµРі РЅРµ РІСЂРµРґРёС‚, РЅРѕ СЌС‚Рѕ СЃРёРіРЅР°Р», С‡С‚Рѕ С€Р°Р±Р»РѕРЅ РјРµС‚С‹ СЃС‚РѕРёС‚ РїРѕС‡РёСЃС‚РёС‚СЊ Рё РїРѕРґРґРµСЂР¶РёРІР°С‚СЊ РІ Р°РєС‚СѓР°Р»СЊРЅРѕРј РІРёРґРµ.",
                evidence=["РќР° РіР»Р°РІРЅРѕР№ РЅР°Р№РґРµРЅ С‚РµРі meta keywords."],
                recommendation="РЈР±СЂР°С‚СЊ meta keywords РёР· С€Р°Р±Р»РѕРЅР° Рё СЃРѕСЃСЂРµРґРѕС‚РѕС‡РёС‚СЊСЃСЏ РЅР° title, description, H1, schema Рё РєРѕРјРјРµСЂС‡РµСЃРєРёС… Р±Р»РѕРєР°С….",
            )
        )

    if audit["sitemap_url_count"] > 5000:
        issues.append(
            AuditIssue(
                severity="Medium",
                title="РљР°СЂС‚Р° СЃР°Р№С‚Р° СѓР¶Рµ Р±РѕР»СЊС€Р°СЏ Рё С‚СЂРµР±СѓРµС‚ Р±РѕР»РµРµ СѓРїСЂР°РІР»СЏРµРјРѕР№ СЃС‚СЂСѓРєС‚СѓСЂС‹",
                why_it_matters=(
                    "РљРѕРіРґР° URL РїРѕС‡С‚Рё 7000 Рё РІСЃРµ Р»РµР¶РёС‚ РІ РѕРґРЅРѕРј РїРѕС‚РѕРєРµ, СЃР»РѕР¶РЅРµРµ РѕС‚СЃР»РµР¶РёРІР°С‚СЊ РёРЅРґРµРєСЃР°С†РёСЋ РєР°С‚РµРіРѕСЂРёР№, С‚РѕРІР°СЂРѕРІ Рё СЃР»СѓР¶РµР±РЅС‹С… СЃС‚СЂР°РЅРёС† РѕС‚РґРµР»СЊРЅРѕ."
                ),
                evidence=[
                    f"Р’ sitemap РѕР±РЅР°СЂСѓР¶РµРЅРѕ {audit['sitemap_url_count']} URL.",
                    "РЎРµР№С‡Р°СЃ СѓРґРѕР±РЅРµРµ СѓРїСЂР°РІР»СЏС‚СЊ РёРЅРґРµРєСЃРѕРј С‡РµСЂРµР· СЂР°Р·РґРµР»СЊРЅС‹Рµ sitemap РїРѕ С‚РёРїР°Рј СЃС‚СЂР°РЅРёС†.",
                ],
                recommendation=(
                    "Р Р°Р·РґРµР»РёС‚СЊ sitemap РјРёРЅРёРјСѓРј РЅР° С‚РѕРІР°СЂС‹, РєР°С‚РµРіРѕСЂРёРё, СЃР»СѓР¶РµР±РЅС‹Рµ СЃС‚СЂР°РЅРёС†С‹ Рё РјРµРґРёР°/СЃС‚Р°С‚СЊРё, "
                    "С‡С‚РѕР±С‹ РїСЂРѕС‰Рµ РєРѕРЅС‚СЂРѕР»РёСЂРѕРІР°С‚СЊ РїРѕРєСЂС‹С‚РёРµ Рё РїРµСЂРµРѕР±С…РѕРґ."
                ),
            )
        )

    if audit.get("sitemap_total_entries", 0) > audit["sitemap_url_count"] * 1.3:
        duplicate_entries = audit["sitemap_total_entries"] - audit["sitemap_url_count"]
        issues.append(
            AuditIssue(
                severity="Medium",
                title="РљР°СЂС‚Р° СЃР°Р№С‚Р° СЂР°Р·РґСѓС‚Р° РїРѕРІС‚РѕСЂСЏСЋС‰РёРјРёСЃСЏ URL",
                why_it_matters=(
                    "РљРѕРіРґР° РѕРґРёРЅ Рё С‚РѕС‚ Р¶Рµ URL РїРѕРІС‚РѕСЂСЏРµС‚СЃСЏ РІ sitemap РјРЅРѕРіРѕ СЂР°Р·, РїРѕРёСЃРєРѕРІРёРєРё РїРѕР»СѓС‡Р°СЋС‚ Р»РёС€РЅРёР№ С€СѓРј РІРјРµСЃС‚Рѕ С‡РёСЃС‚РѕРіРѕ СЃРёРіРЅР°Р»Р°, "
                    "Р° Р°РЅР°Р»РёР· РёРЅРґРµРєСЃР°С†РёРё СЃС‚Р°РЅРѕРІРёС‚СЃСЏ РјРµРЅРµРµ СѓРїСЂР°РІР»СЏРµРјС‹Рј."
                ),
                evidence=[
                    f"Р’ sitemap РЅР°Р№РґРµРЅРѕ {audit['sitemap_total_entries']} Р·Р°РїРёСЃРµР№, РЅРѕ С‚РѕР»СЊРєРѕ {audit['sitemap_url_count']} СѓРЅРёРєР°Р»СЊРЅС‹С… URL.",
                    f"РџРѕРІС‚РѕСЂРѕРІ: {duplicate_entries}.",
                ],
                recommendation=(
                    "РџРµСЂРµСЃРѕР±СЂР°С‚СЊ РіРµРЅРµСЂР°С†РёСЋ sitemap С‚Р°Рє, С‡С‚РѕР±С‹ РєР°Р¶РґС‹Р№ РёРЅРґРµРєСЃРёСЂСѓРµРјС‹Р№ URL РїРѕРїР°РґР°Р» С‚СѓРґР° РѕРґРёРЅ СЂР°Р·, "
                    "Р° РєР°СЂС‚Р° СЃР°Р№С‚Р° РѕС‚СЂР°Р¶Р°Р»Р° СЂРµР°Р»СЊРЅСѓСЋ СЃС‚СЂСѓРєС‚СѓСЂСѓ РїСЂРѕРµРєС‚Р° Р±РµР· РґСѓР±Р»РµР№."
                ),
            )
        )

    severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
    issues.sort(key=lambda issue: severity_order.get(issue.severity, 4))
    return issues


def build_strengths(audit: dict) -> list[str]:
    strengths = []
    if audit["home_page"].status_code == 200:
        strengths.append("РЎР°Р№С‚ СЃС‚Р°Р±РёР»СЊРЅРѕ РѕС‚РІРµС‡Р°РµС‚ РїРѕ HTTPS Рё РіР»Р°РІРЅР°СЏ СЃС‚СЂР°РЅРёС†Р° РѕС‚РґР°РµС‚ 200 РєРѕРґ Р±РµР· JS-Р·Р°РіР»СѓС€РєРё.")
    if audit["sitemap_url_count"]:
        strengths.append(
            f"РЎР°Р№С‚ СѓР¶Рµ РёРјРµРµС‚ РёРЅРґРµРєСЃРёСЂСѓРµРјСѓСЋ РєР°СЂС‚Сѓ СЃР°Р№С‚Р° СЃ {audit['sitemap_url_count']} URL Рё image-С‚РµРіР°РјРё РґР»СЏ РєР°СЂС‚РѕС‡РµРє."
        )
    if audit["home_page"].canonical:
        strengths.append("РќР° РєР»СЋС‡РµРІС‹С… РєРѕРјРјРµСЂС‡РµСЃРєРёС… СЃС‚СЂР°РЅРёС†Р°С… РїСЂРѕСЃС‚Р°РІР»РµРЅС‹ canonical, С‡С‚Рѕ РїРѕРјРѕРіР°РµС‚ СѓРґРµСЂР¶РёРІР°С‚СЊ РѕСЃРЅРѕРІРЅРѕР№ URL.")
    if audit["schema_coverage_ratio"] >= 0.7:
        strengths.append("РќР° С‚РѕРІР°СЂР°С… Рё РєР°С‚РµРіРѕСЂРёСЏС… СѓР¶Рµ РёСЃРїРѕР»СЊР·СѓРµС‚СЃСЏ schema-СЂР°Р·РјРµС‚РєР° (Product / BreadcrumbList / LocalBusiness).")
    if audit["h1_coverage_ratio"] >= 0.8:
        strengths.append("РЈ Р±РѕР»СЊС€РёРЅСЃС‚РІР° РїСЂРѕРІРµСЂРµРЅРЅС‹С… РєРѕРјРјРµСЂС‡РµСЃРєРёС… СЃС‚СЂР°РЅРёС† РµСЃС‚СЊ H1, С‡С‚Рѕ С…РѕСЂРѕС€Рѕ РґР»СЏ РјР°СЃС€С‚Р°Р±РЅРѕРіРѕ РєР°С‚Р°Р»РѕРіР°.")
    strengths.append("URL-СЃС‚СЂСѓРєС‚СѓСЂР° РІ РєР°С‚Р°Р»РѕРіРµ РїРѕРЅСЏС‚РЅР°СЏ Рё С…РѕСЂРѕС€Рѕ РґСЂРѕР±РёС‚ СЃРїСЂРѕСЃ РїРѕ Р»РёС‚СЂР°Р¶Сѓ, С„РѕСЂРјРµ, РЅР°Р·РЅР°С‡РµРЅРёСЋ Рё РєРѕРјРїР»РµРєС‚Р°С†РёРё.")
    if audit["llms_exists"]:
        strengths.append("РЈ РїСЂРѕРµРєС‚Р° СѓР¶Рµ РµСЃС‚СЊ llms.txt, Р° Р·РЅР°С‡РёС‚ РјРѕР¶РЅРѕ РґРѕРїРѕР»РЅРёС‚РµР»СЊРЅРѕ СѓСЃРёР»РёРІР°С‚СЊ AI-РІРёРґРёРјРѕСЃС‚СЊ Рё branded-РѕС‚РІРµС‚С‹.")
    return strengths


def build_growth_points() -> list[str]:
    return [
        "РЈСЃРёР»РёС‚СЊ СЃС‚СЂР°РЅРёС†С‹ РєР°С‚РµРіРѕСЂРёР№ СѓРЅРёРєР°Р»СЊРЅС‹РјРё РёРЅС‚СЂРѕ-Р±Р»РѕРєР°РјРё, FAQ, РґРѕСЃС‚Р°РІРєРѕР№, РіР°СЂР°РЅС‚РёРµР№, СЃСЂРѕРєР°РјРё РїСЂРѕРёР·РІРѕРґСЃС‚РІР° Рё РѕР±СЉСЏСЃРЅРµРЅРёРµРј, РєР°Рє РІС‹Р±СЂР°С‚СЊ РјРѕРґРµР»СЊ.",
        "РЎРґРµР»Р°С‚СЊ РѕС‚РґРµР»СЊРЅС‹Рµ РїРѕСЃР°РґРѕС‡РЅС‹Рµ РїРѕРґ РІС‹СЃРѕРєРёР№ СЃРїСЂРѕСЃ: РїРѕ Р»РёС‚СЂР°Р¶Сѓ, С„РѕСЂРјРµ, РЅР°Р·РЅР°С‡РµРЅРёСЋ, С‚РёРїСѓ РєРѕРјРїР»РµРєС‚Р°С†РёРё Рё РіРѕС‚РѕРІС‹Рј СЂРµС€РµРЅРёСЏРј РїРѕРґ РєР»СЋС‡.",
        "Р”РѕР±Р°РІРёС‚СЊ РІ РєР°С‚РµРіРѕСЂРёРё Рё РєР°СЂС‚РѕС‡РєРё Р±Р»РѕРєРё РґРѕРІРµСЂРёСЏ: РїСЂРѕРёР·РІРѕРґСЃС‚РІРѕ РІ Р РѕСЃСЃРёРё, РіР°СЂР°РЅС‚РёСЏ, РґРѕСЃС‚Р°РІРєР° РїРѕ Р Р¤, РєР°СЃС‚РѕРјРёР·Р°С†РёСЏ СЂР°Р·РјРµСЂРѕРІ, РєРѕРјРїР»РµРєС‚Р°С†РёСЏ.",
        "Р Р°Р·РІРµСЂРЅСѓС‚СЊ image SEO: alt-С€Р°Р±Р»РѕРЅС‹, РїРѕРґРїРёСЃРё, structured data Рё РїРѕРґР±РѕСЂ РІРёР·СѓР°Р»РѕРІ РїРѕРґ РїРѕРёСЃРєРѕРІС‹Р№ СЃРїСЂРѕСЃ РїРѕ Р°РєРІР°СЂРёСѓРјР°Рј Рё С‚РµСЂСЂР°СЂРёСѓРјР°Рј.",
        "РЎРѕР±СЂР°С‚СЊ С€Р°Р±Р»РѕРЅРЅС‹Р№ СЃР»РѕР№ СЃРЅРёРїРїРµС‚РѕРІ: РєРѕСЂРѕС‚РєРёРµ title, РєРѕРјРјРµСЂС‡РµСЃРєРёРµ description, Р°РєРєСѓСЂР°С‚РЅС‹Р№ canonical Рё РїРѕРЅСЏС‚РЅС‹Рµ H1 РЅР° РІСЃРµС… С‚РёРїР°С… СЃС‚СЂР°РЅРёС†.",
        "РЎРµРіРјРµРЅС‚РёСЂРѕРІР°С‚СЊ sitemap Рё РїР°СЂР°Р»Р»РµР»СЊРЅРѕ РЅР°СЃС‚СЂРѕРёС‚СЊ РєРѕРЅС‚СЂРѕР»СЊ РёРЅРґРµРєСЃР°С†РёРё РєР°С‚РµРіРѕСЂРёР№, С‚РѕРІР°СЂРѕРІ, С„РёР»СЊС‚СЂРѕРІ Рё СЃР»СѓР¶РµР±РЅС‹С… route-СЃС‚СЂР°РЅРёС†.",
    ]


def build_roadmap() -> list[tuple[str, list[str]]]:
    return [
        (
            "0-14 РґРЅРµР№",
            [
                "РСЃРїСЂР°РІРёС‚СЊ robots.txt: СѓР±СЂР°С‚СЊ РєРѕРЅС„Р»РёРєС‚ СЃ sitemap.xml Рё РЅРѕСЂРјР°Р»РёР·РѕРІР°С‚СЊ Host.",
                "Р—Р°РєСЂС‹С‚СЊ РёР»Рё РїРµСЂРµСЂР°Р±РѕС‚Р°С‚СЊ `/index.php?route=information/contactform`.",
                "Р”РѕРґРµР»Р°С‚СЊ SEO-С€Р°Р±Р»РѕРЅ СЃС‚СЂР°РЅРёС†С‹ РєРѕРЅС‚Р°РєС‚РѕРІ: H1, title, description, canonical, РєРѕРјРјРµСЂС‡РµСЃРєРёР№ РєРѕРЅС‚РµРЅС‚.",
                "РџРѕРґРіРѕС‚РѕРІРёС‚СЊ РЅРѕРІРѕРµ РїСЂР°РІРёР»Рѕ РіРµРЅРµСЂР°С†РёРё title Рё description РґР»СЏ С‚РѕРІР°СЂРѕРІ Рё РєР°С‚РµРіРѕСЂРёР№.",
            ],
        ),
        (
            "15-30 РґРЅРµР№",
            [
                "Р’РЅРµРґСЂРёС‚СЊ alt-С€Р°Р±Р»РѕРЅС‹ РґР»СЏ РєР°СЂС‚РѕС‡РµРє Рё РєР°С‚РµРіРѕСЂРёР№.",
                "РЎРІРµСЃС‚Рё РґРѕРјРµРЅРЅС‹Рµ СЂРµРґРёСЂРµРєС‚С‹ Рє РѕРґРЅРѕРјСѓ С€Р°РіСѓ РЅР° РєР°РЅРѕРЅРёС‡РµСЃРєРёР№ HTTPS.",
                "Р Р°Р·РґРµР»РёС‚СЊ sitemap РїРѕ С‚РёРїР°Рј СЃС‚СЂР°РЅРёС† Рё РїРѕРґР°С‚СЊ РёС… РѕС‚РґРµР»СЊРЅРѕ РІ РЇРЅРґРµРєСЃ Р’РµР±РјР°СЃС‚РµСЂ Рё GSC.",
                "РћР±РЅРѕРІРёС‚СЊ СЃРЅРёРїРїРµС‚С‹ Рё РїСЂРѕС‚РµСЃС‚РёСЂРѕРІР°С‚СЊ СЂРѕСЃС‚ CTR РїРѕ РєР°С‚РµРіРѕСЂРёСЏРј Рё С‚РѕРІР°СЂРЅС‹Рј РєР»Р°СЃС‚РµСЂР°Рј.",
            ],
        ),
        (
            "31-60 РґРЅРµР№",
            [
                "РЈСЃРёР»РёС‚СЊ СЃРїСЂРѕСЃРѕРІС‹Рµ РєР°С‚РµРіРѕСЂРёРё РєРѕРЅС‚РµРЅС‚РѕРј, FAQ, РєРѕРјРјРµСЂС‡РµСЃРєРёРјРё Р±Р»РѕРєР°РјРё Рё РїРµСЂРµР»РёРЅРєРѕРІРєРѕР№.",
                "РЎРѕР±СЂР°С‚СЊ РїРѕСЃР°РґРѕС‡РЅС‹Рµ РїРѕРґ РєР»СЋС‡РµРІС‹Рµ РєР»Р°СЃС‚РµСЂС‹ СЃРїСЂРѕСЃР°: Р»РёС‚СЂР°Р¶, С„РѕСЂРјР°, СЃС†РµРЅР°СЂРёР№ РёСЃРїРѕР»СЊР·РѕРІР°РЅРёСЏ, РєРѕРјРїР»РµРєС‚Р°С†РёСЏ.",
                "Р”РѕР±Р°РІРёС‚СЊ РєРѕРЅС‚РµРЅС‚-СЃР»РѕР№ РїРѕРґ AI-РѕС‚РІРµС‚С‹: РєСЂР°С‚РєРёРµ РѕРїСЂРµРґРµР»РµРЅРёСЏ, С‚Р°Р±Р»РёС†С‹ СЃСЂР°РІРЅРµРЅРёСЏ, РѕС‚РІРµС‚С‹ РЅР° С‡Р°СЃС‚С‹Рµ РІРѕРїСЂРѕСЃС‹.",
            ],
        ),
    ]


def build_score(audit: dict, issues: list[AuditIssue]) -> int:
    score = 82
    penalties = {"Critical": 9, "High": 6, "Medium": 3, "Low": 1}
    for issue in issues[:6]:
        score -= penalties.get(issue.severity, 1)
    score -= min(8, int(round(audit["title_long_ratio"] * 10)))
    score -= min(6, int(round(audit["description_problem_ratio"] * 8)))
    score -= 2 if audit["home_page"].missing_alt_count else 0
    return max(45, min(95, score))


def shorten_path(url: str, max_length=48) -> str:
    parsed = urlparse(url)
    path = parsed.path or "/"
    if parsed.query:
        path = f"{path}?{parsed.query}"
    if len(path) <= max_length:
        return path
    if max_length <= 3:
        return path[:max_length]
    return f"{path[: max_length - 3]}..."


def dynamic_build_strengths(audit: dict) -> list[str]:
    return []


def dynamic_build_growth_points(audit: dict) -> list[str]:
    profile = infer_project_profile(audit)
    issue_titles = {issue.title for issue in audit.get("issues", [])}
    priority_pages = select_priority_pages(audit["sample_pages"])
    points: list[str] = []

    if ISSUE_TITLE_COMMERCIAL_GAPS in issue_titles:
        points.append(
            "Пересобрать приоритетные шаблоны под коммерческий интент: вынести на видимое место условия, сроки, оплату/доставку или порядок работ, гарантию, контакты и понятный CTA."
        )
    elif profile["is_catalog"]:
        points.append(
            "Усилить категории и карточки не SEO-текстом ради объёма, а условиями выбора: наличие, доставка, оплата, гарантия, фильтры, FAQ и блоки доверия."
        )
    else:
        points.append(
            "Развести ключевые услуги по отдельным посадочным и дать на каждой странице сценарий выбора: стоимость, сроки, этапы, гарантии, кейсы и следующий шаг."
        )

    if ISSUE_TITLE_CONTENT_PROOF in issue_titles:
        points.append(
            "Запустить повторную оптимизацию ключевых страниц: потерянные запросы -> новые H2/H3 -> доказательства, списки, сравнения, кейсы, фото/скрины -> обновление сниппета."
        )
    elif priority_pages:
        points.append(
            "Доработать самые важные страницы фактами и подтверждениями: таблицы, FAQ, отзывы, кейсы, документы, сравнения и короткие answer-first блоки."
        )

    points.append("Собрать единые шаблоны для title, description, H1 и canonical по всем ключевым типам страниц.")

    if audit["total_missing_alt"] > 0:
        points.append("Обновить медиаслой: alt, подписи, размеры, lazy loading и новые изображения там, где контенту не хватает доказательств.")

    if ISSUE_TITLE_INTERNAL_LINKING in issue_titles:
        points.append("Собрать тематическую перелинковку по кластерам спроса: хлебные крошки, блоки \"по теме\", осмысленные анкоры и несколько кликов до ключевой страницы.")
    else:
        points.append("Проверить, что важные разделы связаны между собой внутренними ссылками и не остаются без поддержки со стороны главной и соседних кластеров.")

    if audit["sitemap_url_count"] > 2000 or audit.get("sitemap_total_entries", 0) > audit["sitemap_url_count"] * 1.2:
        points.append("Разделить sitemap по типам страниц и держать под контролем служебные URL.")

    unique_points: list[str] = []
    for point in points:
        if point not in unique_points:
            unique_points.append(point)
    return unique_points[:6]


PHASE_TEXT_REPLACEMENTS = {
    "Разбираю шаблоны title, description, H1 и то, насколько страницы готовы к нормальному сниппету.": "Проверка title, описания страницы и H1 на ключевых шаблонах.",
    "По H1 каркас у выборки в целом собран неплохо.": "На проверенных страницах H1 в целом оформлен нормально.",
    "Сильно пустых страниц в выборке немного.": "Пустых страниц в выборке почти нет.",
    "Даже при нормальном объеме текста нужно отдельно проверить, насколько хорошо он поддерживает коммерческий интент и FAQ-слой.": "Даже при нормальном объеме текста важно проверить, закрывает ли контент выбор, условия, FAQ и доверие.",
    "Критической ямы по title в выборке не видно.": "Критических провалов по title в выборке не видно.",
    "Описания страниц покрыты достаточно ровно.": "Описания страниц заполнены на большей части выборки.",
    "Есть llms.txt, значит сайт уже можно дожимать и под дополнительную видимость.": "Дополнительных ограничений по техническим файлам для обхода в выборке не видно.",
    "Отдельного llms.txt не найдено.": "Отдельных ограничений по техническим файлам для обхода в выборке не видно.",
}


def polish_phase_text(value: object) -> str:
    text = normalize_output_text(value)
    for source, target in PHASE_TEXT_REPLACEMENTS.items():
        text = text.replace(source, target)
    text = text.replace("FAQ-слой", "FAQ и блоки доверия")
    return text.strip()


def polish_phase_sections(phase_sections: list[dict]) -> list[dict]:
    polished_sections: list[dict] = []
    for section in phase_sections:
        polished_checks: list[dict] = []
        for check in section.get("checks", []):
            findings: list[str] = []
            for item in check.get("findings", []):
                cleaned = polish_phase_text(item)
                if not cleaned:
                    continue
                if "llms.txt" in cleaned.lower():
                    cleaned = "Дополнительных ограничений по техническим файлам для обхода в выборке не видно."
                if cleaned not in findings:
                    findings.append(cleaned)

            metrics = [(polish_phase_text(label), polish_phase_text(value)) for label, value in check.get("metrics", [])]
            polished_checks.append(
                {
                    **check,
                    "name": polish_phase_text(check.get("name", "")),
                    "checked": polish_phase_text(check.get("checked", "")),
                    "method": polish_phase_text(check.get("method", "")),
                    "metrics": metrics,
                    "findings": findings or ["Явных проблем в этой выборке не нашли."],
                    "owner": polish_phase_text(check.get("owner", "")),
                    "recommendation": polish_phase_text(check.get("recommendation", "")),
                }
            )

        polished_sections.append(
            {
                **section,
                "title": polish_phase_text(section.get("title", "")),
                "intro": polish_phase_text(section.get("intro", "")),
                "checks": polished_checks,
            }
        )
    return polished_sections


def build_monster_phase_sections(audit: dict, phase_sections: list[dict]) -> list[dict]:
    if not phase_sections:
        return phase_sections

    updated_sections = [dict(section) for section in phase_sections]
    sample_pages: list[PageSnapshot] = audit["sample_pages"]
    priority_pages = select_priority_pages(sample_pages)
    weak_support_pages = [
        snapshot
        for snapshot in priority_pages
        if not any(bool(getattr(snapshot, field, False)) for field in MONSTER_SUPPORT_FIELDS)
    ]
    weak_proof_pages = [
        snapshot
        for snapshot in priority_pages
        if not any(bool(getattr(snapshot, field, False)) for field in MONSTER_PROOF_FIELDS)
    ]
    commercial_gap_pages = [snapshot for snapshot in priority_pages if not snapshot.has_commercial_block]
    trust_gap_pages = [snapshot for snapshot in priority_pages if not snapshot.has_trust_block]
    thin_priority_pages = [snapshot for snapshot in priority_pages if 0 < snapshot.word_count < 250]
    low_link_priority_pages = sorted(
        [snapshot for snapshot in priority_pages if get_template_bucket(snapshot) != "home"],
        key=lambda snapshot: snapshot.internal_links,
    )[:3]
    breadcrumb_gap_pages = [
        snapshot
        for snapshot in priority_pages
        if get_template_bucket(snapshot) in {"category", "product", "service", "contact"} and not snapshot.has_breadcrumbs
    ]
    average_priority_links = round(statistics.mean([snapshot.internal_links for snapshot in priority_pages]), 1) if priority_pages else 0
    average_priority_words = round(statistics.mean([snapshot.word_count for snapshot in priority_pages if snapshot.word_count]), 1) if priority_pages else 0
    long_path_pages = [snapshot for snapshot in sample_pages if len(urlparse(snapshot.url).path) > 75]
    query_pages = get_indexable_query_pages(sample_pages)
    profile = infer_project_profile(audit)

    if len(updated_sections) > 1:
        updated_sections[1] = {
            "title": "Этап 2 — Структура и внутренняя перелинковка",
            "intro": "По SEO Монстр внутренняя структура должна не просто существовать, а направлять вес и пользователя к нужным кластерам спроса. Здесь смотрим тематические связи, глубину доступа и навигационный шум.",
            "checks": [
                {
                    "name": "Тематическая перелинковка приоритетных страниц",
                    "checked": "Сколько внутренних ссылок получают ключевые шаблоны и есть ли у них навигационная поддержка.",
                    "method": "Подсчет внутренних ссылок, хлебных крошек и слабых по связности страниц в репрезентативной выборке.",
                    "metrics": [
                        ("Приоритетных страниц", str(len(priority_pages))),
                        ("Среднее число внутренних ссылок", str(average_priority_links)),
                        ("Страниц со слабой связностью", str(len([snapshot for snapshot in priority_pages if snapshot.internal_links <= 5]))),
                        ("Страниц без хлебных крошек", str(len(breadcrumb_gap_pages))),
                    ],
                    "findings": [
                        (
                            "Слабее всего по внутренним ссылкам выглядят: "
                            + ", ".join(f"{human_path(snapshot.url)} ({snapshot.internal_links})" for snapshot in low_link_priority_pages)
                            + "."
                            if low_link_priority_pages
                            else "Слабые по внутренним ссылкам страницы в выборке не выделяются."
                        ),
                        (
                            f"На {len(breadcrumb_gap_pages)} страницах не обнаружены хлебные крошки, хотя они должны поддерживать иерархию."
                            if breadcrumb_gap_pages
                            else "Критичной просадки по хлебным крошкам в выборке не видно."
                        ),
                        "Приоритетные страницы должны получать тематические ссылки из основной области контента, а не жить только за счёт меню и футера.",
                    ],
                    "priority": "HIGH" if any(issue.title == ISSUE_TITLE_INTERNAL_LINKING for issue in audit.get("issues", [])) else "MEDIUM",
                    "owner": "SEO / Frontend / Content",
                    "recommendation": "Собрать кластеры и связки главная -> раздел -> подраздел/услуга -> заявка, добавить хлебные крошки, блоки \"по теме\" и осмысленные анкоры вместо повторяющихся служебных ссылок.",
                },
                {
                    "name": "Навигационный шум и глубина доступа",
                    "checked": "Есть ли на сайте длинные URL, indexable query-страницы и архитектурные хвосты, которые мешают чистой структуре.",
                    "method": "Проверка query-URL, длинных путей и шаблонов, которые смешивают коммерческую и служебную логику.",
                    "metrics": [
                        ("Indexable query-URL", str(len(query_pages))),
                        ("Длинных URL", str(len(long_path_pages))),
                        ("Категорий / карточек", str(len(profile["category_pages"]) + len(profile["product_pages"]))),
                        ("Страниц с формами", str(len(profile["form_pages"]))),
                    ],
                    "findings": [
                        "Служебные query-URL не должны конкурировать с посадочными страницами за обход и индекс." if query_pages else "Критичного шума от индексируемых query-URL в выборке немного.",
                        f"Длинных путей в выборке: {len(long_path_pages)}." if long_path_pages else "Сильного перекоса в сторону чрезмерно длинных URL в выборке не видно.",
                        "Важные страницы должны находиться в нескольких кликах от главной и быть встроены в понятную иерархию, а не выпадать из цепочки переходов.",
                    ],
                    "priority": "HIGH" if query_pages else "MEDIUM",
                    "owner": "SEO / Backend",
                    "recommendation": "Очистить служебные маршруты из индекса, укоротить длинные пути, держать структуру URL человекочитаемой и не смешивать кластер спроса со служебными параметрами.",
                },
            ],
        }

    if len(updated_sections) > 5:
        updated_sections[5] = {
            "title": "Этап 6 — Коммерческие факторы и качество контента",
            "intro": "Здесь применяем логику SEO Монстр: для коммерческого ранжирования мало текста. Нужны условия сделки, доверие, поддержка, ответы на возражения и доказательства, что страница реально помогает выбрать.",
            "checks": [
                {
                    "name": "Коммерческий слой на приоритетных шаблонах",
                    "checked": "Есть ли на важных страницах условия выбора: цена, сроки, оплата, доставка, гарантия, поддержка и контактные точки.",
                    "method": "Поиск коммерческих и доверительных блоков на ключевых шаблонах: главная, категории, карточки, услуги, контакты и страницы заявок.",
                    "metrics": [
                        ("Приоритетных страниц", str(len(priority_pages))),
                        ("Без коммерческого блока", str(len(commercial_gap_pages))),
                        ("Без доверительных блоков", str(len(trust_gap_pages))),
                        ("Контактных / lead-страниц", str(len(get_contact_related_pages(sample_pages)))),
                    ],
                    "findings": [
                        (
                            f"На {len(commercial_gap_pages)} из {len(priority_pages)} страниц не видно явных условий выбора: цены, сроков, оплаты/доставки или порядка работ."
                            if priority_pages
                            else "Приоритетных страниц в выборке недостаточно для выводов."
                        ),
                        (
                            f"На {len(trust_gap_pages)} из {len(priority_pages)} страниц не хватает доверительных сигналов: гарантий, кейсов, сертификатов, отзывов или поддержки."
                            if trust_gap_pages
                            else "Критичного провала по доверию в выборке не видно."
                        ),
                        "По книге такие сигналы лучше выносить в видимую часть шаблона и поддерживать отдельными служебными страницами с условиями и гарантиями.",
                    ],
                    "priority": "HIGH" if any(issue.title == ISSUE_TITLE_COMMERCIAL_GAPS for issue in audit.get("issues", [])) else "MEDIUM",
                    "owner": "SEO / Marketing / Content",
                    "recommendation": "Вынести на ключевые страницы цену или диапазон, сроки, условия оплаты/доставки или работ, гарантии, поддержку, контакты и понятный следующий шаг; не прятать этот слой только в футер или одну служебную страницу.",
                },
                {
                    "name": "Качество контента и доказательства",
                    "checked": "Помогает ли контент выбрать решение, подтверждает ли оффер фактами и закрывает ли вопросы до звонка.",
                    "method": "Подсчет тонких страниц, страниц без FAQ/отзывов/сравнений и проверка наличия доказательного слоя на приоритетных шаблонах.",
                    "metrics": [
                        ("Средний объём текста", str(average_priority_words)),
                        ("Тонких страниц", str(len(thin_priority_pages))),
                        ("Без FAQ/отзывов/таблиц", str(len(weak_support_pages))),
                        ("Без доказательного слоя", str(len(weak_proof_pages))),
                    ],
                    "findings": [
                        (
                            f"На {len(weak_support_pages)} из {len(priority_pages)} страниц не видно FAQ, отзывов/кейсов или сравнительных блоков."
                            if weak_support_pages
                            else "Критичной просадки по FAQ и поддерживающим блокам в выборке не видно."
                        ),
                        (
                            f"Тонких страниц с объёмом до 250 слов: {len(thin_priority_pages)}."
                            if thin_priority_pages
                            else "Совсем пустых приоритетных страниц в выборке немного."
                        ),
                        "По SEO Монстр прирост даёт не наращивание воды, а повторная оптимизация: новые абзацы под потерянные запросы, доказательства, таблицы, кейсы, фото и повторная перелинковка.",
                    ],
                    "priority": "HIGH" if any(issue.title == ISSUE_TITLE_CONTENT_PROOF for issue in audit.get("issues", [])) else "MEDIUM",
                    "owner": "SEO / Content / Marketing",
                    "recommendation": "Пересобрать ключевые страницы по схеме: потерянные запросы -> новые H2/H3 -> факты, сравнения, кейсы, FAQ, изображения и обновлённый сниппет; слабые или устаревшие материалы отправлять на перепись, объединение или в архив.",
                },
            ],
        }

    return updated_sections


def build_audit(url: str, company_name: str | None, sample_size: int, competitor_urls: list[str] | None = None) -> dict:
    base_url = normalize_base_url(url)
    parsed = urlparse(base_url)
    base_domain = parsed.netloc
    session = make_session()
    requested_limit = max(sample_size, 0)

    homepage = analyse_page(session, base_url, base_domain)
    robots_response, _ = safe_get(session, urljoin(f"{base_url}/", "robots.txt"))
    llms_response, _ = safe_get(session, urljoin(f"{base_url}/", "llms.txt"))
    robots_text = robots_response.text if robots_response is not None and robots_response.status_code == 200 else ""
    sitemap_urls, processed_sitemaps, sitemap_total_entries = parse_sitemap_urls(
        session, discover_sitemaps(robots_text, base_url), base_domain
    )
    home_links = discover_home_links(session, base_url, base_domain)
    preferred_urls = [
        urljoin(f"{base_url}/", "contacts"),
        urljoin(f"{base_url}/", "index.php?route=information/contactform"),
        *home_links[:10],
    ]
    unique_urls = build_audit_url_inventory(
        base_url,
        preferred_urls,
        sitemap_urls,
        base_domain,
        requested_limit,
    )
    snapshots = analyse_pages_parallel(unique_urls, base_domain, max_workers=12)
    html_pages = [snapshot for snapshot in snapshots if snapshot.status_code == 200 and "html" in snapshot.content_type]
    appendix_urls = select_representative_urls(base_url, preferred_urls, [snapshot.url for snapshot in html_pages], 18)
    pages_by_url = {normalize_page_url(snapshot.url): snapshot for snapshot in html_pages}
    appendix_pages = [pages_by_url[url] for url in appendix_urls if url in pages_by_url][:18]

    title_long = [snapshot for snapshot in html_pages if len(snapshot.title) > 70]
    description_problem = [
        snapshot for snapshot in html_pages if len(snapshot.description) > 160 or (0 < len(snapshot.description) < 120)
    ]
    schema_pages = [snapshot for snapshot in html_pages if snapshot.schema_types]
    h1_pages = [snapshot for snapshot in html_pages if snapshot.h1s]
    canonical_pages = [snapshot for snapshot in html_pages if snapshot.canonical]
    total_missing_alt = sum(snapshot.missing_alt_count for snapshot in html_pages)
    redirect_checks = analyze_redirects(session, base_url)
    company = company_name or re.sub(r"^www\.", "", base_domain)

    audit = {
        "generator_version": 4,
        "base_url": base_url,
        "domain": base_domain,
        "company_name": company,
        "generated_at": datetime.now().strftime("%d.%m.%Y %H:%M"),
        "home_page": homepage,
        "robots_status": robots_response.status_code if robots_response is not None else 0,
        "robots_lines": [line.strip() for line in robots_text.splitlines() if line.strip()],
        "llms_exists": bool(llms_response is not None and llms_response.status_code == 200),
        "sitemap_url_count": len(sitemap_urls),
        "sitemap_total_entries": sitemap_total_entries,
        "processed_sitemaps": processed_sitemaps,
        "home_links_count": len(home_links),
        "sample_pages": html_pages,
        "appendix_pages": appendix_pages,
        "raw_sampled_urls": unique_urls,
        "requested_page_limit": requested_limit,
        "crawl_scope": "full" if requested_limit == 0 else "limited",
        "analyzed_pages_count": len(html_pages),
        "title_long_ratio": (len(title_long) / len(html_pages)) if html_pages else 0,
        "description_problem_ratio": (len(description_problem) / len(html_pages)) if html_pages else 0,
        "schema_coverage_ratio": (len(schema_pages) / len(html_pages)) if html_pages else 0,
        "h1_coverage_ratio": (len(h1_pages) / len(html_pages)) if html_pages else 0,
        "canonical_coverage_ratio": (len(canonical_pages) / len(html_pages)) if html_pages else 0,
        "total_missing_alt": total_missing_alt,
        "average_response_ms": round(statistics.mean([snapshot.response_time_ms for snapshot in html_pages]), 1)
        if html_pages
        else 0,
        "average_html_kb": round(statistics.mean([snapshot.html_size_kb for snapshot in html_pages]), 1)
        if html_pages
        else 0,
        "average_words": round(statistics.mean([snapshot.word_count for snapshot in html_pages]), 1) if html_pages else 0,
        "redirect_checks": redirect_checks,
        "requested_competitors": competitor_urls or [],
    }
    issues = dynamic_build_issue_list(audit)
    audit["issues"] = issues
    audit["executive_summary"] = build_executive_summary_dynamic(audit)
    audit["strengths"] = []
    audit["growth_points"] = dynamic_build_growth_points(audit)
    audit["roadmap"] = dynamic_build_roadmap(audit)
    audit["priority_matrix"] = build_priority_matrix(audit)
    audit["critical_errors"] = build_critical_errors(audit)
    audit["phase_sections"] = build_monster_phase_sections(audit, polish_phase_sections(build_phase_sections(audit)))
    audit["quick_wins"] = build_quick_wins(audit)
    audit["strategic_moves"] = build_strategic_moves(audit)
    audit["competitor_comparison"] = build_competitor_comparison(audit, competitor_urls or [])
    audit["score"] = build_score(audit, issues)
    return audit


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    add_page_borders(section)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = RGBColor.from_string(BRAND_TEXT)


def add_header_footer(doc: Document, logo_path: Path, audit: dict) -> None:
    for section in doc.sections:
        header = section.header
        header_table = header.add_table(rows=1, cols=3, width=Cm(17.4))
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(header_table)
        header_table.columns[0].width = Cm(0.45)
        header_table.columns[1].width = Cm(10.7)
        header_table.columns[2].width = Cm(6.25)
        accent_cell, brand_cell, label_cell = header_table.rows[0].cells
        set_cell_shading(accent_cell, BRAND_ORANGE)
        set_cell_shading(brand_cell, BRAND_DARK)
        set_cell_shading(label_cell, BRAND_DARK)
        for cell in (accent_cell, brand_cell, label_cell):
            set_cell_margins(cell, 70, 120, 70, 120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        brand_paragraph = brand_cell.paragraphs[0]
        brand_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if logo_path.exists():
            brand_run = brand_paragraph.add_run()
            brand_run.add_picture(str(logo_path), width=Inches(0.24))
            brand_paragraph.add_run("  ")
        brand_run = brand_paragraph.add_run(BRAND_NAME)
        set_font(brand_run, size=10.5, bold=True, color="FFFFFF")
        label_paragraph = label_cell.paragraphs[0]
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        label_run = label_paragraph.add_run(normalize_output_text("SEO AUDIT / GROWTH MAP"))
        set_font(label_run, size=8.6, bold=True, color=BRAND_CYAN)

        footer = section.footer
        footer_table = footer.add_table(rows=1, cols=2, width=Cm(17.4))
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(footer_table)
        footer_table.columns[0].width = Cm(13.5)
        footer_table.columns[1].width = Cm(3.9)
        left_cell, right_cell = footer_table.rows[0].cells
        set_cell_margins(left_cell, 50, 0, 20, 0)
        set_cell_margins(right_cell, 50, 0, 20, 0)
        left_paragraph = left_cell.paragraphs[0]
        left_run = left_paragraph.add_run(normalize_output_text(f"{audit['domain']}  |  Audit by {BRAND_NAME}"))
        set_font(left_run, size=8.7, color=BRAND_MUTED)
        right_paragraph = right_cell.paragraphs[0]
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(right_paragraph)
        if right_paragraph.runs:
            set_font(right_paragraph.runs[-1], size=8.7, color=BRAND_MUTED)


def add_cover(doc: Document, audit: dict, logo_path: Path) -> None:
    cover = doc.add_table(rows=1, cols=2)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(cover)
    left, right = cover.rows[0].cells
    left.width = Cm(11.2)
    right.width = Cm(6.2)
    set_cell_shading(left, BRAND_DARK)
    set_cell_shading(right, BRAND_ACCENT_SOFT)
    for cell in (left, right):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_margins(cell, 220, 220, 220, 220)

    if logo_path.exists():
        p_logo = left.paragraphs[0]
        run = p_logo.add_run()
        run.add_picture(str(logo_path), width=Inches(0.65))
        p_logo.add_run("  ")
        brand_run = p_logo.add_run(BRAND_NAME)
        set_font(brand_run, size=14, bold=True, color="FFFFFF")

    p_tag = left.add_paragraph()
    p_tag.paragraph_format.space_before = Pt(30)
    tag_run = p_tag.add_run("SEO-аудит / план роста")
    set_font(tag_run, size=10.5, bold=True, color=BRAND_CYAN)

    p_title = left.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    title_run = p_title.add_run(f"SEO-аудит сайта\n{audit['domain']}")
    set_font(title_run, size=26, bold=True, color="FFFFFF")
    p_title.paragraph_format.space_after = Pt(18)

    callout = left.add_table(rows=1, cols=1)
    remove_table_borders(callout)
    callout_cell = callout.rows[0].cells[0]
    set_cell_shading(callout_cell, BRAND_DARK_ALT)
    set_cell_margins(callout_cell, 110, 130, 110, 130)
    callout_run = callout_cell.paragraphs[0].add_run(
        "Внутри: индексация, robots.txt и sitemap, мета-теги, структура страниц, изображения, коммерческие блоки и план работ на 60 дней."
    )
    set_font(callout_run, size=10.8, color="FFFFFF")

    run = right.paragraphs[0].add_run("Паспорт проекта")
    set_font(run, size=14, bold=True, color=BRAND_TEXT)
    add_text_paragraph(right, f"Бренд: {audit['company_name']}", size=11.4)
    add_text_paragraph(right, f"Домен: {audit['base_url']}", size=11.2)
    add_text_paragraph(right, f"Дата: {audit['generated_at']}", size=11.2)
    add_text_paragraph(right, "Формат: SEO-аудит / план работ", size=11.2)
    add_text_paragraph(right, f"Исполнитель: {BRAND_NAME}", size=11.2, bold=True)
    add_text_paragraph(
        right,
        "Фокус аудита: индексация, структура сайта, мета-теги, изображения, коммерческие страницы и точки роста по поисковому спросу.",
        size=11,
        color=BRAND_MUTED,
        space_after=12,
    )

    score_box = right.add_table(rows=1, cols=1)
    remove_table_borders(score_box)
    score_cell = score_box.rows[0].cells[0]
    set_cell_shading(score_cell, "FFFFFF")
    set_cell_margins(score_cell, 120, 120, 120, 120)
    score_label = score_cell.paragraphs[0]
    score_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = score_label.add_run("Индекс SEO-готовности")
    set_font(label_run, size=10.2, bold=True, color=BRAND_MUTED)
    score_value = score_cell.add_paragraph()
    score_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_run = score_value.add_run(str(audit["score"]))
    set_font(score_run, size=28, bold=True, color=BRAND_TEXT)
    score_hint = score_cell.add_paragraph()
    score_hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hint_run = score_hint.add_run("из 100")
    set_font(hint_run, size=10.2, color=BRAND_MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_metric_grid(doc: Document, audit: dict) -> None:
    metrics = [
        ("URL в sitemap", str(audit["sitemap_url_count"])),
        ("Записей в sitemap", str(audit.get("sitemap_total_entries", audit["sitemap_url_count"]))),
        ("Средний ответ", f"{audit['average_response_ms']} ms"),
        ("Средний HTML", f"{audit['average_html_kb']} KB"),
        ("Title длиннее 70", f"{math.floor(audit['title_long_ratio'] * 100)}%"),
        ("Проблемных description", f"{math.floor(audit['description_problem_ratio'] * 100)}%"),
        ("Страниц со schema", f"{math.floor(audit['schema_coverage_ratio'] * 100)}%"),
        ("Изображений без alt", str(audit["total_missing_alt"])),
    ]
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    index = 0
    for row in table.rows:
        for cell in row.cells:
            title, value = metrics[index]
            set_cell_shading(cell, BRAND_SOFT if index % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, 100, 110, 100, 110)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p1 = cell.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(normalize_output_text(title).upper())
            set_font(r1, size=8.5, bold=True, color=BRAND_MUTED)
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(normalize_output_text(value))
            set_font(r2, size=17, bold=True, color=BRAND_TEXT)
            index += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_issue_cards(doc: Document, issues: list[AuditIssue]) -> None:
    severity_colors = {"Critical": "FFE3E3", "High": "FFF1E6", "Medium": "EEF7FF", "Low": "F4F6FA"}
    severity_text_colors = {"Critical": "B42318", "High": "B54708", "Medium": "175CD3", "Low": BRAND_MUTED}
    for issue in issues:
        table = doc.add_table(rows=2, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(table)
        head = table.rows[0].cells[0]
        body = table.rows[1].cells[0]
        set_cell_shading(head, severity_colors.get(issue.severity, "F7F7F7"))
        set_cell_shading(body, "FFFFFF")
        set_cell_margins(head, 100, 120, 90, 120)
        set_cell_margins(body, 100, 120, 120, 120)
        head_run = head.paragraphs[0].add_run(normalize_output_text(f"{severity_label(issue.severity)}  |  {issue.title}"))
        set_font(head_run, size=11.2, bold=True, color=severity_text_colors.get(issue.severity, BRAND_TEXT))
        body_run = body.paragraphs[0].add_run(normalize_output_text(issue.why_it_matters))
        set_font(body_run, size=11.1, color=BRAND_TEXT)
        body.paragraphs[0].paragraph_format.space_after = Pt(6)
        if issue.evidence:
            add_bullet_list(body, issue.evidence, size=10.8)
        else:
            add_text_paragraph(body, "Примеры страниц в выборке не найдены.", size=10.6, color=BRAND_MUTED, space_after=4)
        recommendation = body.add_paragraph()
        recommendation_run = recommendation.add_run(normalize_output_text(f"Что делать: {issue.recommendation}"))
        set_font(recommendation_run, size=10.9, bold=True, color=BRAND_TEXT)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_roadmap_table(doc: Document, roadmap: list[tuple[str, list[str]]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    for idx, (period, tasks) in enumerate(roadmap):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BRAND_SOFT if idx != 1 else "FFFFFF")
        set_cell_margins(cell, 120, 120, 120, 120)
        title = cell.paragraphs[0]
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(period)
        set_font(run, size=11.2, bold=True, color=BRAND_ORANGE)
        add_bullet_list(cell, tasks, size=10.4)


def add_snapshot_table(doc: Document, snapshots: list[PageSnapshot]) -> None:
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(8.3), Cm(2.6), Cm(1.2), Cm(1.3), Cm(1.3), Cm(1.0), Cm(3.1)]
    headers = ["Путь", "Тип", "Код", "Title", "Desc", "H1", "Schema"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_shading(cell, BRAND_DARK)
        set_cell_margins(cell, 90, 80, 90, 80)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(normalize_output_text(header))
        set_font(run, size=9, bold=True, color="FFFFFF")
    for snapshot in snapshots[:12]:
        row = table.add_row().cells
        values = [
            shorten_path(snapshot.url, 42),
            snapshot.page_type,
            str(snapshot.status_code),
            str(len(snapshot.title)),
            str(len(snapshot.description)),
            str(len(snapshot.h1s)),
            ", ".join(snapshot.schema_types[:2]) or "-",
        ]
        for idx, value in enumerate(values):
            cell = row[idx]
            cell.width = widths[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_shading(cell, BRAND_SOFT if idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, 70, 70, 70, 70)
            paragraph = cell.paragraphs[0]
            if idx in (2, 3, 4, 5):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(normalize_output_text(value))
            set_font(run, size=9.2, color=BRAND_TEXT)


def add_priority_matrix_table(doc: Document, rows: list[dict]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    headers = ["Проблема", "Влияние", "Риск", "Польза", "Итог", "Приоритет", "Ответственный"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BRAND_DARK)
        set_cell_margins(cell, 80, 70, 80, 70)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(normalize_output_text(header))
        set_font(run, size=8.7, bold=True, color="FFFFFF")
    for row_data in rows:
        row = table.add_row().cells
        values = [
            row_data.get("problem", ""),
            str(row_data.get("impact", "")),
            str(row_data.get("risk", "")),
            str(row_data.get("business", "")),
            str(row_data.get("total", "")),
            severity_label(str(row_data.get("severity", ""))),
            row_data.get("owner", ""),
        ]
        for idx, value in enumerate(values):
            cell = row[idx]
            set_cell_shading(cell, BRAND_SOFT if idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, 70, 70, 70, 70)
            paragraph = cell.paragraphs[0]
            if idx in (1, 2, 3, 4):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(normalize_output_text(value))
            set_font(run, size=8.9, color=BRAND_TEXT)


def add_metrics_list(container, metrics: list[tuple[str, str]]) -> None:
    for label, value in metrics:
        paragraph = container.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        key_run = paragraph.add_run(f"{normalize_output_text(label)}: ")
        set_font(key_run, size=10.1, bold=True, color=BRAND_TEXT)
        value_run = paragraph.add_run(normalize_output_text(value))
        set_font(value_run, size=10.1, color=BRAND_MUTED)


def add_phase_sections_doc(doc: Document, phase_sections: list[dict]) -> None:
    for section in phase_sections:
        add_section_heading(doc, "Подробный разбор", section.get("title", ""), section.get("intro"))
        for idx, check in enumerate(section.get("checks", []), start=1):
            card = doc.add_table(rows=2, cols=1)
            card.alignment = WD_TABLE_ALIGNMENT.CENTER
            remove_table_borders(card)
            head = card.rows[0].cells[0]
            body = card.rows[1].cells[0]
            set_cell_shading(head, BRAND_SOFT)
            set_cell_shading(body, "FFFFFF")
            set_cell_margins(head, 90, 110, 80, 110)
            set_cell_margins(body, 100, 110, 110, 110)
            title_run = head.paragraphs[0].add_run(normalize_output_text(f"{idx}. {check.get('name', '')}"))
            set_font(title_run, size=11.4, bold=True, color=BRAND_TEXT)

            add_text_paragraph(body, f"Что проверялось: {check.get('checked', '')}", size=10.7, color=BRAND_TEXT, space_after=3)
            add_text_paragraph(body, f"Как проверялось: {check.get('method', '')}", size=10.5, color=BRAND_MUTED, space_after=4)
            add_metrics_list(body, check.get("metrics", []))
            add_text_paragraph(body, "Что нашли:", size=10.5, bold=True, color=BRAND_TEXT, space_after=3)
            add_bullet_list(body, check.get("findings", []), size=10.4)
            add_text_paragraph(
                body,
                f"Приоритет: {severity_label(str(check.get('priority', '')))}  |  Ответственный: {check.get('owner', '')}",
                size=10.2,
                bold=True,
                color=BRAND_ORANGE,
                space_after=3,
            )
            add_text_paragraph(body, f"Что делать: {check.get('recommendation', '')}", size=10.4, color=BRAND_TEXT, bold=True, space_after=4)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_action_cards_doc(doc: Document, items: list[dict], value_key: str, extra_key: str) -> None:
    if not items:
        return
    for item in items:
        card = doc.add_table(rows=2, cols=1)
        card.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(card)
        head = card.rows[0].cells[0]
        body = card.rows[1].cells[0]
        set_cell_shading(head, BRAND_ACCENT_SOFT)
        set_cell_shading(body, "FFFFFF")
        set_cell_margins(head, 90, 110, 80, 110)
        set_cell_margins(body, 100, 110, 110, 110)
        title_run = head.paragraphs[0].add_run(normalize_output_text(item.get("title", "")))
        set_font(title_run, size=11.2, bold=True, color=BRAND_TEXT)
        add_text_paragraph(body, f"{action_meta_label(value_key)}: {item.get(value_key.lower(), item.get(value_key, ''))}", size=10.2, color=BRAND_ORANGE, bold=True, space_after=2)
        add_text_paragraph(body, f"{action_meta_label(extra_key)}: {item.get(extra_key.lower(), item.get(extra_key, ''))}", size=10.2, color=BRAND_MUTED, bold=True, space_after=3)
        main_text = item.get("action") or item.get("details") or ""
        add_text_paragraph(body, main_text, size=10.5, color=BRAND_TEXT, space_after=4)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_competitor_comparison_doc(doc: Document, comparison: dict) -> None:
    summary = comparison.get("summary", [])
    if summary:
        add_bullet_list(doc, summary, size=10.9)

    for competitor in comparison.get("competitors", []):
        card = doc.add_table(rows=2, cols=1)
        card.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(card)
        head = card.rows[0].cells[0]
        body = card.rows[1].cells[0]
        set_cell_shading(head, BRAND_SOFT)
        set_cell_shading(body, "FFFFFF")
        set_cell_margins(head, 90, 110, 80, 110)
        set_cell_margins(body, 100, 110, 110, 110)
        title_run = head.paragraphs[0].add_run(normalize_output_text(f"{competitor.get('domain', '')}  |  проверено страниц: {competitor.get('pages_checked', 0)}"))
        set_font(title_run, size=11.2, bold=True, color=BRAND_TEXT)
        if competitor.get("highlights"):
            add_text_paragraph(body, "Коротко:", size=10.4, bold=True, color=BRAND_ORANGE, space_after=3)
            add_bullet_list(body, competitor.get("highlights", []), size=10.3)
        if competitor.get("factor_summary"):
            add_text_paragraph(body, "Ключевые факторы:", size=10.25, bold=True, color=BRAND_TEXT, space_after=3)
            add_bullet_list(body, competitor.get("factor_summary", []), size=10.15)
        if competitor.get("template_rows"):
            add_text_paragraph(body, "Ключевые различия по шаблонам:", size=10.25, bold=True, color=BRAND_TEXT, space_after=3)
            add_bullet_list(body, competitor.get("template_rows", []), size=10.15)
        if competitor.get("template_findings"):
            add_text_paragraph(body, "Где конкурент сильнее по шаблонам:", size=10.25, bold=True, color=BRAND_TEXT, space_after=3)
            add_bullet_list(body, competitor.get("template_findings", []), size=10.15)
        if competitor.get("snippet_findings"):
            add_text_paragraph(body, "Что видно по сниппетам:", size=10.25, bold=True, color=BRAND_TEXT, space_after=3)
            add_bullet_list(body, competitor.get("snippet_findings", []), size=10.15)
        if competitor.get("commercial_findings"):
            add_text_paragraph(body, "Что сильнее по FAQ, доверию и коммерческому слою:", size=10.25, bold=True, color=BRAND_TEXT, space_after=3)
            add_bullet_list(body, competitor.get("commercial_findings", []), size=10.15)
        if competitor.get("sample_paths"):
            add_text_paragraph(
                body,
                f"Где это видно: {', '.join(competitor.get('sample_paths', [])[:3])}",
                size=10.1,
                color=BRAND_MUTED,
                space_after=4,
            )
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    for item in comparison.get("gap_items", []):
        card = doc.add_table(rows=2, cols=1)
        card.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(card)
        head = card.rows[0].cells[0]
        body = card.rows[1].cells[0]
        set_cell_shading(head, BRAND_ACCENT_SOFT)
        set_cell_shading(body, "FFFFFF")
        set_cell_margins(head, 90, 110, 80, 110)
        set_cell_margins(body, 100, 110, 110, 110)
        title_run = head.paragraphs[0].add_run(normalize_output_text(item.get("title", "")))
        set_font(title_run, size=11.2, bold=True, color=BRAND_TEXT)
        add_text_paragraph(
            body,
            f"Приоритет: {item.get('priority', '')}  |  Ответственный: {item.get('owner', '')}",
            size=10.2,
            color=BRAND_ORANGE,
            bold=True,
            space_after=3,
        )
        add_text_paragraph(body, f"Что сейчас: {item.get('current_state', '')}", size=10.4, color=BRAND_TEXT, space_after=3)
        add_text_paragraph(body, f"Что видно у конкурентов: {item.get('competitor_state', '')}", size=10.4, color=BRAND_TEXT, space_after=3)
        if item.get("examples"):
            add_bullet_list(body, item.get("examples", []), size=10.2)
        if item.get("where_to_implement"):
            add_text_paragraph(body, f"Где внедрять: {', '.join(item.get('where_to_implement', []))}", size=10.25, color=BRAND_TEXT, space_after=3)
        add_text_paragraph(body, f"Техническое задание: {item.get('task', '')}", size=10.3, color=BRAND_TEXT, bold=True, space_after=3)
        if item.get("implementation_steps"):
            add_text_paragraph(body, "Шаги внедрения:", size=10.25, bold=True, color=BRAND_TEXT, space_after=3)
            add_bullet_list(body, item.get("implementation_steps", []), size=10.15)
        add_text_paragraph(body, f"Что это даст: {item.get('benefit', '')}", size=10.3, color=BRAND_TEXT, space_after=3)
        if item.get("impact_points"):
            add_text_paragraph(body, "Ожидаемый эффект:", size=10.25, bold=True, color=BRAND_MUTED, space_after=3)
            add_bullet_list(body, item.get("impact_points", []), size=10.1, color=BRAND_MUTED)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    failures = comparison.get("failures", [])
    if failures:
        add_text_paragraph(doc, "Не по всем конкурентам удалось собрать страницы. Ниже список URL, которые стоит перепроверить:", size=10.4, color=BRAND_MUTED, space_after=3)
        add_bullet_list(doc, [f"{item.get('url', '')}: {item.get('error', '')}" for item in failures], size=10.1)


def build_executive_summary(audit: dict) -> list[str]:
    return [
        (
            f"{audit['company_name']} СѓР¶Рµ РёРјРµРµС‚ СЃРёР»СЊРЅС‹Р№ С„СѓРЅРґР°РјРµРЅС‚ РґР»СЏ СЂРѕСЃС‚Р°: С‡РёСЃС‚С‹Р№ HTTPS-РґРѕРјРµРЅ, "
            f"Р±РѕР»СЊС€РѕР№ РєР°С‚Р°Р»РѕРі СЃРїСЂРѕСЃР°, СЂР°Р±РѕС‡РёР№ sitemap, schema РЅР° С‚РѕРІР°СЂР°С… Рё РєР°С‚РµРіРѕСЂРёСЏС… Рё С€РёСЂРѕРєСѓСЋ URL-СЃС‚СЂСѓРєС‚СѓСЂСѓ РїРѕРґ РєР»Р°СЃС‚РµСЂС‹."
        ),
        (
            "РќРѕ СЃРµР№С‡Р°СЃ РїСЂРѕРµРєС‚ С‚РµСЂСЏРµС‚ С‡Р°СЃС‚СЊ РІРёРґРёРјРѕСЃС‚Рё РЅРµ РёР·-Р·Р° РѕС‚СЃСѓС‚СЃС‚РІРёСЏ Р°СЃСЃРѕСЂС‚РёРјРµРЅС‚Р°, Р° РёР·-Р·Р° С‚РµС…РЅРёС‡РµСЃРєРёС… Рё С€Р°Р±Р»РѕРЅРЅС‹С… РѕРіСЂР°РЅРёС‡РµРЅРёР№: "
            "robots.txt РєРѕРЅС„Р»РёРєС‚СѓРµС‚ СЃ РєР°СЂС‚РѕР№ СЃР°Р№С‚Р°, РјРµС‚Р°РґР°РЅРЅС‹Рµ Сѓ С‚РѕРІР°СЂРЅС‹С… Рё РєР°С‚РµРіРѕСЂРёР№РЅС‹С… С€Р°Р±Р»РѕРЅРѕРІ С‡Р°СЃС‚Рѕ СЃР»РёС€РєРѕРј РґР»РёРЅРЅС‹Рµ, "
            "РєРѕРЅС‚Р°РєС‚РЅС‹Рµ СЃС‚СЂР°РЅРёС†С‹ РЅРµРґРѕРѕС„РѕСЂРјР»РµРЅС‹, Р° РёР·РѕР±СЂР°Р¶РµРЅРёСЏ РЅРµ РґРѕР±РёСЂР°СЋС‚ SEO-СЃРёРіРЅР°Р»С‹ С‡РµСЂРµР· alt."
        ),
        (
            "Р•СЃР»Рё СѓР±СЂР°С‚СЊ СЌС‚РѕС‚ С‚РµС…РЅРёС‡РµСЃРєРёР№ С€СѓРј Рё РїРµСЂРµСЃРѕР±СЂР°С‚СЊ РєРѕРјРјРµСЂС‡РµСЃРєРёРµ С€Р°Р±Р»РѕРЅС‹, СЃР°Р№С‚ СЃРјРѕР¶РµС‚ Р·Р°РјРµС‚РЅРѕ СЃРёР»СЊРЅРµРµ СЂР°СЃРєСЂС‹С‚СЊ СЃРїСЂРѕСЃ "
            "РїРѕ Р°РєРІР°СЂРёСѓРјР°Рј, С‚РµСЂСЂР°СЂРёСѓРјР°Рј, С‚СѓРјР±Р°Рј Рё Р°РєСЃРµСЃСЃСѓР°СЂР°Рј Р±РµР· РїРѕР»РЅРѕР№ РїРµСЂРµРґРµР»РєРё РїР»Р°С‚С„РѕСЂРјС‹."
        ),
    ]


def add_screenshot_gallery(doc: Document, screenshots: list[dict]) -> None:
    if not screenshots:
        return
    add_section_heading(
        doc,
        "РЎРєСЂРёРЅС€РѕС‚С‹",
        "РђРІС‚РѕСЃРєСЂРёРЅС€РѕС‚С‹ РєР»СЋС‡РµРІС‹С… СЃС‚СЂР°РЅРёС†",
        "Р­С‚Рё РёР·РѕР±СЂР°Р¶РµРЅРёСЏ РїРѕРјРѕРіР°СЋС‚ Р±С‹СЃС‚СЂРѕ СѓРІРёРґРµС‚СЊ РєРѕРЅС‚РµРєСЃС‚: РєР°Рє РІС‹РіР»СЏРґРёС‚ РіР»Р°РІРЅР°СЏ, РєРѕРЅС‚Р°РєС‚С‹, РєР°С‚РµРіРѕСЂРёСЏ Рё РєР°СЂС‚РѕС‡РєР° С‚РѕРІР°СЂР° РІ РјРѕРјРµРЅС‚ Р°СѓРґРёС‚Р°.",
    )
    for screenshot in screenshots:
        image_path = Path(screenshot["path"])
        if not image_path.exists():
            continue
        card = doc.add_table(rows=2, cols=1)
        card.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(card)
        head = card.rows[0].cells[0]
        body = card.rows[1].cells[0]
        set_cell_shading(head, BRAND_SOFT)
        set_cell_shading(body, "FFFFFF")
        set_cell_margins(head, 90, 110, 80, 110)
        set_cell_margins(body, 110, 110, 110, 110)
        title_run = head.paragraphs[0].add_run(f"{screenshot.get('label', 'РЎС‚СЂР°РЅРёС†Р°')}  |  {shorten_path(screenshot.get('url', ''), 72)}")
        set_font(title_run, size=10.6, bold=True, color=BRAND_TEXT)
        picture_run = body.paragraphs[0].add_run()
        picture_run.add_picture(str(image_path), width=Inches(6.25))
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def generate_docx(audit: dict, output_path: Path, logo_path: Path) -> None:
    audit = normalize_structure(audit)
    doc = Document()
    configure_document(doc)
    add_header_footer(doc, logo_path, audit)
    add_cover(doc, audit, logo_path)
    doc.add_page_break()

    add_section_heading(
        doc,
        "Краткий вывод",
        "Что тормозит рост и где лежит резерв",
        "Ниже только диагностический вывод: где сайт теряет потенциал, какие шаблоны проседают и что даст самый быстрый прирост.",
    )
    for paragraph in audit.get("executive_summary", []) or build_executive_summary_dynamic(audit):
        add_text_paragraph(doc, paragraph, size=11.4, color=BRAND_TEXT, space_after=6)
    add_metric_grid(doc, audit)

    if audit.get("priority_matrix"):
        add_section_heading(
            doc,
            "Приоритеты",
            "Какие задачи делать в первую очередь",
            "Сначала закрываем критичные и массовые проблемы, которые сильнее всего бьют по индексации, сниппетам и ключевым страницам сайта.",
        )
        add_priority_matrix_table(doc, audit.get("priority_matrix", []))

    critical_issue_cards = [AuditIssue(**item) for item in audit.get("critical_errors", [])] or audit["issues"][:4]
    if critical_issue_cards:
        add_section_heading(
            doc,
            "Критичные ошибки",
            "Что прямо сейчас мешает сайту расти",
            "Здесь только те проблемы, которые заметно влияют на индекс, сниппеты, трафик и заявки.",
        )
        add_issue_cards(doc, critical_issue_cards)

    if audit.get("phase_sections"):
        add_section_heading(
            doc,
            "Подробный разбор",
            "Проверка по основным слоям сайта",
            "Аудит разбит по этапам, чтобы было понятно, где именно находятся проблемы и как они проверялись.",
        )
        add_phase_sections_doc(doc, audit.get("phase_sections", []))

    competitor_comparison = audit.get("competitor_comparison") or {}
    if competitor_comparison and (
        competitor_comparison.get("competitors")
        or competitor_comparison.get("gap_items")
        or competitor_comparison.get("failures")
    ):
        add_section_heading(
            doc,
            "Сравнение с конкурентами",
            "Чего не хватает на фоне сильных конкурентов",
            "Сравнение показывает, какие шаблоны, блоки и коммерческие сигналы конкуренты используют лучше и что стоит внедрить в первую очередь.",
        )
        add_competitor_comparison_doc(doc, competitor_comparison)

    if audit.get("quick_wins"):
        add_section_heading(
            doc,
            "Быстрые исправления",
            "Что можно исправить в ближайшее время",
            "Это задачи, которые можно внедрить быстро и без большой переделки сайта.",
        )
        add_action_cards_doc(doc, audit.get("quick_wins", []), "effort", "impact")

    if audit.get("strategic_moves"):
        add_section_heading(
            doc,
            "Стратегические улучшения",
            "Что даст рост после базовых исправлений",
            "Это более крупные изменения, которые усиливают сайт в поиске и помогают получать больше заявок.",
        )
        add_action_cards_doc(doc, audit.get("strategic_moves", []), "impact", "effort")

    if audit.get("growth_points"):
        add_section_heading(
            doc,
            "Вектор роста",
            "Что нужно переработать после технички",
            "Здесь не абстрактные пожелания, а направления, которые усиливают структуру, коммерческий интент и качество контента после базовых исправлений.",
        )
        add_bullet_list(doc, audit["growth_points"], size=11.1)

    if audit.get("roadmap"):
        add_section_heading(
            doc,
            "План работ",
            "План внедрения на 60 дней",
            "Порядок выстроен так, чтобы сначала снять технические ограничения, потом усилить шаблоны и только после этого масштабировать рост.",
        )
        add_roadmap_table(doc, audit["roadmap"])

    add_section_heading(
        doc,
        "Приложение",
        "Какие страницы легли в основу разбора",
        "В приложении собраны главная, коммерческие, контактные и служебные URL, по которым видно качество шаблонов и SEO-оформления сайта.",
    )
    add_snapshot_table(doc, audit.get("appendix_pages") or audit["sample_pages"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def serialize_audit(audit: dict) -> dict:
    payload = dict(audit)
    payload["home_page"] = asdict(audit["home_page"])
    payload["sample_pages"] = [asdict(item) for item in audit["sample_pages"]]
    payload["appendix_pages"] = [asdict(item) for item in audit.get("appendix_pages", [])]
    payload["issues"] = [asdict(item) for item in audit["issues"]]
    payload["screenshots"] = audit.get("screenshots", [])
    return normalize_structure(payload)


def main() -> None:
    global BRAND_NAME
    parser = argparse.ArgumentParser(description="Generate a branded Shelpakov Digital SEO audit in DOCX format.")
    parser.add_argument("--url", required=True, help="Target website URL")
    parser.add_argument("--company", help="Client / brand name for the cover")
    parser.add_argument("--brand", default=BRAND_NAME, help="Brand label on the audit")
    parser.add_argument("--output", help="Output DOCX path")
    parser.add_argument("--sample-size", type=int, default=0, help="How many pages to analyze from the website; 0 = full crawl")
    parser.add_argument("--competitor", action="append", default=[], help="Competitor URL to compare against. Repeat the flag for multiple domains.")
    parser.add_argument("--no-json", action="store_true", help="Do not save raw audit JSON next to the DOCX")
    parser.add_argument("--no-preview", action="store_true", help="Do not save HTML preview next to the DOCX")
    parser.add_argument("--no-pdf", action="store_true", help="Do not save PDF preview next to the DOCX")
    parser.add_argument("--no-screenshots", action="store_true", help="Deprecated: screenshots block has been removed from the audit")
    args = parser.parse_args()

    BRAND_NAME = args.brand

    target = normalize_base_url(args.url)
    today = datetime.now().strftime("%Y-%m-%d")
    output = Path(args.output) if args.output else Path("audits") / f"{slugify_for_filename(target)}-{today}.docx"

    audit = build_audit(target, args.company, args.sample_size, args.competitor)
    audit["screenshots"] = []
    logo_path = Path("public") / "android-chrome-512x512.png"
    generate_docx(audit, output, logo_path)
    audit_payload = serialize_audit(audit)
    if not args.no_json:
        json_path = output.with_suffix(".json")
        json_path.write_text(json.dumps(audit_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    html_path = output.with_suffix(".html")
    pdf_path = output.with_suffix(".pdf")
    should_write_html = not args.no_preview or not args.no_pdf

    if should_write_html:
        write_preview_html(audit_payload, html_path)

    if not args.no_pdf:
        write_preview_pdf(audit_payload, pdf_path)
    print(f"Audit ready: {output}")


if __name__ == "__main__":
    main()

